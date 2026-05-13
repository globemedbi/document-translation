import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_02.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p02.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key):
    return lookup.get(key, "")

def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_run_with_format(para, text, bold=False, italic=False, font_size=10,
                        color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def add_job_header(doc, title_key, employer_key, period_key, location_key):
    title = get_value(title_key)
    employer = get_value(employer_key)
    period = get_value(period_key)
    location = get_value(location_key)

    table = doc.add_table(rows=1, cols=2)
    set_table_borders_none(table)
    table.autofit = False

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_width(left_cell, 12.0)
    set_cell_width(right_cell, 5.0)

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(left_para, before=0, after=0)

    add_run_with_format(left_para, title, bold=True, font_size=11, color=(26, 43, 76))
    add_run_with_format(left_para, "  |  ", bold=False, font_size=11, color=(180, 180, 180))
    add_run_with_format(left_para, employer, bold=False, font_size=11, color=(26, 43, 76))

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(right_para, before=0, after=0)
    add_run_with_format(right_para, period, bold=True, font_size=10, color=(26, 43, 76))

    loc_para = left_cell.add_paragraph()
    loc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(loc_para, before=0, after=0)
    add_run_with_format(loc_para, location, italic=True, font_size=9, color=(120, 120, 120))

    return table

def add_bullet(doc, text, font_size=9.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before=0, after=20, line=276)

    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)

    bullet_run = para.add_run("\u2022  ")
    bullet_run.font.size = Pt(font_size)
    bullet_run.font.color.rgb = RGBColor(0, 0, 0)

    text_run = para.add_run(text)
    text_run.font.size = Pt(font_size)
    text_run.font.color.rgb = RGBColor(0, 0, 0)

    return para

def add_spacer(doc, space_before=100):
    para = doc.add_paragraph()
    set_para_spacing(para, before=space_before, after=0)
    return para

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Job 1: AI Engineer | Inheaden ──────────────────────────────────────────
add_job_header(doc,
               "Intitulé du poste 1",
               "Employeur 1",
               "Période 1",
               "Lieu 1")

add_bullet(doc, get_value("Point 1.1"))
add_bullet(doc, get_value("Point 1.2"))
add_bullet(doc, get_value("Point 1.3"))
add_bullet(doc, get_value("Point 1.4"))

add_spacer(doc, 120)

# ── Job 2: Machine Learning Engineer | GlobeMed ───────────────────────────
add_job_header(doc,
               "Intitulé du poste 2",
               "Employeur 2",
               "Période 2",
               "Lieu 2")

add_bullet(doc, get_value("Point 2.1"))
add_bullet(doc, get_value("Point 2.2"))
add_bullet(doc, get_value("Point 2.3"))

add_spacer(doc, 120)

# ── Job 3: AI Engineer | Augment Digital ──────────────────────────────────
add_job_header(doc,
               "Intitulé du poste 3",
               "Employeur 3",
               "Période 3",
               "Lieu 3")

add_bullet(doc, get_value("Point 3.1"))
add_bullet(doc, get_value("Point 3.2"))

add_spacer(doc, 120)

# ── Job 4: Machine Learning Engineer Intern | Quantmetry ─────────────────
add_job_header(doc,
               "Intitulé du poste 4",
               "Employeur 4",
               "Période 4",
               "Lieu 4")

add_bullet(doc, get_value("Point 4.1"))
add_bullet(doc, get_value("Point 4.2"))
add_bullet(doc, get_value("Point 4.3"))

# ── Page number ───────────────────────────────────────────────────────────
page_para = doc.add_paragraph()
page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(page_para, before=200, after=0)
run_page_label = page_para.add_run("Page ")
run_page_label.font.size = Pt(8)
run_page_label.font.color.rgb = RGBColor(150, 150, 150)

fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")

run_num = page_para.add_run()
run_num.font.size = Pt(8)
run_num.font.color.rgb = RGBColor(150, 150, 150)
run_num._r.append(fldChar1)
run_num._r.append(instrText)
run_num._r.append(fldChar2)

doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")