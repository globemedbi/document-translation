import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p02.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_02.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()

s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def para_in_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def merge_row_cells(table, row_idx):
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[len(row.cells)-1])

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def no_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_section_header(text):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(18)
    no_table_border(t)
    r = t.rows[0]
    c = r.cells[0]
    set_cell_bg(c, "2E5B8A")
    para_in_cell(c, text, bold=True, size=9, color=(255, 255, 255))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ── HEADER ──────────────────────────────────────────────────────────────────
t = doc.add_table(rows=1, cols=2)
t.width = Cm(18)
no_table_border(t)

left_c = t.rows[0].cells[0]
left_c.width = Cm(11)
remove_cell_borders(left_c)

logo_p = left_c.paragraphs[0]
logo_p.paragraph_format.space_before = Pt(0)
logo_p.paragraph_format.space_after = Pt(2)
lr = logo_p.add_run("MC  ")
lr.bold = True
lr.font.size = Pt(14)
lr.font.color.rgb = RGBColor(255, 255, 255)

title_p = left_c.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(1)
tr = title_p.add_run(lookup["Formulaire de demande de remboursement médical"])
tr.bold = True
tr.font.size = Pt(13)
tr.font.color.rgb = RGBColor(30, 60, 120)

sub_p = left_c.add_paragraph()
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(0)
sr = sub_p.add_run("Soins primaires – à compléter par le membre et le médecin traitant")
sr.italic = True
sr.font.size = Pt(8)
sr.font.color.rgb = RGBColor(80, 80, 80)

right_c = t.rows[0].cells[1]
right_c.width = Cm(7)
remove_cell_borders(right_c)

info_t = doc.add_table(rows=4, cols=2)
info_t.width = Cm(7)
set_table_border(info_t)

info_rows = [
    ("N° de demande", lookup["N° de demande"]),
    ("Réf. d'approbation", lookup["Réf. d'approbation"]),
    ("Date du traitement", lookup["Date du traitement"]),
    ("Type", lookup["Type"]),
]
for i, (lbl, val) in enumerate(info_rows):
    rc = info_t.rows[i]
    lc = rc.cells[0]
    vc = rc.cells[1]
    lc.width = Cm(3.5)
    vc.width = Cm(3.5)
    para_in_cell(lc, lbl, bold=False, size=8)
    para_in_cell(vc, val, bold=True, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── SECTION A ────────────────────────────────────────────────────────────────
add_section_header(lookup["SECTION A - DONNÉES DU MEMBRE ASSURÉ"])

tA = doc.add_table(rows=5, cols=2)
tA.width = Cm(18)
set_table_border(tA)

def fill_a(row_idx, lbl1, val1, lbl2, val2):
    r = tA.rows[row_idx]
    lc = r.cells[0]
    rc = r.cells[1]
    lc.width = Cm(9)
    rc.width = Cm(9)
    remove_cell_borders(lc)
    remove_cell_borders(rc)
    p1 = lc.paragraphs[0]
    p1.paragraph_format.space_before = Pt(1)
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run(lbl1)
    run1.bold = False
    run1.font.size = Pt(7)
    run1.font.color.rgb = RGBColor(80, 80, 80)
    p1v = lc.add_paragraph()
    p1v.paragraph_format.space_before = Pt(0)
    p1v.paragraph_format.space_after = Pt(2)
    r1v = p1v.add_run(val1)
    r1v.bold = True
    r1v.font.size = Pt(9)

    p2 = rc.paragraphs[0]
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(lbl2)
    run2.bold = False
    run2.font.size = Pt(7)
    run2.font.color.rgb = RGBColor(80, 80, 80)
    p2v = rc.add_paragraph()
    p2v.paragraph_format.space_before = Pt(0)
    p2v.paragraph_format.space_after = Pt(2)
    r2v = p2v.add_run(val2)
    r2v.bold = True
    r2v.font.size = Pt(9)

fill_a(0, "NOM DU MEMBRE / PATIENT", lookup["NOM DU MEMBRE / PATIENT"], "NUMÉRO DE MEMBRE / CARTE", lookup["NUMÉRO DE MEMBRE / CARTE"])

r1 = tA.rows[1]
lc1 = r1.cells[0]
rc1 = r1.cells[1]
lc1.width = Cm(9)
rc1.width = Cm(9)
remove_cell_borders(lc1)
remove_cell_borders(rc1)

p_dob_lbl = lc1.paragraphs[0]
p_dob_lbl.paragraph_format.space_before = Pt(1)
p_dob_lbl.paragraph_format.space_after = Pt(0)
r_dob = p_dob_lbl.add_run("DATE DE NAISSANCE")
r_dob.font.size = Pt(7)
r_dob.font.color.rgb = RGBColor(80, 80, 80)
p_dob_val = lc1.add_paragraph()
p_dob_val.paragraph_format.space_before = Pt(0)
p_dob_val.paragraph_format.space_after = Pt(2)
r_dob_v = p_dob_val.add_run(lookup["DATE DE NAISSANCE"])
r_dob_v.bold = True
r_dob_v.font.size = Pt(9)

p_sex_lbl = rc1.paragraphs[0]
p_sex_lbl.paragraph_format.space_before = Pt(1)
p_sex_lbl.paragraph_format.space_after = Pt(0)
r_sex_lbl = p_sex_lbl.add_run("SEXE")
r_sex_lbl.font.size = Pt(7)
r_sex_lbl.font.color.rgb = RGBColor(80, 80, 80)

m_val = "☐" if "unchecked" in lookup["SEXE - M"] else "☑"
f_val = "☑" if "checked" in lookup["SEXE - F"] else "☐"

p_sex_m = rc1.add_paragraph()
p_sex_m.paragraph_format.space_before = Pt(1)
p_sex_m.paragraph_format.space_after = Pt(0)
r_sex_m = p_sex_m.add_run(f"{m_val} M")
r_sex_m.font.size = Pt(9)

p_sex_f = rc1.add_paragraph()
p_sex_f.paragraph_format.space_before = Pt(1)
p_sex_f.paragraph_format.space_after = Pt(2)
r_sex_f = p_sex_f.add_run(f"{f_val} F")
r_sex_f.font.size = Pt(9)

fill_a(2, "COMPAGNIE D'ASSURANCE / TPA", lookup["COMPAGNIE D'ASSURANCE / TPA"], "TITULAIRE DE LA POLICE", lookup["TITULAIRE DE LA POLICE"])
fill_a(3, "CPR / PASSEPORT", lookup["CPR / PASSEPORT"], "TÉLÉPHONE DU MEMBRE", lookup["TÉLÉPHONE DU MEMBRE"])

r4 = tA.rows[4]
tA.rows[4].cells[0].merge(tA.rows[4].cells[1])
remove_cell_borders(tA.rows[4].cells[0])
tA.rows[4].cells[0].paragraphs[0].paragraph_format.space_after = Pt(1)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── SECTION B ────────────────────────────────────────────────────────────────
add_section_header(lookup["SECTION B - INFORMATIONS MÉDICALES"])

tB = doc.add_table(rows=1, cols=2)
tB.width = Cm(18)
no_table_border(tB)

left_b = tB.rows[0].cells[0]
left_b.width = Cm(9)
right_b = tB.rows[0].cells[1]
right_b.width = Cm(9)
remove_cell_borders(left_b)
remove_cell_borders(right_b)

p_ts = left_b.paragraphs[0]
p_ts.paragraph_format.space_before = Pt(2)
p_ts.paragraph_format.space_after = Pt(3)
r_ts = p_ts.add_run(lookup["Type de demande"])
r_ts.bold = True
r_ts.font.size = Pt(9)

hosp_val = "☐" if "unchecked" in lookup["Hospitalisation"] else "☑"
amb_val = "☑" if "checked" in lookup["Ambulatoire"] else "☐"
urg_val = "☐" if "unchecked" in lookup["Urgence"] else "☑"

for cb_text in [f"{hosp_val} Hospitalisation", f"{amb_val} Ambulatoire", f"{urg_val} Urgence"]:
    p_cb = left_b.add_paragraph()
    p_cb.paragraph_format.space_before = Pt(2)
    p_cb.paragraph_format.space_after = Pt(2)
    r_cb = p_cb.add_run(cb_text)
    r_cb.font.size = Pt(9)

details_data = [
    ("Champ", "Détails"),
    ("Date du traitement", lookup["Date du traitement"]),
    ("Fournisseur / clinique", lookup["Fournisseur / clinique"]),
    ("Médecin traitant", lookup["Médecin traitant"]),
    ("Numéro de dossier médical", lookup["Numéro de dossier médical"]),
]

tDet = doc.add_table(rows=5, cols=2)
tDet.width = Cm(8.5)
set_table_border(tDet)

for i, (k, v) in enumerate(details_data):
    rc_d = tDet.rows[i]
    c0 = rc_d.cells[0]
    c1 = rc_d.cells[1]
    c0.width = Cm(4)
    c1.width = Cm(4.5)
    if i == 0:
        para_in_cell(c0, k, bold=True, size=8)
        para_in_cell(c1, v, bold=True, size=8)
        set_cell_bg(c0, "EEEEEE")
        set_cell_bg(c1, "EEEEEE")
    else:
        para_in_cell(c0, k, bold=False, size=8)
        para_in_cell(c1, v, bold=False, size=9)

left_b.add_paragraph()
from docx.oxml import OxmlElement as OE
left_b._tc.append(tDet._tbl)

p_cv = right_b.paragraphs[0]
p_cv.paragraph_format.space_before = Pt(2)
p_cv.paragraph_format.space_after = Pt(3)
r_cv = p_cv.add_run(lookup["Conditions et signes vitaux"])
r_cv.bold = True
r_cv.font.size = Pt(9)

cond_data = [
    ("Affection préexistante", "Affection préexistante - Oui", "Affection préexistante - Non"),
    ("Accident du travail", "Accident du travail - Oui", "Accident du travail - Non"),
    ("Maternité", "Maternité - Oui", "Maternité - Non"),
]

tCond = doc.add_table(rows=4, cols=3)
tCond.width = Cm(8.5)
set_table_border(tCond)

hdr_row = tCond.rows[0]
set_cell_bg(hdr_row.cells[0], "EEEEEE")
set_cell_bg(hdr_row.cells[1], "EEEEEE")
set_cell_bg(hdr_row.cells[2], "EEEEEE")
para_in_cell(hdr_row.cells[0], "Condition", bold=True, size=8)
para_in_cell(hdr_row.cells[1], "Oui", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
para_in_cell(hdr_row.cells[2], "Non", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (label, oui_key, non_key) in enumerate(cond_data):
    oui_val = "☐" if "unchecked" in lookup[oui_key] else "☑"
    non_val = "☑" if "checked" in lookup[non_key] else "☐"
    cr = tCond.rows[i + 1]
    cr.cells[0].width = Cm(4.5)
    cr.cells[1].width = Cm(2)
    cr.cells[2].width = Cm(2)
    para_in_cell(cr.cells[0], label, bold=False, size=8)
    para_in_cell(cr.cells[1], oui_val, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    para_in_cell(cr.cells[2], non_val, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

p_temp_lbl = right_b.add_paragraph()
p_temp_lbl.paragraph_format.space_before = Pt(4)
p_temp_lbl.paragraph_format.space_after = Pt(0)
r_temp_lbl = p_temp_lbl.add_run("Tension / pouls / température")
r_temp_lbl.font.size = Pt(8)
r_temp_lbl.font.color.rgb = RGBColor(80, 80, 80)

p_temp_val = right_b.add_paragraph()
p_temp_val.paragraph_format.space_before = Pt(0)
p_temp_val.paragraph_format.space_after = Pt(2)
r_temp_val = p_temp_val.add_run(lookup["Tension / pouls / température"])
r_temp_val.font.size = Pt(9)

right_b._tc.append(tCond._tbl)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

tMD = doc.add_table(rows=1, cols=2)
tMD.width = Cm(18)
set_table_border(tMD)
tMD.rows[0].cells[0].width = Cm(9)
tMD.rows[0].cells[1].width = Cm(9)

mc = tMD.rows[0].cells[0]
dc = tMD.rows[0].cells[1]

p_motif_lbl = mc.paragraphs[0]
p_motif_lbl.paragraph_format.space_before = Pt(1)
p_motif_lbl.paragraph_format.space_after = Pt(2)
r_motif_lbl = p_motif_lbl.add_run("MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS")
r_motif_lbl.bold = True
r_motif_lbl.font.size = Pt(8)

p_motif_val = mc.add_paragraph()
p_motif_val.paragraph_format.space_before = Pt(0)
p_motif_val.paragraph_format.space_after = Pt(2)
r_motif_val = p_motif_val.add_run(lookup["MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS"])
r_motif_val.font.size = Pt(8)

p_diag_lbl = dc.paragraphs[0]
p_diag_lbl.paragraph_format.space_before = Pt(1)
p_diag_lbl.paragraph_format.space_after = Pt(2)
r_diag_lbl = p_diag_lbl.add_run("DIAGNOSTIC FINAL")
r_diag_lbl.bold = True
r_diag_lbl.font.size = Pt(8)

p_diag_val = dc.add_paragraph()
p_diag_val.paragraph_format.space_before = Pt(0)
p_diag_val.paragraph_format.space_after = Pt(2)
r_diag_val = p_diag_val.add_run(lookup["DIAGNOSTIC FINAL"])
r_diag_val.font.size = Pt(8)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── SECTION C ────────────────────────────────────────────────────────────────
add_section_header(lookup["SECTION C - TRAITEMENT, INVESTIGATIONS ET COÛT ESTIMÉ"])

tC = doc.add_table(rows=5, cols=5)
tC.width = Cm(18)
set_table_border(tC)

col_widths = [5.5, 3.5, 1.5, 2.5, 5.0]
headers_c = ["Traitement / service", "Catégorie", "Qté", "Coût brut", "Justification médicale"]

hdr = tC.rows[0]
for i, (h, w) in enumerate(zip(headers_c, col_widths)):
    hdr.cells[i].width = Cm(w)
    set_cell_bg(hdr.cells[i], "EEEEEE")
    para_in_cell(hdr.cells[i], h, bold=True, size=8)

service_rows = [
    ("Consultation ORL", "Traitement / service - Consultation ORL"),
    ("Injection Lincocin", "Traitement / service - Injection Lincocin"),
    ("Injection Depo-Medrol", "Traitement / service - Injection Depo-Medrol"),
]

for i, (svc_name, key) in enumerate(service_rows):
    val = lookup[key]
    parts = [p.strip() for p in val.split("|")]
    cat = parts[0].split(":", 1)[1].strip() if ":" in parts[0] else parts[0]
    qty = parts[1].split(":", 1)[1].strip() if len(parts) > 1 and ":" in parts[1] else "1"
    cost = parts[2].split(":", 1)[1].strip() if len(parts) > 2 and ":" in parts[2] else ""
    just = parts[3].split(":", 1)[1].strip() if len(parts) > 3 and ":" in parts[3] else ""
    r = tC.rows[i + 1]
    for j, w in enumerate(col_widths):
        r.cells[j].width = Cm(w)
    para_in_cell(r.cells[0], svc_name, bold=False, size=9)
    para_in_cell(r.cells[1], cat, bold=False, size=9)
    para_in_cell(r.cells[2], qty, bold=False, size=9)
    para_in_cell(r.cells[3], cost, bold=False, size=9)
    para_in_cell(r.cells[4], just, bold=False, size=9)

total_parts = lookup["Coût prévu / total brut"].split("|")
total_cost = total_parts[0].strip()
total_note = total_parts[1].strip() if len(total_parts) > 1 else ""

tr_last = tC.rows[4]
for j, w in enumerate(col_widths):
    tr_last.cells[j].width = Cm(w)

tr_last.cells[0].merge(tr_last.cells[2])
para_in_cell(tr_last.cells[0], "Coût prévu / total brut", bold=True, size=9)
para_in_cell(tr_last.cells[3], total_cost, bold=True, size=9)
para_in_cell(tr_last.cells[4], total_note, bold=True, size=8)

doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── DECLARATIONS ─────────────────────────────────────────────────────────────
tDecl = doc.add_table(rows=1, cols=2)
tDecl.width = Cm(18)
no_table_border(tDecl)

dcl = tDecl.rows[0].cells[0]
dcr = tDecl.rows[0].cells[1]
dcl.width = Cm(9)
dcr.width = Cm(9)
remove_cell_borders(dcl)
remove_cell_borders(dcr)

p_dl = dcl.paragraphs[0]
p_dl.paragraph_format.space_before = Pt(1)
p_dl.paragraph_format.space_after = Pt(2)
r_dl = p_dl.add_run("Déclaration du membre")
r_dl.bold = True
r_dl.font.size = Pt(9)

p_dlv = dcl.add_paragraph()
p_dlv.paragraph_format.space_before = Pt(0)
p_dlv.paragraph_format.space_after = Pt(6)
r_dlv = p_dlv.add_run(lookup["Déclaration du membre"])
r_dlv.font.size = Pt(8)

p_sig_l = dcl.add_paragraph()
p_sig_l.paragraph_format.space_before = Pt(10)
p_sig_l.paragraph_format.space_after = Pt(0)
r_sig_l = p_sig_l.add_run("_" * 40)
r_sig_l.font.size = Pt(9)

p_sig_ll = dcl.add_paragraph()
p_sig_ll.paragraph_format.space_before = Pt(1)
p_sig_ll.paragraph_format.space_after = Pt(2)
r_sig_ll = p_sig_ll.add_run("Signature du membre / date")
r_sig_ll.font.size = Pt(8)
r_sig_ll.font.color.rgb = RGBColor(120, 120, 120)

p_dr = dcr.paragraphs[0]
p_dr.paragraph_format.space_before = Pt(1)
p_dr.paragraph_format.space_after = Pt(2)
r_dr = p_dr.add_run("Déclaration du prestataire médical")
r_dr.bold = True
r_dr.font.size = Pt(9)

p_drv = dcr.add_paragraph()
p_drv.paragraph_format.space_before = Pt(0)
p_drv.paragraph_format.space_after = Pt(6)
r_drv = p_drv.add_run(lookup["Déclaration du prestataire médical"])
r_drv.font.size = Pt(8)

p_sig_r = dcr.add_paragraph()
p_sig_r.paragraph_format.space_before = Pt(10)
p_sig_r.paragraph_format.space_after = Pt(0)
r_sig_r = p_sig_r.add_run("_" * 40)
r_sig_r.font.size = Pt(9)

p_sig_rl = dcr.add_paragraph()
p_sig_rl.paragraph_format.space_before = Pt(1)
p_sig_rl.paragraph_format.space_after = Pt(2)
r_sig_rl = p_sig_rl.add_run("Nom du médecin, signature et cachet")
r_sig_rl.font.size = Pt(8)
r_sig_rl.font.color.rgb = RGBColor(120, 120, 120)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)