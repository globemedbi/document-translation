import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p02.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_02.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def cell_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def no_space_para(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

# HEADER BLOCK
ht = doc.add_table(rows=1, cols=2)
ht.width = Cm(18)
from docx.shared import Inches
ht.columns[0].width = Cm(12)
ht.columns[1].width = Cm(6)

left_cell = ht.cell(0, 0)
left_cell.text = ""
lp1 = left_cell.paragraphs[0]
lp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = lp1.add_run("MC  ")
r1.bold = True
r1.font.size = Pt(16)
r1.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
r2 = lp1.add_run(lookup["Formulaire de demande de remboursement médical"])
r2.bold = True
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

lp2 = left_cell.add_paragraph("Soins primaires - à compléter par le membre et le médecin traitant")
lp2.runs[0].font.size = Pt(8)
lp2.runs[0].italic = True
lp2.paragraph_format.space_before = Pt(2)

right_cell = ht.cell(0, 1)
right_cell.text = ""
meta_t = right_cell.tables[0] if right_cell.tables else None

inner = right_cell.add_table(rows=4, cols=2)
inner.width = Cm(6)
set_table_border(inner)
meta_labels = ["N° de demande", "Réf. d'approbation", "Date du traitement", "Type"]
for i, lbl in enumerate(meta_labels):
    cell_para(inner.rows[i].cells[0], lbl, bold=False, size=8)
    cell_para(inner.rows[i].cells[1], lookup[lbl], bold=True, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# SECTION A
sa_t = doc.add_table(rows=1, cols=1)
sa_t.width = Cm(18)
set_cell_bg(sa_t.cell(0, 0), "2E5E99")
cell_para(sa_t.cell(0, 0), lookup["SECTION A - DONNÉES DU MEMBRE ASSURÉ"], bold=True, size=10)
sa_t.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

a1 = doc.add_table(rows=1, cols=2)
a1.width = Cm(18)
set_table_border(a1)
c0 = a1.cell(0, 0)
c0.text = ""
p = c0.paragraphs[0]
r = p.add_run("NOM DU MEMBRE / PATIENT")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c0.add_paragraph(lookup["NOM DU MEMBRE / PATIENT"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

c1 = a1.cell(0, 1)
c1.text = ""
p = c1.paragraphs[0]
r = p.add_run("NUMÉRO DE MEMBRE / CARTE")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c1.add_paragraph(lookup["NUMÉRO DE MEMBRE / CARTE"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

a2 = doc.add_table(rows=1, cols=2)
a2.width = Cm(18)
set_table_border(a2)
c0 = a2.cell(0, 0)
c0.text = ""
p = c0.paragraphs[0]
r = p.add_run("DATE DE NAISSANCE")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c0.add_paragraph(lookup["DATE DE NAISSANCE"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

c1 = a2.cell(0, 1)
c1.text = ""
p = c1.paragraphs[0]
r = p.add_run("SEXE")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
m_sym = "☐" if "unchecked" in lookup["SEXE - M"] else "☑"
f_sym = "☐" if "unchecked" in lookup["SEXE - F"] else "☑"
pm = c1.add_paragraph(f"{m_sym} M")
pm.runs[0].font.size = Pt(9)
pf = c1.add_paragraph(f"{f_sym} F")
pf.runs[0].font.size = Pt(9)

a3 = doc.add_table(rows=1, cols=2)
a3.width = Cm(18)
set_table_border(a3)
c0 = a3.cell(0, 0)
c0.text = ""
p = c0.paragraphs[0]
r = p.add_run("COMPAGNIE D'ASSURANCE / TPA")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c0.add_paragraph(lookup["COMPAGNIE D'ASSURANCE / TPA"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

c1 = a3.cell(0, 1)
c1.text = ""
p = c1.paragraphs[0]
r = p.add_run("TITULAIRE DE LA POLICE")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c1.add_paragraph(lookup["TITULAIRE DE LA POLICE"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

a4 = doc.add_table(rows=1, cols=2)
a4.width = Cm(18)
set_table_border(a4)
c0 = a4.cell(0, 0)
c0.text = ""
p = c0.paragraphs[0]
r = p.add_run("CPR / PASSEPORT")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c0.add_paragraph(lookup["CPR / PASSEPORT"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

c1 = a4.cell(0, 1)
c1.text = ""
p = c1.paragraphs[0]
r = p.add_run("TÉLÉPHONE DU MEMBRE")
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p2 = c1.add_paragraph(lookup["TÉLÉPHONE DU MEMBRE"])
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# SECTION B
sb_t = doc.add_table(rows=1, cols=1)
sb_t.width = Cm(18)
set_cell_bg(sb_t.cell(0, 0), "2E5E99")
cell_para(sb_t.cell(0, 0), lookup["SECTION B - INFORMATIONS MÉDICALES"], bold=True, size=10)
sb_t.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

b_main = doc.add_table(rows=1, cols=2)
b_main.width = Cm(18)
set_table_border(b_main)
b_main.columns[0].width = Cm(9)
b_main.columns[1].width = Cm(9)

left_b = b_main.cell(0, 0)
left_b.text = ""
ph = left_b.paragraphs[0]
rh = ph.add_run(lookup["Type de demande"])
rh.bold = True
rh.font.size = Pt(9)

hosp_sym = "☐" if "unchecked" in lookup["Hospitalisation"] else "☑"
amb_sym = "☐" if "unchecked" in lookup["Ambulatoire"] else "☑"
urg_sym = "☐" if "unchecked" in lookup["Urgence"] else "☑"

for sym, lbl in [(hosp_sym, "Hospitalisation"), (amb_sym, "Ambulatoire"), (urg_sym, "Urgence")]:
    pp = left_b.add_paragraph(f"{sym} {lbl}")
    pp.runs[0].font.size = Pt(9)
    pp.paragraph_format.space_before = Pt(2)

det_t = left_b.add_table(rows=4, cols=2)
det_t.width = Cm(9)
set_table_border(det_t)
det_rows = [
    ("Champ", "Détails"),
    ("Date du traitement", lookup["Date du traitement"]),
    ("Fournisseur / clinique", lookup["Fournisseur / clinique"]),
    ("Médecin traitant", lookup["Médecin traitant"]),
]
for i, (k, v) in enumerate(det_rows):
    cell_para(det_t.rows[i].cells[0], k, bold=(i == 0), size=9)
    cell_para(det_t.rows[i].cells[1], v, bold=(i == 0), size=9)

num_row = det_t.add_row()
cell_para(num_row.cells[0], "Numéro de dossier médical", bold=False, size=9)
cell_para(num_row.cells[1], lookup["Numéro de dossier médical"], bold=False, size=9)

right_b = b_main.cell(0, 1)
right_b.text = ""
ph2 = right_b.paragraphs[0]
rh2 = ph2.add_run(lookup["Conditions et signes vitaux"])
rh2.bold = True
rh2.font.size = Pt(9)

cond_t = right_b.add_table(rows=4, cols=3)
cond_t.width = Cm(9)
set_table_border(cond_t)

cell_para(cond_t.rows[0].cells[0], "Affection préexistante", bold=False, size=9)
oui_sym = "☐" if "unchecked" in lookup["Affection préexistante - Oui"] else "☑"
non_sym = "☐" if "unchecked" in lookup["Affection préexistante - Non"] else "☑"
cell_para(cond_t.rows[0].cells[1], f"{oui_sym} Oui", size=9)
cell_para(cond_t.rows[0].cells[2], f"{non_sym} Non", size=9)

cell_para(cond_t.rows[1].cells[0], "Accident du travail", bold=False, size=9)
oui2 = "☐" if "unchecked" in lookup["Accident du travail - Oui"] else "☑"
non2 = "☐" if "unchecked" in lookup["Accident du travail - Non"] else "☑"
cell_para(cond_t.rows[1].cells[1], f"{oui2} Oui", size=9)
cell_para(cond_t.rows[1].cells[2], f"{non2} Non", size=9)

cell_para(cond_t.rows[2].cells[0], "Maternité", bold=False, size=9)
oui3 = "☐" if "unchecked" in lookup["Maternité - Oui"] else "☑"
non3 = "☐" if "unchecked" in lookup["Maternité - Non"] else "☑"
cell_para(cond_t.rows[2].cells[1], f"{oui3} Oui", size=9)
cell_para(cond_t.rows[2].cells[2], f"{non3} Non", size=9)

cell_para(cond_t.rows[3].cells[0], "Tension / pouls / température", bold=False, size=9)
cell_para(cond_t.rows[3].cells[1], lookup["Tension / pouls / température"], size=9)
cell_para(cond_t.rows[3].cells[2], "", size=9)

b_bot = doc.add_table(rows=1, cols=2)
b_bot.width = Cm(18)
set_table_border(b_bot)

motif_cell = b_bot.cell(0, 0)
motif_cell.text = ""
pm1 = motif_cell.paragraphs[0]
rm1 = pm1.add_run("MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS")
rm1.bold = True
rm1.font.size = Pt(8)
pm2 = motif_cell.add_paragraph(lookup["MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS"])
pm2.runs[0].font.size = Pt(8)

diag_cell = b_bot.cell(0, 1)
diag_cell.text = ""
pd1 = diag_cell.paragraphs[0]
rd1 = pd1.add_run("DIAGNOSTIC FINAL")
rd1.bold = True
rd1.font.size = Pt(8)
pd2 = diag_cell.add_paragraph(lookup["DIAGNOSTIC FINAL"])
pd2.runs[0].font.size = Pt(8)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# SECTION C
sc_t = doc.add_table(rows=1, cols=1)
sc_t.width = Cm(18)
set_cell_bg(sc_t.cell(0, 0), "2E5E99")
cell_para(sc_t.cell(0, 0), lookup["SECTION C - TRAITEMENT, INVESTIGATIONS ET COÛT ESTIMÉ"], bold=True, size=10)
sc_t.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

c_tbl = doc.add_table(rows=1, cols=5)
c_tbl.width = Cm(18)
set_table_border(c_tbl)
c_tbl.columns[0].width = Cm(5)
c_tbl.columns[1].width = Cm(4)
c_tbl.columns[2].width = Cm(1.5)
c_tbl.columns[3].width = Cm(2.5)
c_tbl.columns[4].width = Cm(5)

hdrs = ["Traitement / service", "Catégorie", "Qté", "Coût brut", "Justification médicale"]
for i, h in enumerate(hdrs):
    cell_para(c_tbl.rows[0].cells[i], h, bold=True, size=9)
    set_cell_bg(c_tbl.rows[0].cells[i], "E8E8E8")

svc_keys = [
    "Traitement / service - Consultation ORL",
    "Traitement / service - Injection Lincocin",
    "Traitement / service - Injection Depo-Medrol",
]
svc_names = ["Consultation ORL", "Injection Lincocin", "Injection Depo-Medrol"]

for sn, sk in zip(svc_names, svc_keys):
    val = lookup[sk]
    parts = {}
    for part in val.split(" | "):
        if " : " in part:
            k2, v2 = part.split(" : ", 1)
            parts[k2.strip()] = v2.strip()
    nr = c_tbl.add_row()
    cell_para(nr.cells[0], sn, size=9)
    cell_para(nr.cells[1], parts.get("Catégorie", ""), size=9)
    cell_para(nr.cells[2], parts.get("Qté", ""), size=9)
    cell_para(nr.cells[3], parts.get("Coût brut", ""), size=9)
    cell_para(nr.cells[4], parts.get("Justification médicale", ""), size=9)

total_val = lookup["Coût prévu / total brut"]
tv_parts = total_val.split(" | ")
total_cost = tv_parts[0] if tv_parts else total_val
total_note = tv_parts[1] if len(tv_parts) > 1 else ""

fr = c_tbl.add_row()
cell_para(fr.cells[0], "Coût prévu / total brut", bold=True, size=9)
cell_para(fr.cells[1], "", size=9)
cell_para(fr.cells[2], "", size=9)
cell_para(fr.cells[3], total_cost, bold=True, size=9)
cell_para(fr.cells[4], total_note, bold=True, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# DECLARATIONS
decl_t = doc.add_table(rows=1, cols=2)
decl_t.width = Cm(18)
set_table_border(decl_t)

dc0 = decl_t.cell(0, 0)
dc0.text = ""
pd0 = dc0.paragraphs[0]
rd0 = pd0.add_run("Déclaration du membre")
rd0.bold = True
rd0.font.size = Pt(9)
pt0 = dc0.add_paragraph(lookup["Déclaration du membre"])
pt0.runs[0].font.size = Pt(8)
ps0 = dc0.add_paragraph("\n_________________________________")
ps0.runs[0].font.size = Pt(8)
pl0 = dc0.add_paragraph(lookup["Signature du membre / date"] or "Signature du membre / date")
pl0.runs[0].font.size = Pt(7)
pl0.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

dc1 = decl_t.cell(0, 1)
dc1.text = ""
pd1 = dc1.paragraphs[0]
rd1 = pd1.add_run("Déclaration du prestataire médical")
rd1.bold = True
rd1.font.size = Pt(9)
pt1 = dc1.add_paragraph(lookup["Déclaration du prestataire médical"])
pt1.runs[0].font.size = Pt(8)
ps1 = dc1.add_paragraph("\n_________________________________")
ps1.runs[0].font.size = Pt(8)
pl1 = dc1.add_paragraph(lookup["Nom du médecin, signature et cachet"] or "Nom du médecin, signature et cachet")
pl1.runs[0].font.size = Pt(7)
pl1.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)