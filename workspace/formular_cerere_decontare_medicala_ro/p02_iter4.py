import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p02.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_02.docx"
DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(par, text, bold=False, size=9, color=None):
    r = par.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r

def section_header(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(18)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "2E4057")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, text, bold=True, size=9, color=(255, 255, 255))
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def cb(val):
    return "☑" if "checked" in val and "unchecked" not in val else "☐"

# Header table
ht = doc.add_table(rows=1, cols=2)
ht.width = Cm(18)
ht.columns[0].width = Cm(11)
ht.columns[1].width = Cm(7)
lc = ht.rows[0].cells[0]
rc = ht.rows[0].cells[1]
lp = lc.paragraphs[0]
add_run(lp, "Formular de cerere de decontare medicală", bold=True, size=14, color=(30, 80, 160))
lp2 = lc.add_paragraph()
add_run(lp2, "Îngrijire primară - se completează de membru și de medicul curant", size=8)

info_items = [
    ("Nr. cerere", lookup["N° de demande"]),
    ("Ref. aprobare", lookup["Réf. d'approbation"]),
    ("Data tratamentului", lookup["Date du traitement"]),
    ("Tip", lookup["Type"]),
]
it = rc.add_table(rows=4, cols=2)
it.width = Cm(7)
for i, (k, v) in enumerate(info_items):
    r = it.rows[i]
    r.cells[0].width = Cm(3.2)
    r.cells[1].width = Cm(3.8)
    p0 = r.cells[0].paragraphs[0]
    add_run(p0, k, size=8)
    p1 = r.cells[1].paragraphs[0]
    add_run(p1, v, bold=True, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Section A
section_header(doc, "SECȚIUNEA A - DATELE MEMBRULUI ASIGURAT")
ta = doc.add_table(rows=5, cols=2)
ta.width = Cm(18)
for col in ta.columns:
    col.width = Cm(9)

def fill_cell(cell, label, value):
    p1 = cell.paragraphs[0]
    add_run(p1, label, size=8, color=(100, 100, 100))
    p2 = cell.add_paragraph()
    add_run(p2, value, bold=True, size=9)
    set_cell_borders(cell, top="single", bottom="single", left="single", right="single")

fill_cell(ta.rows[0].cells[0], "NUME MEMBRU / PACIENT", lookup["NOM DU MEMBRE / PATIENT"])
fill_cell(ta.rows[0].cells[1], "NUMĂR MEMBRU / CARD", lookup["NUMÉRO DE MEMBRE / CARTE"])
fill_cell(ta.rows[1].cells[0], "DATA NAȘTERII", lookup["DATE DE NAISSANCE"])
sex_cell = ta.rows[1].cells[1]
sp = sex_cell.paragraphs[0]
add_run(sp, "SEX", size=8, color=(100, 100, 100))
sp2 = sex_cell.add_paragraph()
add_run(sp2, cb(lookup["SEXE - M"]) + " M", size=9)
sp3 = sex_cell.add_paragraph()
add_run(sp3, cb(lookup["SEXE - F"]) + " F", size=9)
set_cell_borders(sex_cell, top="single", bottom="single", left="single", right="single")

fill_cell(ta.rows[2].cells[0], "COMPANIE DE ASIGURĂRI / TPA", lookup["COMPAGNIE D'ASSURANCE / TPA"])
fill_cell(ta.rows[2].cells[1], "DEȚINATOR POLIȚĂ", lookup["TITULAIRE DE LA POLICE"])
fill_cell(ta.rows[3].cells[0], "CPR / PAȘAPORT", lookup["CPR / PASSEPORT"])
fill_cell(ta.rows[3].cells[1], "TELEFON MEMBRU", lookup["TÉLÉPHONE DU MEMBRE"])

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Section B
section_header(doc, "SECȚIUNEA B - INFORMAȚII MEDICALE")

tb = doc.add_table(rows=1, cols=2)
tb.width = Cm(18)
tb.columns[0].width = Cm(9)
tb.columns[1].width = Cm(9)

left_cell = tb.rows[0].cells[0]
right_cell = tb.rows[0].cells[1]

# Left side: Tip solicitare
lp = left_cell.paragraphs[0]
add_run(lp, "Tip solicitare", bold=True, size=10)

for key, label in [("Hospitalisation", "Internare"), ("Ambulatoire", "Ambulatoriu"), ("Urgence", "Urgență")]:
    p = left_cell.add_paragraph()
    add_run(p, cb(lookup[key]) + " " + label, size=9)

# Câmp / Detalii sub-table
p_space = left_cell.add_paragraph()
p_space.paragraph_format.space_before = Pt(4)
add_run(p_space, "", size=4)

detail_t = left_cell.add_table(rows=5, cols=2)
detail_t.width = Cm(8.5)
headers_row = detail_t.rows[0]
set_cell_bg(headers_row.cells[0], "E8E8E8")
set_cell_bg(headers_row.cells[1], "E8E8E8")
h0 = headers_row.cells[0].paragraphs[0]
add_run(h0, "Câmp", bold=True, size=8)
h1 = headers_row.cells[1].paragraphs[0]
add_run(h1, "Detalii", bold=True, size=8)

detail_rows = [
    ("Data tratamentului", lookup["Date du traitement"]),
    ("Furnizor / clinică", lookup["Fournisseur / clinique"]),
    ("Medic curant", lookup["Médecin traitant"]),
    ("Număr fișă medicală", lookup["Numéro de dossier médical"]),
]
for i, (k, v) in enumerate(detail_rows):
    dr = detail_t.rows[i + 1]
    p0 = dr.cells[0].paragraphs[0]
    add_run(p0, k, size=8)
    p1 = dr.cells[1].paragraphs[0]
    add_run(p1, v, size=8)
    set_cell_borders(dr.cells[0], top="single", bottom="single", left="single", right="single")
    set_cell_borders(dr.cells[1], top="single", bottom="single", left="single", right="single")

# Right side: Condiții și semne vitale
rp = right_cell.paragraphs[0]
add_run(rp, "Condiții și semne vitale", bold=True, size=10)

cond_t = right_cell.add_table(rows=4, cols=2)
cond_t.width = Cm(8.5)
cond_header = cond_t.rows[0]
set_cell_bg(cond_header.cells[0], "E8E8E8")
set_cell_bg(cond_header.cells[1], "E8E8E8")
ch0 = cond_header.cells[0].paragraphs[0]
add_run(ch0, "Condiție", bold=True, size=8)
ch1 = cond_header.cells[1].paragraphs[0]
add_run(ch1, "Da / Nu", bold=True, size=8)

conditions = [
    ("Afecțiune preexistentă", "Affection préexistante - Oui", "Affection préexistante - Non"),
    ("Accident de muncă", "Accident du travail - Oui", "Accident du travail - Non"),
    ("Maternitate", "Maternité - Oui", "Maternité - Non"),
]
for i, (label, k_da, k_nu) in enumerate(conditions):
    cr = cond_t.rows[i + 1]
    p0 = cr.cells[0].paragraphs[0]
    add_run(p0, label, size=8)
    p1 = cr.cells[1].paragraphs[0]
    add_run(p1, cb(lookup[k_da]) + " Da  " + cb(lookup[k_nu]) + " Nu", size=8)
    set_cell_borders(cr.cells[0], top="single", bottom="single", left="single", right="single")
    set_cell_borders(cr.cells[1], top="single", bottom="single", left="single", right="single")

rp2 = right_cell.add_paragraph()
rp2.paragraph_format.space_before = Pt(4)
add_run(rp2, "Tensiune / puls / temperatură", size=8, color=(100, 100, 100))
rp3 = right_cell.add_paragraph()
add_run(rp3, lookup["Tension / pouls / température"], size=8)

set_cell_borders(left_cell, top="single", bottom="single", left="single", right="single")
set_cell_borders(right_cell, top="single", bottom="single", left="single", right="single")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Motiv + Diagnostic
tb2 = doc.add_table(rows=1, cols=2)
tb2.width = Cm(18)
tb2.columns[0].width = Cm(9)
tb2.columns[1].width = Cm(9)

mc = tb2.rows[0].cells[0]
dc = tb2.rows[0].cells[1]
set_cell_borders(mc, top="single", bottom="single", left="single", right="single")
set_cell_borders(dc, top="single", bottom="single", left="single", right="single")

mp = mc.paragraphs[0]
add_run(mp, "MOTIV CONSULTAȚIE / SIMPTOME PREZENTATE", bold=True, size=8)
mp2 = mc.add_paragraph()
add_run(mp2, lookup["MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS"], size=8)

dp = dc.paragraphs[0]
add_run(dp, "DIAGNOSTIC FINAL", bold=True, size=8)
diag_val = lookup["DIAGNOSTIC FINAL"]
for line in diag_val.split("\n"):
    dp2 = dc.add_paragraph()
    add_run(dp2, line, size=8)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Section C
section_header(doc, "SECȚIUNEA C - TRATAMENT, INVESTIGAȚII ȘI COST ESTIMAT")

tc = doc.add_table(rows=1, cols=5)
tc.width = Cm(18)
widths = [Cm(5), Cm(4), Cm(1.5), Cm(2.5), Cm(5)]
col_headers = ["Tratament / serviciu", "Categorie", "Cant.", "Cost brut", "Justificare medicală"]
for i, (w, h) in enumerate(zip(widths, col_headers)):
    c = tc.rows[0].cells[i]
    c.width = w
    set_cell_bg(c, "E8E8E8")
    p = c.paragraphs[0]
    add_run(p, h, bold=True, size=9)
    set_cell_borders(c, top="single", bottom="single", left="single", right="single")

service_rows = [
    ("Traitement / service - Consultation ORL", "Consultație ORL", "Consultație specialist", "1", "15.00 BD", "Evaluare clinică"),
    ("Traitement / service - Injection Lincocin", "Injecție Lincocin", "Injecții și perfuzii", "1", "10.00 BD", "Indicat medical"),
    ("Traitement / service - Injection Depo-Medrol", "Injecție Depo-Medrol", "Injecții și perfuzii", "1", "10.00 BD", "Indicat medical"),
]

for key, name, cat, qty, cost, just in service_rows:
    nr = tc.add_row()
    vals = [name, cat, qty, cost, just]
    for i, v in enumerate(vals):
        c = nr.cells[i]
        c.width = widths[i]
        p = c.paragraphs[0]
        add_run(p, v, size=9)
        set_cell_borders(c, top="single", bottom="single", left="single", right="single")

total_row = tc.add_row()
tc0 = total_row.cells[0]
tc0.merge(total_row.cells[1]).merge(total_row.cells[2])
tp = tc0.paragraphs[0]
add_run(tp, "Cost anticipat / total brut", bold=True, size=9)
set_cell_borders(tc0, top="single", bottom="single", left="single", right="single")

tc3 = total_row.cells[3]
tp3 = tc3.paragraphs[0]
add_run(tp3, "35.00 BD", bold=True, size=9)
set_cell_borders(tc3, top="single", bottom="single", left="single", right="single")

tc4 = total_row.cells[4]
tp4 = tc4.paragraphs[0]
add_run(tp4, "Subiect al revizuirii de eligibilitate", bold=True, size=9)
set_cell_borders(tc4, top="single", bottom="single", left="single", right="single")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Declarations
td = doc.add_table(rows=1, cols=2)
td.width = Cm(18)
td.columns[0].width = Cm(9)
td.columns[1].width = Cm(9)

dec_m = td.rows[0].cells[0]
dec_f = td.rows[0].cells[1]

dm_p = dec_m.paragraphs[0]
add_run(dm_p, "Declarația membrului", bold=True, size=9)
dm_p2 = dec_m.add_paragraph()
add_run(dm_p2, lookup["Déclaration du membre"], size=8)
dm_p3 = dec_m.add_paragraph()
dm_p3.paragraph_format.space_before = Pt(20)
add_run(dm_p3, "_" * 40, size=8)
dm_p4 = dec_m.add_paragraph()
add_run(dm_p4, "Semnătură membru / dată", size=8, color=(150, 150, 150))

df_p = dec_f.paragraphs[0]
add_run(df_p, "Declarația furnizorului medical", bold=True, size=9)
df_p2 = dec_f.add_paragraph()
add_run(df_p2, lookup["Déclaration du prestataire médical"], size=8)
df_p3 = dec_f.add_paragraph()
df_p3.paragraph_format.space_before = Pt(20)
add_run(df_p3, "_" * 40, size=8)
df_p4 = dec_f.add_paragraph()
add_run(df_p4, "Nume medic, semnătură și ștampilă", size=8, color=(150, 150, 150))

set_cell_borders(dec_m, top="single", bottom="single", left="single", right="single")
set_cell_borders(dec_f, top="single", bottom="single", left="single", right="single")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, "Formular claim medical - RO | pag. 2/3", size=7, color=(150, 150, 150))

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)