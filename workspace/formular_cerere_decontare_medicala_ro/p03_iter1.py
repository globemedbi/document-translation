import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p03.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_03.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

NAVY = RGBColor(0x1F, 0x35, 0x64)
GREY_BG = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
GREEN = RGBColor(0x70, 0xAD, 0x47)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    hex_color = "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_bg_hex(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"val": "none", "sz": 0, "color": "FFFFFF"},
                bottom={"val": "none", "sz": 0, "color": "FFFFFF"},
                left={"val": "none", "sz": 0, "color": "FFFFFF"},
                right={"val": "none", "sz": 0, "color": "FFFFFF"})


def cell_para(cell, text, size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    return p


def add_para(doc, text, size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    return p


def section_header(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(17)
    c = t.rows[0].cells[0]
    set_cell_bg_hex(c, "1F3564")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.2)

for para in doc.paragraphs:
    para.clear()

# BAND 1 - HEADER
header_t = doc.add_table(rows=1, cols=2)
header_t.width = Cm(17)
no_borders(header_t)

left_cell = header_t.rows[0].cells[0]
right_cell = header_t.rows[0].cells[1]
left_cell.width = Cm(11)
right_cell.width = Cm(6)

# MC logo + title in left cell
left_inner = left_cell.add_table(rows=1, cols=2)
no_borders(left_inner)
logo_cell = left_inner.rows[0].cells[0]
logo_cell.width = Cm(2.2)
title_cell = left_inner.rows[0].cells[1]
title_cell.width = Cm(8.8)

set_cell_bg_hex(logo_cell, "1F3564")
lp = logo_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lp.paragraph_format.space_before = Pt(8)
lp.paragraph_format.space_after = Pt(2)
lr = lp.add_run("MC")
lr.font.size = Pt(14)
lr.font.bold = True
lr.font.color.rgb = WHITE

tp = title_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp.paragraph_format.space_before = Pt(2)
tp.paragraph_format.space_after = Pt(2)
tr = tp.add_run(lookup["Rapport d'approbation et détails de couverture"])
tr.font.size = Pt(14)
tr.font.bold = True
tr.font.color.rgb = NAVY

subtitle_p = title_cell.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
subtitle_p.paragraph_format.space_after = Pt(2)
sr = subtitle_p.add_run("Résumé pour la compagnie d'assurance - éligibilité, documents et conditions de paiement")
sr.font.size = Pt(7.5)
sr.font.italic = True
sr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# Right info box
set_cell_bg_hex(right_cell, "F2F2F2")
info_keys = ["Réf. approbation", "Date d'approbation", "Valable jusqu'au", "Statut"]
for k in info_keys:
    rp = right_cell.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp.paragraph_format.space_before = Pt(1)
    rp.paragraph_format.space_after = Pt(1)
    rk = rp.add_run(k + "  ")
    rk.font.size = Pt(8)
    rk.font.color.rgb = DARK_GREY
    rv = rp.add_run(lookup[k])
    rv.font.size = Pt(9)
    rv.font.bold = True
    rv.font.color.rgb = NAVY

sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(4)

# BAND 2 - DATE DE REFERINTA
section_header(doc, lookup["DONNÉES DE RÉFÉRENCE"])

ref_t = doc.add_table(rows=3, cols=2)
ref_t.width = Cm(17)
no_borders(ref_t)

ref_fields_left = ["NUMÉRO INDIVIDUEL", "DATE DE NAISSANCE", "MÉDECIN / PRESTATAIRE"]
ref_fields_right = ["NOM DU TITULAIRE DE CARTE", "COMPAGNIE D'ASSURANCE", "SPÉCIALITÉ"]

for i, (lk, rk) in enumerate(zip(ref_fields_left, ref_fields_right)):
    lc = ref_t.rows[i].cells[0]
    rc = ref_t.rows[i].cells[1]
    set_cell_bg_hex(lc, "F7F7F7")
    set_cell_bg_hex(rc, "F7F7F7")
    set_cell_border(lc, bottom={"val": "single", "sz": 4, "color": "DDDDDD"})
    set_cell_border(rc, bottom={"val": "single", "sz": 4, "color": "DDDDDD"})

    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after = Pt(1)
    lr = lp.add_run(lk)
    lr.font.size = Pt(7.5)
    lr.font.color.rgb = DARK_GREY

    lp2 = lc.add_paragraph()
    lp2.paragraph_format.space_after = Pt(3)
    lr2 = lp2.add_run(lookup[lk])
    lr2.font.size = Pt(9)
    lr2.font.bold = True
    lr2.font.color.rgb = NAVY

    rp = rc.paragraphs[0]
    rp.paragraph_format.space_before = Pt(2)
    rp.paragraph_format.space_after = Pt(1)
    rr = rp.add_run(rk)
    rr.font.size = Pt(7.5)
    rr.font.color.rgb = DARK_GREY

    rp2 = rc.add_paragraph()
    rp2.paragraph_format.space_after = Pt(3)
    rr2 = rp2.add_run(lookup[rk])
    rr2.font.size = Pt(9)
    rr2.font.bold = True
    rr2.font.color.rgb = NAVY

sp2 = doc.add_paragraph()
sp2.paragraph_format.space_after = Pt(4)

# BAND 3 - SERVICII APROBATE
section_header(doc, lookup["SERVICES APPROUVÉS"])

svc_t = doc.add_table(rows=4, cols=6)
svc_t.width = Cm(17)

# Header row
hdr_labels = ["Catégorie de service", "Élément approuvé", "Qté\ndemandée", "Qté\napprouvée", "Statut", "Motif / note"]
for i, h in enumerate(hdr_labels):
    c = svc_t.rows[0].cells[i]
    set_cell_bg_hex(c, "F2F2F2")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h)
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = NAVY

svc_rows = [
    lookup["Catégorie de service - Ligne 1"].split(" | "),
    lookup["Catégorie de service - Ligne 2"].split(" | "),
    lookup["Catégorie de service - Ligne 3"].split(" | "),
]

for ri, row_data in enumerate(svc_rows):
    for ci, val in enumerate(row_data):
        c = svc_t.rows[ri + 1].cells[ci]
        set_cell_bg_hex(c, "FFFFFF")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if ci == 4:
            r = p.add_run(val.strip())
            r.font.size = Pt(8)
            r.font.bold = True
            r.font.color.rgb = NAVY
            set_cell_bg_hex(c, "E8F0E9")
        else:
            r = p.add_run(val.strip())
            r.font.size = Pt(8)

sp3 = doc.add_paragraph()
sp3.paragraph_format.space_after = Pt(4)

# BAND 4 - DOCUMENTE + DIAGNOSTIC (two columns)
band4_t = doc.add_table(rows=1, cols=2)
band4_t.width = Cm(17)
no_borders(band4_t)
lc4 = band4_t.rows[0].cells[0]
rc4 = band4_t.rows[0].cells[1]
lc4.width = Cm(8.3)
rc4.width = Cm(8.7)

# Left: documents
set_cell_bg_hex(lc4, "F2F2F2")
lc4_hdr = lc4.paragraphs[0]
lc4_hdr.paragraph_format.space_before = Pt(3)
lc4_hdr.paragraph_format.space_after = Pt(3)
lhdr_r = lc4_hdr.add_run(lookup["DOCUMENTS REQUIS POUR LA SOUMISSION DE LA DEMANDE DE REMBOURSEMENT"])
lhdr_r.font.size = Pt(8)
lhdr_r.font.bold = True
lhdr_r.font.color.rgb = NAVY

docs_text = lookup["Liste des documents requis"]
for line in docs_text.split("\n"):
    lp = lc4.add_paragraph()
    lp.paragraph_format.space_before = Pt(1)
    lp.paragraph_format.space_after = Pt(1)
    lr = lp.add_run(line)
    lr.font.size = Pt(8)

# Right: diagnostic
set_cell_bg_hex(rc4, "FFFFFF")
rc4_hdr = rc4.paragraphs[0]
rc4_hdr.paragraph_format.space_before = Pt(3)
rc4_hdr.paragraph_format.space_after = Pt(3)
rhdr_r = rc4_hdr.add_run(lookup["DIAGNOSTIC ET INFORMATIONS MÉDICALES"])
rhdr_r.font.size = Pt(8)
rhdr_r.font.bold = True
rhdr_r.font.color.rgb = NAVY

diag_inner = rc4.add_table(rows=5, cols=2)
no_borders(diag_inner)
diag_fields = [
    ("Diagnostic principal", lookup["Diagnostic principal"]),
    ("Date de la visite", lookup["Date de la visite"]),
    ("Tension / pouls / température", lookup["Tension / pouls / température"]),
    ("Durée de la maladie", lookup["Durée de la maladie"]),
    ("Fichiers téléchargés", lookup["Fichiers téléchargés"]),
]
for i, (k, v) in enumerate(diag_fields):
    kc = diag_inner.rows[i].cells[0]
    vc = diag_inner.rows[i].cells[1]
    set_cell_bg_hex(kc, "F7F7F7")
    set_cell_bg_hex(vc, "FFFFFF")
    kp = kc.paragraphs[0]
    kp.paragraph_format.space_before = Pt(1)
    kp.paragraph_format.space_after = Pt(1)
    kr = kp.add_run(k)
    kr.font.size = Pt(8)
    kr.font.color.rgb = DARK_GREY
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(1)
    vp.paragraph_format.space_after = Pt(1)
    vr = vp.add_run(v)
    vr.font.size = Pt(8)

sp4 = doc.add_paragraph()
sp4.paragraph_format.space_after = Pt(4)

# BAND 5 - CONDITIONS DE PAIEMENT
section_header(doc, lookup["CONDITIONS D'ÉVALUATION ET DE PAIEMENT"])

cond_text = lookup["Texte des conditions d'évaluation"]
for para_text in cond_text.split("\n"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    if para_text.startswith("Note TVA"):
        bold_part = "Note TVA :"
        rest = para_text[len(bold_part):]
        rb = p.add_run(bold_part)
        rb.font.size = Pt(8.5)
        rb.font.bold = True
        rn = p.add_run(rest)
        rn.font.size = Pt(8.5)
    else:
        r = p.add_run(para_text)
        r.font.size = Pt(8.5)

sp5 = doc.add_paragraph()
sp5.paragraph_format.space_after = Pt(4)

# BAND 6 - UTILIZARE INTERNA + CONTROL CALITATE
band6_t = doc.add_table(rows=1, cols=2)
band6_t.width = Cm(17)
no_borders(band6_t)
lc6 = band6_t.rows[0].cells[0]
rc6 = band6_t.rows[0].cells[1]
lc6.width = Cm(8.5)
rc6.width = Cm(8.5)

set_cell_bg_hex(lc6, "F7F7F7")
set_cell_bg_hex(rc6, "FFFFFF")

lc6_hdr = lc6.paragraphs[0]
lc6_hdr.paragraph_format.space_before = Pt(3)
lc6_hdr.paragraph_format.space_after = Pt(3)
l6hr = lc6_hdr.add_run(lookup["Usage interne - compagnie d'assurance"])
l6hr.font.size = Pt(9)
l6hr.font.bold = True
l6hr.font.color.rgb = NAVY

dec_p = lc6.add_paragraph()
dec_p.paragraph_format.space_before = Pt(2)
dec_p.paragraph_format.space_after = Pt(1)
dec_r = dec_p.add_run("Décision")
dec_r.font.size = Pt(8)
dec_r.font.bold = True
dec_r.font.color.rgb = DARK_GREY

app_p = lc6.add_paragraph()
app_p.paragraph_format.space_before = Pt(1)
app_p.paragraph_format.space_after = Pt(1)
app_r = app_p.add_run("☑ Approuvé")
app_r.font.size = Pt(9)

res_p = lc6.add_paragraph()
res_p.paragraph_format.space_before = Pt(1)
res_p.paragraph_format.space_after = Pt(3)
res_r = res_p.add_run("☐ Refusé")
res_r.font.size = Pt(9)

nr_p = lc6.add_paragraph()
nr_p.paragraph_format.space_before = Pt(1)
nr_p.paragraph_format.space_after = Pt(1)
nr_label = nr_p.add_run("N° d'approbation   ")
nr_label.font.size = Pt(8)
nr_label.font.bold = True
nr_val = nr_p.add_run(lookup["N° d'approbation"])
nr_val.font.size = Pt(9)

com_p = lc6.add_paragraph()
com_p.paragraph_format.space_before = Pt(1)
com_p.paragraph_format.space_after = Pt(1)
com_label = com_p.add_run("Commentaires   ")
com_label.font.size = Pt(8)
com_label.font.bold = True
com_val = com_p.add_run(lookup["Commentaires"])
com_val.font.size = Pt(9)

# Right: control calitate
rc6_hdr = rc6.paragraphs[0]
rc6_hdr.paragraph_format.space_before = Pt(3)
rc6_hdr.paragraph_format.space_after = Pt(2)
r6hr = rc6_hdr.add_run(lookup["Contrôle qualité de la demande"])
r6hr.font.size = Pt(9)
r6hr.font.bold = True
r6hr.font.color.rgb = NAVY

desc_p = rc6.add_paragraph()
desc_p.paragraph_format.space_before = Pt(1)
desc_p.paragraph_format.space_after = Pt(8)
desc_r = desc_p.add_run(lookup["Contrôle qualité de la demande - description"])
desc_r.font.size = Pt(8)
desc_r.font.color.rgb = DARK_GREY

sig_line_p = rc6.add_paragraph()
sig_line_p.paragraph_format.space_before = Pt(10)
sig_line_p.paragraph_format.space_after = Pt(2)
sig_r = sig_line_p.add_run("_" * 45)
sig_r.font.size = Pt(9)

sig_label_p = rc6.add_paragraph()
sig_label_p.paragraph_format.space_before = Pt(1)
sig_label_p.paragraph_format.space_after = Pt(2)
slr = sig_label_p.add_run("Nom de l'évaluateur / signature / date")
slr.font.size = Pt(8)
slr.font.color.rgb = DARK_GREY

sp6 = doc.add_paragraph()
sp6.paragraph_format.space_after = Pt(6)

# FOOTER
footer_t = doc.add_table(rows=1, cols=2)
footer_t.width = Cm(17)
no_borders(footer_t)
fc_l = footer_t.rows[0].cells[0]
fc_r = footer_t.rows[0].cells[1]
set_cell_bg_hex(fc_l, "F2F2F2")
set_cell_bg_hex(fc_r, "F2F2F2")

fl_p = fc_l.paragraphs[0]
fl_p.paragraph_format.space_before = Pt(2)
fl_p.paragraph_format.space_after = Pt(2)
flr = fl_p.add_run("Formular claim medical - RO | pag. 3/3")
flr.font.size = Pt(7)
flr.font.color.rgb = DARK_GREY

fr_p = fc_r.paragraphs[0]
fr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr_p.paragraph_format.space_before = Pt(2)
fr_p.paragraph_format.space_after = Pt(2)
frr = fr_p.add_run("Document model cu date anonimizate")
frr.font.size = Pt(7)
frr.font.color.rgb = DARK_GREY

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)