import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p01.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_01.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()

s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.2)

TEAL = RGBColor(0x00, 0x7B, 0x8A)
DARK = RGBColor(0x2C, 0x2C, 0x2C)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_TEAL = RGBColor(0xE0, 0xF4, 0xF7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY_BG = RGBColor(0xF5, 0xF5, 0xF5)
BLUE_BANNER = RGBColor(0xEA, 0xF6, 0xF8)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "CCCCCC")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders_none(t):
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_borders_light(t):
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def para_in_cell(cell, text, bold=False, size=9, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_para(text, bold=False, size=9, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def set_col_width(t, col_idx, width_cm):
    for row in t.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ── HEADER BAR ──────────────────────────────────────────────────────────────
t = doc.add_table(rows=1, cols=2)
t.width = Cm(18)
set_table_borders_none(t)

left_cell = t.rows[0].cells[0]
left_cell.width = Cm(12)
set_cell_bg(left_cell, "FFFFFF")

inner = left_cell.add_table(rows=1, cols=2)
set_table_borders_none(inner)
logo_cell = inner.rows[0].cells[0]
logo_cell.width = Cm(1.8)
set_cell_bg(logo_cell, "007B8A")
p = logo_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MC")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = WHITE

title_cell = inner.rows[0].cells[1]
title_cell.width = Cm(10)
p = title_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(lookup["Reçu médical et récapitulatif des coûts"])
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = DARK
p2 = title_cell.add_paragraph()
run2 = p2.add_run(lookup["Dossier de remboursement - services ORL, soins primaires / ambulatoire"])
run2.font.size = Pt(8)
run2.font.color.rgb = GREY

right_cell = t.rows[0].cells[1]
right_cell.width = Cm(6)
set_cell_bg(right_cell, "FFFFFF")

meta_labels = ["N° de dossier", "Date", "Devise", "Statut"]
meta_t = right_cell.add_table(rows=4, cols=2)
set_table_borders_none(meta_t)
for i, lbl in enumerate(meta_labels):
    lc = meta_t.rows[i].cells[0]
    vc = meta_t.rows[i].cells[1]
    lc.width = Cm(2.5)
    vc.width = Cm(3.5)
    para_in_cell(lc, lbl, bold=True, size=8, color=DARK)
    para_in_cell(vc, lookup[lbl], bold=(lbl in ["N° de dossier", "Statut"]), size=8, color=DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph()

# ── BLUE BANNER ─────────────────────────────────────────────────────────────
tb = doc.add_table(rows=1, cols=1)
tb.width = Cm(18)
set_table_borders_none(tb)
bc = tb.rows[0].cells[0]
set_cell_bg(bc, "EAF6F8")
p = bc.paragraphs[0]
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Objet du document: ")
run.bold = True
run.font.size = Pt(8.5)
run.font.color.rgb = DARK
run2 = p.add_run(lookup["Objet du document"])
run2.font.size = Pt(8.5)
run2.font.color.rgb = DARK

doc.add_paragraph()

# ── PROVIDER / INSURANCE TWO-COLUMN ─────────────────────────────────────────
tp = doc.add_table(rows=1, cols=2)
tp.width = Cm(18)
set_table_borders_none(tp)

pc = tp.rows[0].cells[0]
pc.width = Cm(8.5)
set_cell_bg(pc, "F9F9F9")

ic = tp.rows[0].cells[1]
ic.width = Cm(9.5)
set_cell_bg(ic, "F9F9F9")

def fill_provider(cell):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(lookup["Prestataire médical"])
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK

    provider_fields = [
        ("Clinique", lookup["Clinique"]),
        ("Spécialité", lookup["Spécialité"]),
        ("Adresse", lookup["Adresse"]),
        ("Téléphone", lookup["Téléphone"]),
    ]
    inner = cell.add_table(rows=len(provider_fields), cols=2)
    set_table_borders_none(inner)
    for i, (lbl, val) in enumerate(provider_fields):
        lc = inner.rows[i].cells[0]
        vc = inner.rows[i].cells[1]
        lc.width = Cm(2.5)
        vc.width = Cm(6)
        para_in_cell(lc, lbl, bold=True, size=8.5, color=TEAL)
        para_in_cell(vc, val, size=8.5, color=DARK)

fill_provider(pc)

def fill_insurance(cell):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(lookup["Assurance / gestionnaire de sinistre"])
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK

    ins_fields = [
        ("Compagnie / TPA", lookup["Compagnie / TPA"]),
        ("Titulaire de la police", lookup["Titulaire de la police"]),
        ("Membre", lookup["Membre"]),
        ("Type de demande", lookup["Type de demande"]),
    ]
    inner = cell.add_table(rows=len(ins_fields), cols=2)
    set_table_borders_none(inner)
    for i, (lbl, val) in enumerate(ins_fields):
        lc = inner.rows[i].cells[0]
        vc = inner.rows[i].cells[1]
        lc.width = Cm(3)
        vc.width = Cm(6.5)
        para_in_cell(lc, lbl, bold=True, size=8.5, color=TEAL)
        para_in_cell(vc, val, bold=(lbl == "Compagnie / TPA"), size=8.5, color=DARK)

fill_insurance(ic)

doc.add_paragraph()

# ── FINANCIAL SUMMARY HEADER ─────────────────────────────────────────────────
tf = doc.add_table(rows=1, cols=1)
tf.width = Cm(18)
set_table_borders_none(tf)
fhc = tf.rows[0].cells[0]
set_cell_bg(fhc, "007B8A")
p = fhc.paragraphs[0]
run = p.add_run(lookup["SYNTHÈSE FINANCIÈRE"])
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = WHITE

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── THREE KPI BOXES ──────────────────────────────────────────────────────────
tk = doc.add_table(rows=1, cols=3)
tk.width = Cm(18)
set_table_borders_none(tk)

kpi_data = [
    ("TOTAL BRUT", lookup["TOTAL BRUT"], "Services facturés par la clinique", "F0F0F0"),
    ("RÉDUCTIONS CONTRACTUELLES", lookup["RÉDUCTIONS CONTRACTUELLES"], "Appliquées sur consultation et injections", "F0F0F0"),
    ("MONTANT NET DEMANDÉ", lookup["MONTANT NET DEMANDÉ"], "Après réduction et déduction/quote-part", "F0F0F0"),
]

for i, (label, value, sub, bg) in enumerate(kpi_data):
    c = tk.rows[0].cells[i]
    c.width = Cm(6)
    set_cell_bg(c, bg)
    p1 = c.paragraphs[0]
    r1 = p1.add_run(label)
    r1.bold = True
    r1.font.size = Pt(7.5)
    r1.font.color.rgb = GREY
    p2 = c.add_paragraph()
    r2 = p2.add_run(value)
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = TEAL
    p3 = c.add_paragraph()
    r3 = p3.add_run(sub)
    r3.font.size = Pt(7.5)
    r3.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── SERVICES TABLE ───────────────────────────────────────────────────────────
headers = ["Service / traitement", "Qté", "Tarif brut", "Réduction", "Valeur éligible", "Observations"]
col_widths = [5.5, 1.2, 2.3, 2.3, 2.5, 2.7]

rows_raw = lookup["Tableau des services"].split("\n")
service_rows = []
for row_str in rows_raw:
    parts = [p.strip() for p in row_str.split("|")]
    while len(parts) < 6:
        parts.append("")
    service_rows.append(parts)

ts = doc.add_table(rows=1 + len(service_rows), cols=6)
ts.width = Cm(18)
set_table_borders_light(ts)

for ci, (h, w) in enumerate(zip(headers, col_widths)):
    hc = ts.rows[0].cells[ci]
    hc.width = Cm(w)
    set_cell_bg(hc, "E8E8E8")
    para_in_cell(hc, h, bold=True, size=8.5, color=DARK)

for ri, row_data in enumerate(service_rows):
    is_total = row_data[0].strip().lower() == "total"
    bg = "F5F5F5" if is_total else "FFFFFF"
    for ci, val in enumerate(row_data):
        c = ts.rows[ri + 1].cells[ci]
        c.width = Cm(col_widths[ci])
        set_cell_bg(c, bg)
        para_in_cell(c, val, bold=is_total, size=8.5, color=DARK)

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ── DEDUCTION / NET ROWS ──────────────────────────────────────────────────────
td = doc.add_table(rows=2, cols=2)
td.width = Cm(18)
set_table_borders_light(td)

ded_label = td.rows[0].cells[0]
ded_label.width = Cm(14)
set_cell_bg(ded_label, "FAFAFA")
para_in_cell(ded_label, lookup["Déduction / quote-part du membre"], bold=True, size=8.5, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)

ded_val = td.rows[0].cells[1]
ded_val.width = Cm(4)
set_cell_bg(ded_val, "FAFAFA")
para_in_cell(ded_val, lookup["Déduction / quote-part du membre".replace("Déduction / quote-part du membre", "Déduction / quote-part du membre")], bold=True, size=8.5, color=DARK)
ded_val.paragraphs[0].clear()
run = ded_val.paragraphs[0].add_run(lookup["Déduction / quote-part du membre"])
run.bold = True
run.font.size = Pt(8.5)

net_label = td.rows[1].cells[0]
net_label.width = Cm(14)
set_cell_bg(net_label, "FAFAFA")
para_in_cell(net_label, lookup["Montant net à rembourser / à percevoir"], bold=True, size=8.5, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)

net_val = td.rows[1].cells[1]
net_val.width = Cm(4)
set_cell_bg(net_val, "E0F4F7")
p = net_val.paragraphs[0]
p.clear()
run = p.add_run(lookup["Montant net à rembourser / à percevoir"])
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = TEAL

# fix deduction value cell
ded_val.paragraphs[0].clear()
r = ded_val.paragraphs[0].add_run(lookup["Déduction / quote-part du membre"])
r.bold = True
r.font.size = Pt(8.5)
r.font.color.rgb = DARK

doc.add_paragraph()

# ── SIGNATURE BLOCK ───────────────────────────────────────────────────────────
tsig = doc.add_table(rows=1, cols=2)
tsig.width = Cm(18)
set_table_borders_none(tsig)

sc1 = tsig.rows[0].cells[0]
sc1.width = Cm(8.5)
set_cell_bg(sc1, "F9F9F9")
p = sc1.paragraphs[0]
run = p.add_run(lookup["Confirmation du prestataire médical"])
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = DARK
p2 = sc1.add_paragraph()
run2 = p2.add_run(lookup["Confirmation du prestataire médical - texte"])
run2.font.size = Pt(8)
run2.font.color.rgb = GREY
sc1.add_paragraph()
sc1.add_paragraph()
p3 = sc1.add_paragraph()
run3 = p3.add_run(lookup["Nom, signature et cachet"] if lookup["Nom, signature et cachet"] else "Nom, signature et cachet")
run3.font.size = Pt(8)
run3.font.color.rgb = GREY
p3.paragraph_format.space_before = Pt(20)

sc2 = tsig.rows[0].cells[1]
sc2.width = Cm(9.5)
set_cell_bg(sc2, "F9F9F9")
p = sc2.paragraphs[0]
run = p.add_run(lookup["Confirmation du membre / patient"])
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = DARK
p2 = sc2.add_paragraph()
run2 = p2.add_run(lookup["Confirmation du membre / patient - texte"])
run2.font.size = Pt(8)
run2.font.color.rgb = GREY
sc2.add_paragraph()
sc2.add_paragraph()
p3 = sc2.add_paragraph()
run3 = p3.add_run(lookup["Nom et signature du membre"] if lookup["Nom et signature du membre"] else "Nom et signature du membre")
run3.font.size = Pt(8)
run3.font.color.rgb = GREY
p3.paragraph_format.space_before = Pt(20)

doc.add_paragraph()

# ── FOOTER ────────────────────────────────────────────────────────────────────
fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(10)
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
section_footer = doc.sections[0]
footer = section_footer.footer
fp2 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp2.clear()
fp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_l = fp2.add_run("Formulaire claim médical - FR | pag. 1/3")
run_l.font.size = Pt(7.5)
run_l.font.color.rgb = GREY
tab_run = fp2.add_run("\t\t\t\t")
tab_run.font.size = Pt(7.5)
run_r = fp2.add_run("Document modèle avec données anonymisées")
run_r.font.size = Pt(7.5)
run_r.font.color.rgb = GREY

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)