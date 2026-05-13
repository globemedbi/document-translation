import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p02.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_02.docx"
DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(cell, text, bold=False, size=9, italic=False):
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def section_header(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(18)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "1F3864")
    p = add_para(c, text, bold=True, size=10)
    p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

# HEADER
ht = doc.add_table(rows=1, cols=2)
ht.width = Cm(18)
ht.columns[0].width = Cm(11)
ht.columns[1].width = Cm(7)
lc = ht.rows[0].cells[0]
rc = ht.rows[0].cells[1]

p = lc.paragraphs[0]
run = p.add_run("MC  ")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
lc_bg_p = p
set_cell_bg(lc, "1F3864")

p2 = lc.add_paragraph()
run2 = p2.add_run(lookup["Formulaire de demande de remboursement médical"])
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p3 = lc.add_paragraph()
run3 = p3.add_run("Soins primaires - à remplir par le membre et le médecin traitant")
run3.italic = True
run3.font.size = Pt(9)
run3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

info_keys = ["N° de demande", "Réf. d'approbation", "Date du traitement", "Type"]
info_labels = ["Nr. cerere", "Ref. aprobare", "Data tratamentului", "Tip"]
it = doc.add_table(rows=4, cols=2)
for i, (k, lbl) in enumerate(zip(info_keys, info_labels)):
    r = it.rows[i]
    r.cells[0].paragraphs[0].add_run(lbl).font.size = Pt(8)
    vrun = r.cells[1].paragraphs[0].add_run(lookup[k])
    vrun.bold = True
    vrun.font.size = Pt(9)

rc_p = rc.paragraphs[0]
rc_p.add_run("")
set_table_border(it)

inner_xml = it._tbl.xml
from lxml import etree
it_elem = it._tbl
rc._tc.append(it_elem)

doc.add_paragraph()

# SECTION A
section_header(doc, lookup["SECTION A - DONNÉES DU MEMBRE ASSURÉ"])

at = doc.add_table(rows=5, cols=2)
at.width = Cm(18)
set_table_border(at)

def fill_a(row_idx, left_label, left_val, right_label, right_val):
    lc = at.rows[row_idx].cells[0]
    rc = at.rows[row_idx].cells[1]
    add_para(lc, left_label, bold=False, size=8)
    add_para(lc, left_val, bold=True, size=10)
    add_para(rc, right_label, bold=False, size=8)
    add_para(rc, right_val, bold=True, size=10)

fill_a(0, "NOM DU MEMBRE / PATIENT", lookup["NOM DU MEMBRE / PATIENT"], "NUMÉRO DE MEMBRE / CARTE", lookup["NUMÉRO DE MEMBRE / CARTE"])

lc2 = at.rows[1].cells[0]
add_para(lc2, "DATE DE NAISSANCE", bold=False, size=8)
add_para(lc2, lookup["DATE DE NAISSANCE"], bold=True, size=10)

rc2 = at.rows[1].cells[1]
add_para(rc2, "SEXE", bold=False, size=8)
sm = lookup["SEXE - M"]
sf = lookup["SEXE - F"]
add_para(rc2, ("☑" if "checked" in sm else "☐") + " M", bold=False, size=10)
add_para(rc2, ("☑" if "checked" in sf else "☐") + " F", bold=False, size=10)

fill_a(2, "COMPAGNIE D'ASSURANCE / TPA", lookup["COMPAGNIE D'ASSURANCE / TPA"], "TITULAIRE DE LA POLICE", lookup["TITULAIRE DE LA POLICE"])
fill_a(3, "CPR / PASSEPORT", lookup["CPR / PASSEPORT"], "TÉLÉPHONE DU MEMBRE", lookup["TÉLÉPHONE DU MEMBRE"])

doc.add_paragraph()

# SECTION B
section_header(doc, lookup["SECTION B - INFORMATIONS MÉDICALES"])

bt = doc.add_table(rows=1, cols=2)
bt.width = Cm(18)
set_table_border(bt)
bt.columns[0].width = Cm(9)
bt.columns[1].width = Cm(9)

blc = bt.rows[0].cells[0]
brc = bt.rows[0].cells[1]

add_para(blc, lookup["Type de demande"], bold=True, size=10)
hosp = lookup["Hospitalisation"]
amb = lookup["Ambulatoire"]
urg = lookup["Urgence"]
add_para(blc, ("☑" if "checked" in hosp else "☐") + " Hospitalisation", size=9)
add_para(blc, ("☑" if "checked" in amb else "☐") + " Ambulatoire", size=9)
add_para(blc, ("☑" if "checked" in urg else "☐") + " Urgence", size=9)

blc.add_paragraph()
ft = OxmlElement("w:tbl")
details_t = doc.add_table(rows=4, cols=2)
details_t.width = Cm(8.5)
set_table_border(details_t)
dfields = [
    ("Champ", "Détails"),
    ("Date du traitement", lookup["Date du traitement"]),
    ("Fournisseur / clinique", lookup["Fournisseur / clinique"]),
    ("Médecin traitant", lookup["Médecin traitant"]),
]
for i, (k, v) in enumerate(dfields):
    details_t.rows[i].cells[0].paragraphs[0].add_run(k).font.size = Pt(8 if i > 0 else 9)
    r2 = details_t.rows[i].cells[1].paragraphs[0].add_run(v)
    r2.font.size = Pt(9)
    if i == 0:
        details_t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        r2.bold = True

nr_row = doc.add_table(rows=1, cols=2)
nr_row.width = Cm(8.5)
set_table_border(nr_row)
nr_row.rows[0].cells[0].paragraphs[0].add_run("Numéro de dossier médical").font.size = Pt(8)
nr_row.rows[0].cells[1].paragraphs[0].add_run(lookup["Numéro de dossier médical"]).font.size = Pt(9)

blc._tc.append(details_t._tbl)
blc._tc.append(nr_row._tbl)

add_para(brc, lookup["Conditions et signes vitaux"], bold=True, size=10)

cond_items = [
    ("Affection préexistante", "Affection préexistante - Oui", "Affection préexistante - Non"),
    ("Accident du travail", "Accident du travail - Oui", "Accident du travail - Non"),
    ("Maternité", "Maternité - Oui", "Maternité - Non"),
]
for label, k_oui, k_non in cond_items:
    add_para(brc, label, bold=False, size=9)
    v_oui = lookup[k_oui]
    v_non = lookup[k_non]
    add_para(brc, ("☑" if "checked" in v_oui else "☐") + " Oui", size=9)
    add_para(brc, ("☑" if "checked" in v_non else "☐") + " Non", size=9)

add_para(brc, "Tension / pouls / température", bold=False, size=8)
add_para(brc, lookup["Tension / pouls / température"], bold=False, size=9)

doc.add_paragraph()

mt = doc.add_table(rows=1, cols=2)
mt.width = Cm(18)
set_table_border(mt)
mlc = mt.rows[0].cells[0]
mrc = mt.rows[0].cells[1]

add_para(mlc, "MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS", bold=True, size=9)
add_para(mlc, lookup["MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS"], bold=False, size=9)

add_para(mrc, "DIAGNOSTIC FINAL", bold=True, size=9)
diag_val = lookup["DIAGNOSTIC FINAL"]
for line in diag_val.split("\n"):
    add_para(mrc, line, bold=False, size=9)

doc.add_paragraph()

# SECTION C
section_header(doc, lookup["SECTION C - TRAITEMENT, INVESTIGATIONS ET COÛT ESTIMÉ"])

ct = doc.add_table(rows=5, cols=5)
ct.width = Cm(18)
set_table_border(ct)

headers_c = ["Traitement / service", "Catégorie", "Qté.", "Coût brut", "Justification médicale"]
for i, h in enumerate(headers_c):
    p = ct.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    set_cell_bg(ct.rows[0].cells[i], "E8E8E8")

services = [
    ("Traitement / service - Consultation ORL", "Consultation ORL"),
    ("Traitement / service - Injection Lincocin", "Injection Lincocin"),
    ("Traitement / service - Injection Depo-Medrol", "Injection Depo-Medrol"),
]

for ri, (k, name) in enumerate(services, 1):
    val = lookup[k]
    parts = {}
    for part in val.split(" | "):
        kp, vp = part.split(" : ", 1)
        parts[kp.strip()] = vp.strip()
    ct.rows[ri].cells[0].paragraphs[0].add_run(name).font.size = Pt(9)
    ct.rows[ri].cells[1].paragraphs[0].add_run(parts.get("Catégorie", "")).font.size = Pt(9)
    ct.rows[ri].cells[2].paragraphs[0].add_run(parts.get("Qté", "")).font.size = Pt(9)
    ct.rows[ri].cells[3].paragraphs[0].add_run(parts.get("Coût brut", "")).font.size = Pt(9)
    ct.rows[ri].cells[4].paragraphs[0].add_run(parts.get("Justification médicale", "")).font.size = Pt(9)

total_val = lookup["Coût prévu / total brut"]
total_parts = total_val.split(" | ")
total_cost = total_parts[0] if total_parts else ""
total_note = total_parts[1] if len(total_parts) > 1 else ""

last_row = ct.rows[4]
mc = last_row.cells[0]
p = mc.paragraphs[0]
run = p.add_run("Coût prévu / total brut")
run.bold = True
run.font.size = Pt(9)
last_row.cells[3].paragraphs[0].add_run(total_cost).font.size = Pt(9)
p4 = last_row.cells[4].paragraphs[0]
r4 = p4.add_run(total_note)
r4.bold = True
r4.font.size = Pt(9)

doc.add_paragraph()

# DECLARATIONS
dt = doc.add_table(rows=1, cols=2)
dt.width = Cm(18)
set_table_border(dt)

dlc = dt.rows[0].cells[0]
drc = dt.rows[0].cells[1]

add_para(dlc, "Déclaration du membre", bold=True, size=10)
add_para(dlc, lookup["Déclaration du membre"], bold=False, size=8)
dlc.add_paragraph()
add_para(dlc, "_" * 45, bold=False, size=9)
add_para(dlc, lookup["Signature du membre / date"] or "Signature du membre / date", bold=False, size=8)

add_para(drc, "Déclaration du prestataire médical", bold=True, size=10)
add_para(drc, lookup["Déclaration du prestataire médical"], bold=False, size=8)
drc.add_paragraph()
add_para(drc, "_" * 45, bold=False, size=9)
add_para(drc, lookup["Nom du médecin, signature et cachet"] or "Nom du médecin, signature et cachet", bold=False, size=8)

doc.add_paragraph()

fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
frun = fp.add_run("Formulaire claim médical - RO | pag. 2/3          Document modèle avec données anonymisées")
frun.font.size = Pt(7)
frun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)