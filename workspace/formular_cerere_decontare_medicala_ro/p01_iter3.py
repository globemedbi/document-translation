import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p01.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_01.docx"
DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.2)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_table_borders(t):
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_borders(t, color="CCCCCC"):
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def cp(p, size=9, bold=False, color=None, italic=False):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

def add_para(cell, text, size=9, bold=False, color=None, align=None, italic=False, space_before=0, space_after=0):
    if len(cell.paragraphs) > 0 and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def set_col_width(t, col_idx, width_cm):
    for row in t.rows:
        row.cells[col_idx].width = Cm(width_cm)

def merge_row(t, row_idx):
    row = t.rows[row_idx]
    row.cells[0].merge(row.cells[-1])
    return row.cells[0]

# ── BAND 1: HEADER ──────────────────────────────────────────────────────────
t1 = doc.add_table(rows=1, cols=2)
t1.width = Cm(18)
remove_table_borders(t1)

c_left = t1.rows[0].cells[0]
c_left.width = Cm(11)
c_right = t1.rows[0].cells[1]
c_right.width = Cm(7)

# MC logo + title
p_mc = c_left.paragraphs[0]
p_mc.clear()
run_mc = p_mc.add_run("MC  ")
run_mc.font.size = Pt(14)
run_mc.font.bold = True
run_mc.font.color.rgb = RGBColor(0x1A, 0x7A, 0x7A)
run_title = p_mc.add_run(lookup["Reçu médical et récapitulatif des coûts"])
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
p_mc.paragraph_format.space_before = Pt(2)
p_mc.paragraph_format.space_after = Pt(2)

p_sub = c_left.add_paragraph()
run_sub = p_sub.add_run(lookup["Dossier de remboursement - services ORL, soins primaires / ambulatoire"])
run_sub.font.size = Pt(8)
run_sub.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(0)

# right meta table
meta_t = doc.add_table(rows=4, cols=2)
meta_t.width = Cm(7)
set_table_borders(meta_t, "CCCCCC")

meta_labels = ["N° de dossier", "Date", "Devise", "Statut"]
meta_vals = [lookup["N° de dossier"], lookup["Date"], lookup["Devise"], lookup["Statut"]]
for i, (lbl, val) in enumerate(zip(meta_labels, meta_vals)):
    lc = meta_t.rows[i].cells[0]
    vc = meta_t.rows[i].cells[1]
    lc.width = Cm(3)
    vc.width = Cm(4)
    add_para(lc, lbl, size=8, bold=True, color="333333")
    add_para(vc, val, size=8, bold=True, color="1A1A2E", align=WD_ALIGN_PARAGRAPH.RIGHT)

# move meta table into right cell
meta_xml = meta_t._tbl
c_right._tc.append(copy.deepcopy(meta_xml))
meta_xml.getparent().remove(meta_xml)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── BAND 2: PURPOSE BANNER ───────────────────────────────────────────────────
t2 = doc.add_table(rows=1, cols=1)
t2.width = Cm(18)
set_table_borders(t2, "B0C4DE")
c2 = t2.rows[0].cells[0]
set_cell_bg(c2, "EEF4FB")
p2 = c2.paragraphs[0]
p2.clear()
run_label = p2.add_run("Objet du document : ")
run_label.font.size = Pt(9)
run_label.font.bold = True
run_label.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run_text = p2.add_run(lookup["Objet du document"])
run_text.font.size = Pt(9)
run_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after = Pt(4)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── BAND 3: PROVIDER / INSURANCE TWO-COLUMN ──────────────────────────────────
t3 = doc.add_table(rows=1, cols=2)
t3.width = Cm(18)
remove_table_borders(t3)

c3l = t3.rows[0].cells[0]
c3l.width = Cm(9)
c3r = t3.rows[0].cells[1]
c3r.width = Cm(9)

# inner left table
tl = doc.add_table(rows=5, cols=2)
tl.width = Cm(8.8)
set_table_borders(tl, "CCCCCC")

add_para(tl.rows[0].cells[0], lookup["Prestataire médical"], size=9, bold=True, color="1A1A2E")
tl.rows[0].cells[0].merge(tl.rows[0].cells[1])

prov_labels = ["Clinique", "Spécialité", "Adresse", "Téléphone"]
prov_vals = [lookup["Clinique"], lookup["Spécialité"], lookup["Adresse"], lookup["Téléphone"]]
for i, (lbl, val) in enumerate(zip(prov_labels, prov_vals)):
    lc = tl.rows[i+1].cells[0]
    vc = tl.rows[i+1].cells[1]
    lc.width = Cm(2.8)
    vc.width = Cm(6)
    add_para(lc, lbl, size=9, bold=True, color="1A7A7A")
    add_para(vc, val, size=9, color="1A1A2E")

tl_xml = tl._tbl
c3l._tc.append(copy.deepcopy(tl_xml))
tl_xml.getparent().remove(tl_xml)

# inner right table
tr = doc.add_table(rows=5, cols=2)
tr.width = Cm(8.8)
set_table_borders(tr, "CCCCCC")

add_para(tr.rows[0].cells[0], lookup["Assurance / gestionnaire de sinistre"], size=9, bold=True, color="1A1A2E")
tr.rows[0].cells[0].merge(tr.rows[0].cells[1])

ins_labels = ["Compagnie / TPA", "Titulaire de la police", "Membre", "Type de demande"]
ins_keys = ["Compagnie / TPA", "Titulaire de la police", "Membre", "Type de demande"]
ins_vals = [lookup[k] for k in ins_keys]
for i, (lbl, val) in enumerate(zip(ins_labels, ins_vals)):
    lc = tr.rows[i+1].cells[0]
    vc = tr.rows[i+1].cells[1]
    lc.width = Cm(3.2)
    vc.width = Cm(5.6)
    add_para(lc, lbl, size=9, bold=True, color="1A7A7A")
    add_para(vc, val, size=9, color="1A1A2E")

tr_xml = tr._tbl
c3r._tc.append(copy.deepcopy(tr_xml))
tr_xml.getparent().remove(tr_xml)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── BAND 4: FINANCIAL SUMMARY HEADER ─────────────────────────────────────────
t4h = doc.add_table(rows=1, cols=1)
t4h.width = Cm(18)
c4h = t4h.rows[0].cells[0]
set_cell_bg(c4h, "1A7A7A")
p4h = add_para(c4h, lookup["SYNTHÈSE FINANCIÈRE"], size=11, bold=True, color="FFFFFF")
p4h.paragraph_format.space_before = Pt(3)
p4h.paragraph_format.space_after = Pt(3)

# ── BAND 4: THREE TOTALS BOXES ────────────────────────────────────────────────
t4s = doc.add_table(rows=1, cols=3)
t4s.width = Cm(18)
set_table_borders(t4s, "CCCCCC")

boxes = [
    ("TOTAL BRUT", lookup["TOTAL BRUT"], "Services facturés par la clinique"),
    ("RÉDUCTIONS CONTRACTUELLES", lookup["RÉDUCTIONS CONTRACTUELLES"], "Appliquées sur consultation et injections"),
    ("MONTANT NET DEMANDÉ", lookup["MONTANT NET DEMANDÉ"], "Après réduction et déduction/quote-part"),
]
for i, (hdr, val, sub) in enumerate(boxes):
    c = t4s.rows[0].cells[i]
    set_cell_bg(c, "F5F9FF")
    add_para(c, hdr, size=8, bold=True, color="555555", space_before=4)
    add_para(c, val, size=16, bold=True, color="1A1A2E", space_after=0)
    add_para(c, sub, size=8, color="777777", space_after=4)

# ── BAND 4: SERVICES TABLE ───────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(2)

t5 = doc.add_table(rows=1, cols=6)
t5.width = Cm(18)
set_table_borders(t5, "CCCCCC")

headers5 = ["Service / traitement", "Qté", "Tarif brut", "Réduction", "Valeur éligible", "Observations"]
widths5 = [5.5, 1.2, 2.2, 2.2, 2.5, 2.9]
for i, (hdr, w) in enumerate(zip(headers5, widths5)):
    c = t5.rows[0].cells[i]
    c.width = Cm(w)
    set_cell_bg(c, "E8E8E8")
    add_para(c, hdr, size=9, bold=True, color="1A1A2E")

svc_data_raw = lookup["Tableau des services"]
svc_rows = [line.strip() for line in svc_data_raw.strip().split("\n")]
for row_str in svc_rows:
    parts = [p.strip() for p in row_str.split("|")]
    while len(parts) < 6:
        parts.append("")
    r = t5.add_row()
    is_total = parts[0].strip().lower() == "total"
    for i, val in enumerate(parts[:6]):
        c = r.cells[i]
        c.width = Cm(widths5[i])
        add_para(c, val, size=9, bold=is_total, color="1A1A2E")
        if is_total:
            set_cell_bg(c, "F0F0F0")

# Deduction row
r_ded = t5.add_row()
for i in range(6):
    r_ded.cells[i].width = Cm(widths5[i])
c_ded_label = r_ded.cells[0]
c_ded_label.merge(r_ded.cells[1])
c_ded_label.merge(r_ded.cells[2])
c_ded_label.merge(r_ded.cells[3])
add_para(c_ded_label, lookup["Déduction / quote-part du membre"], size=9, bold=False, color="1A1A2E", align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(r_ded.cells[4], lookup["Déduction / quote-part du membre".replace("Déduction / quote-part du membre", "Déduction / quote-part du membre")], size=9, bold=False, color="1A1A2E")
r_ded.cells[4].paragraphs[0].clear()
run_d = r_ded.cells[4].paragraphs[0].add_run("8.20 BD")
run_d.font.size = Pt(9)
run_d.font.bold = True

# Net row
r_net = t5.add_row()
for i in range(6):
    r_net.cells[i].width = Cm(widths5[i])
c_net_label = r_net.cells[0]
c_net_label.merge(r_net.cells[1])
c_net_label.merge(r_net.cells[2])
c_net_label.merge(r_net.cells[3])
add_para(c_net_label, lookup["Montant net à rembourser / à percevoir"], size=9, bold=False, color="1A1A2E", align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(r_net.cells[4], lookup["Montant net à rembourser / à percevoir"], size=9, bold=False, color="1A1A2E")
r_net.cells[4].paragraphs[0].clear()
run_n = r_net.cells[4].paragraphs[0].add_run("19.80 BD")
run_n.font.size = Pt(9)
run_n.font.bold = True
set_cell_bg(r_net.cells[4], "E8F5E9")
add_para(r_net.cells[5], "Net", size=9, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r_net.cells[5], "1A7A7A")

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── BAND 5: SIGNATURE BLOCK ──────────────────────────────────────────────────
t6 = doc.add_table(rows=1, cols=2)
t6.width = Cm(18)
set_table_borders(t6, "CCCCCC")

c6l = t6.rows[0].cells[0]
c6r = t6.rows[0].cells[1]
c6l.width = Cm(9)
c6r.width = Cm(9)

add_para(c6l, lookup["Confirmation du prestataire médical"], size=9, bold=True, color="1A1A2E", space_before=4)
add_para(c6l, lookup["Confirmation du prestataire médical - texte"], size=8, color="333333")
add_para(c6l, "", size=8)
add_para(c6l, "_" * 40, size=8, color="555555")
add_para(c6l, lookup["Nom, signature et cachet"] if lookup["Nom, signature et cachet"] else "Nom, signature et cachet", size=8, color="777777", italic=True, space_after=4)

add_para(c6r, lookup["Confirmation du membre / patient"], size=9, bold=True, color="1A1A2E", space_before=4)
add_para(c6r, lookup["Confirmation du membre / patient - texte"], size=8, color="333333")
add_para(c6r, "", size=8)
add_para(c6r, "_" * 40, size=8, color="555555")
add_para(c6r, lookup["Nom et signature du membre"] if lookup["Nom et signature du membre"] else "Nom et signature du membre", size=8, color="777777", italic=True, space_after=4)

# ── FOOTER ───────────────────────────────────────────────────────────────────
p_foot = doc.add_paragraph()
p_foot.paragraph_format.space_before = Pt(12)
run_fl = p_foot.add_run("Formulaire demande médicale - RO | pag. 1/3")
run_fl.font.size = Pt(7)
run_fl.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run_fr = p_foot.add_run("          Document modèle avec données anonymisées")
run_fr.font.size = Pt(7)
run_fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)