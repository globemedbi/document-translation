import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/formular_cerere_decontare_medicala_ro/data_p03.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/formular_cerere_decontare_medicala_ro/page_03.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(k):
    return lookup.get(k, k)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        val = kwargs.get(edge, {})
        tag.set(qn("w:val"), val.get("val", "none"))
        tag.set(qn("w:sz"), str(val.get("sz", 0)))
        tag.set(qn("w:color"), val.get("color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_table_border(table, hex_color="CCCCCC", sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(sz))
        tag.set(qn("w:color"), hex_color)
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def no_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:color"), "auto")
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def add_paragraph_in_cell(cell, text, bold=False, size=9, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    if len(cell.paragraphs) == 0:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_run_to_para(para, text, bold=False, size=9, color="000000"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.0)
section.right_margin = Cm(1.0)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(1.0)

for para in doc.paragraphs:
    para.clear()

NAVY = "1B2A4A"
TEAL = "2B5C6B"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
DARK_GRAY = "555555"
MID_GRAY = "888888"
BLACK = "222222"
GREEN_BADGE = "27AE60"
BORDER_GRAY = "CCCCCC"

def add_section_header(doc, text, color=TEAL):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    no_table_border(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    set_cell_margins(cell, top=60, start=100, bottom=60, end=100)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return tbl

def small_spacer(doc, size=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)

# ── SECTION 1: HEADER ──────────────────────────────────────────────────────────
header_tbl = doc.add_table(rows=1, cols=2)
no_table_border(header_tbl)
set_table_border(header_tbl, hex_color=BORDER_GRAY, sz=4)

# Left cell: MC logo + title
left_cell = header_tbl.cell(0, 0)
set_cell_margins(left_cell, top=80, start=80, bottom=80, end=80)
set_col_width(left_cell, 12.5)

logo_title_tbl = left_cell.add_table(rows=1, cols=2)
no_table_border(logo_title_tbl)
logo_cell = logo_title_tbl.cell(0, 0)
set_col_width(logo_cell, 2.0)
set_cell_bg(logo_cell, NAVY)
set_cell_margins(logo_cell, top=60, start=60, bottom=60, end=60)
p_logo = logo_cell.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(0)
p_logo.paragraph_format.space_after = Pt(0)
r_logo = p_logo.add_run("MC")
r_logo.bold = True
r_logo.font.size = Pt(14)
r_logo.font.color.rgb = RGBColor.from_string(WHITE)

title_cell = logo_title_tbl.cell(0, 1)
set_col_width(title_cell, 10.5)
set_cell_margins(title_cell, top=40, start=100, bottom=40, end=40)
p_title = title_cell.paragraphs[0]
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(2)
r_title = p_title.add_run(get_value("Rapport d'approbation et détails de couverture"))
r_title.bold = True
r_title.font.size = Pt(14)
r_title.font.color.rgb = RGBColor.from_string(NAVY)

p_sub = title_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(0)
r_sub = p_sub.add_run("Résumé pour la compagnie d'assurance - éligibilité, documents et conditions de paiement")
r_sub.bold = False
r_sub.font.size = Pt(7.5)
r_sub.font.color.rgb = RGBColor.from_string(MID_GRAY)

# Right cell: ref box
right_cell = header_tbl.cell(0, 1)
set_col_width(right_cell, 6.5)
set_cell_bg(right_cell, LIGHT_GRAY)
set_cell_margins(right_cell, top=60, start=80, bottom=60, end=80)

ref_tbl = right_cell.add_table(rows=4, cols=2)
no_table_border(ref_tbl)

ref_labels = [
    "Réf. approbation",
    "Date d'approbation",
    "Valable jusqu'au",
    "Statut"
]
ref_keys = [
    "Réf. approbation",
    "Date d'approbation",
    "Valable jusqu'au",
    "Statut"
]

for i, (label, key) in enumerate(zip(ref_labels, ref_keys)):
    lc = ref_tbl.cell(i, 0)
    vc = ref_tbl.cell(i, 1)
    set_col_width(lc, 3.2)
    set_col_width(vc, 3.0)
    set_cell_margins(lc, top=20, start=40, bottom=20, end=20)
    set_cell_margins(vc, top=20, start=20, bottom=20, end=40)
    p_l = lc.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(0)
    p_l.paragraph_format.space_after = Pt(0)
    r_l = p_l.add_run(label)
    r_l.bold = False
    r_l.font.size = Pt(8)
    r_l.font.color.rgb = RGBColor.from_string(MID_GRAY)
    p_v = vc.paragraphs[0]
    p_v.paragraph_format.space_before = Pt(0)
    p_v.paragraph_format.space_after = Pt(0)
    r_v = p_v.add_run(get_value(key))
    r_v.bold = True
    r_v.font.size = Pt(9)
    r_v.font.color.rgb = RGBColor.from_string(DARK_GRAY)

small_spacer(doc, 3)

# ── SECTION 2: DONNÉES DE RÉFÉRENCE ───────────────────────────────────────────
add_section_header(doc, get_value("DONNÉES DE RÉFÉRENCE"), color=TEAL)

ref_data_tbl = doc.add_table(rows=3, cols=2)
set_table_border(ref_data_tbl, hex_color=BORDER_GRAY, sz=4)

ref_fields_left = [
    ("NUMÉRO INDIVIDUEL", "NUMÉRO INDIVIDUEL"),
    ("DATE DE NAISSANCE", "DATE DE NAISSANCE"),
    ("MÉDECIN / PRESTATAIRE", "MÉDECIN / PRESTATAIRE"),
]
ref_fields_right = [
    ("NOM DU TITULAIRE DE LA CARTE", "NOM DU TITULAIRE DE LA CARTE"),
    ("COMPAGNIE D'ASSURANCE", "COMPAGNIE D'ASSURANCE"),
    ("SPÉCIALITÉ", "SPÉCIALITÉ"),
]

for i in range(3):
    lc = ref_data_tbl.cell(i, 0)
    rc = ref_data_tbl.cell(i, 1)
    set_col_width(lc, 9.5)
    set_col_width(rc, 9.5)
    set_cell_bg(lc, WHITE)
    set_cell_bg(rc, WHITE)
    set_cell_margins(lc, top=60, start=100, bottom=60, end=80)
    set_cell_margins(rc, top=60, start=100, bottom=60, end=80)

    label_l, key_l = ref_fields_left[i]
    label_r, key_r = ref_fields_right[i]

    p_l = lc.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(0)
    p_l.paragraph_format.space_after = Pt(1)
    r_ll = p_l.add_run(label_l)
    r_ll.bold = False
    r_ll.font.size = Pt(7.5)
    r_ll.font.color.rgb = RGBColor.from_string(MID_GRAY)
    p_lv = lc.add_paragraph()
    p_lv.paragraph_format.space_before = Pt(0)
    p_lv.paragraph_format.space_after = Pt(0)
    r_lv = p_lv.add_run(get_value(key_l))
    r_lv.bold = True
    r_lv.font.size = Pt(9)
    r_lv.font.color.rgb = RGBColor.from_string(BLACK)

    p_r = rc.paragraphs[0]
    p_r.paragraph_format.space_before = Pt(0)
    p_r.paragraph_format.space_after = Pt(1)
    r_rl = p_r.add_run(label_r)
    r_rl.bold = False
    r_rl.font.size = Pt(7.5)
    r_rl.font.color.rgb = RGBColor.from_string(MID_GRAY)
    p_rv = rc.add_paragraph()
    p_rv.paragraph_format.space_before = Pt(0)
    p_rv.paragraph_format.space_after = Pt(0)
    r_rv = p_rv.add_run(get_value(key_r))
    r_rv.bold = True
    r_rv.font.size = Pt(9)
    r_rv.font.color.rgb = RGBColor.from_string(BLACK)

small_spacer(doc, 3)

# ── SECTION 3: SERVICES APPROUVÉS ─────────────────────────────────────────────
add_section_header(doc, get_value("SERVICES APPROUVÉS"), color=TEAL)

svc_tbl = doc.add_table(rows=4, cols=6)
set_table_border(svc_tbl, hex_color=BORDER_GRAY, sz=4)

col_widths = [3.5, 5.5, 1.8, 1.8, 2.2, 4.2]
headers_svc = ["Catégorie de service", "Élément approuvé", "Qté\ncerută", "Qté\nappr.", "Statut", "Motif / notes"]

for j, (hdr, cw) in enumerate(zip(headers_svc, col_widths)):
    cell = svc_tbl.cell(0, j)
    set_col_width(cell, cw)
    set_cell_bg(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=60, start=80, bottom=60, end=80)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(hdr)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(DARK_GRAY)

rows_data = [
    [
        get_value("Services approuvés - Ligne 1 : Catégorie de service"),
        get_value("Services approuvés - Ligne 1 : Élément approuvé"),
        get_value("Services approuvés - Ligne 1 : Qté demandée"),
        get_value("Services approuvés - Ligne 1 : Qté approuvée"),
        get_value("Services approuvés - Ligne 1 : Statut"),
        get_value("Services approuvés - Ligne 1 : Motif / remarques"),
    ],
    [
        get_value("Services approuvés - Ligne 2 : Catégorie de service"),
        get_value("Services approuvés - Ligne 2 : Élément approuvé"),
        get_value("Services approuvés - Ligne 2 : Qté demandée"),
        get_value("Services approuvés - Ligne 2 : Qté approuvée"),
        get_value("Services approuvés - Ligne 2 : Statut"),
        get_value("Services approuvés - Ligne 2 : Motif / notes"),
    ],
    [
        get_value("Services approuvés - Ligne 3 : Catégorie de service"),
        get_value("Services approuvés - Ligne 3 : Élément approuvé"),
        get_value("Services approuvés - Ligne 3 : Qté demandée"),
        get_value("Services approuvés - Ligne 3 : Qté approuvée"),
        get_value("Services approuvés - Ligne 3 : Statut"),
        get_value("Services approuvés - Ligne 3 : Motif / notes"),
    ],
]

for i, row_vals in enumerate(rows_data):
    for j, val in enumerate(row_vals):
        cell = svc_tbl.cell(i + 1, j)
        set_col_width(cell, col_widths[j])
        set_cell_bg(cell, WHITE)
        set_cell_margins(cell, top=60, start=80, bottom=60, end=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [2, 3] else WD_ALIGN_PARAGRAPH.LEFT
        if j == 4:
            badge_tbl = cell.add_table(rows=1, cols=1)
            no_table_border(badge_tbl)
            bc = badge_tbl.cell(0, 0)
            set_cell_bg(bc, GREEN_BADGE)
            set_cell_margins(bc, top=30, start=60, bottom=30, end=60)
            bp = bc.paragraphs[0]
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(0)
            br = bp.add_run(val)
            br.bold = True
            br.font.size = Pt(8)
            br.font.color.rgb = RGBColor.from_string(WHITE)
        else:
            r = p.add_run(val)
            r.bold = False
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string(BLACK)

small_spacer(doc, 3)

# ── SECTION 4: TWO-COLUMN BLOCK (Documents + Diagnostic) ──────────────────────
two_col_tbl = doc.add_table(rows=1, cols=2)
no_table_border(two_col_tbl)

left_panel = two_col_tbl.cell(0, 0)
right_panel = two_col_tbl.cell(0, 1)
set_col_width(left_panel, 9.5)
set_col_width(right_panel, 9.5)
set_cell_margins(left_panel, top=0, start=0, bottom=0, end=60)
set_cell_margins(right_panel, top=0, start=60, bottom=0, end=0)

# Left: Documents nécessaires
docs_outer = left_panel.add_table(rows=2, cols=1)
set_table_border(docs_outer, hex_color=BORDER_GRAY, sz=4)

docs_hdr_cell = docs_outer.cell(0, 0)
set_cell_bg(docs_hdr_cell, TEAL)
set_cell_margins(docs_hdr_cell, top=60, start=100, bottom=60, end=100)
p_dh = docs_hdr_cell.paragraphs[0]
p_dh.paragraph_format.space_before = Pt(0)
p_dh.paragraph_format.space_after = Pt(0)
r_dh = p_dh.add_run(get_value("DOCUMENTS NÉCESSAIRES POUR LA SOUMISSION DE LA DEMANDE"))
r_dh.bold = True
r_dh.font.size = Pt(9)
r_dh.font.color.rgb = RGBColor.from_string(WHITE)

docs_body_cell = docs_outer.cell(1, 0)
set_cell_bg(docs_body_cell, WHITE)
set_cell_margins(docs_body_cell, top=80, start=100, bottom=80, end=100)

doc_items = [
    get_value("Documents - 1"),
    get_value("Documents - 2"),
    get_value("Documents - 3"),
    get_value("Documents - 4"),
    get_value("Documents - 5"),
    get_value("Documents - 6"),
]

for idx, item in enumerate(doc_items):
    if idx == 0:
        p_item = docs_body_cell.paragraphs[0]
    else:
        p_item = docs_body_cell.add_paragraph()
    p_item.paragraph_format.space_before = Pt(1)
    p_item.paragraph_format.space_after = Pt(1)
    r_item = p_item.add_run(f"{idx + 1}. {item}")
    r_item.bold = False
    r_item.font.size = Pt(8.5)
    r_item.font.color.rgb = RGBColor.from_string(BLACK)

# Right: Diagnostic et informations médicales
diag_outer = right_panel.add_table(rows=2, cols=1)
set_table_border(diag_outer, hex_color=BORDER_GRAY, sz=4)

diag_hdr_cell = diag_outer.cell(0, 0)
set_cell_bg(diag_hdr_cell, TEAL)
set_cell_margins(diag_hdr_cell, top=60, start=100, bottom=60, end=100)
p_diagh = diag_hdr_cell.paragraphs[0]
p_diagh.paragraph_format.space_before = Pt(0)
p_diagh.paragraph_format.space_after = Pt(0)
r_diagh = p_diagh.add_run(get_value("DIAGNOSTIC ET INFORMATIONS MÉDICALES"))
r_diagh.bold = True
r_diagh.font.size = Pt(9)
r_diagh.font.color.rgb = RGBColor.from_string(WHITE)

diag_body_cell = diag_outer.cell(1, 0)
set_cell_bg(diag_body_cell, WHITE)
set_cell_margins(diag_body_cell, top=40, start=40, bottom=40, end=40)

diag_inner = diag_body_cell.add_table(rows=6, cols=2)
set_table_border(diag_inner, hex_color=BORDER_GRAY, sz=4)

diag_col_hdr = ["Champ", "Valeur"]
for j, hdr in enumerate(diag_col_hdr):
    c = diag_inner.cell(0, j)
    set_cell_bg(c, LIGHT_GRAY)
    set_cell_margins(c, top=40, start=80, bottom=40, end=80)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(hdr)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(DARK_GRAY)

diag_rows = [
    ("Diagnostic principal", "Diagnostic principal"),
    ("Date de la visite", "Date de la visite"),
    ("Tension / pouls / température", "Tension / pouls / température"),
    ("Durée de la maladie", "Durée de la maladie"),
    ("Fichiers téléchargés", "Fichiers téléchargés"),
]

for i, (field_label, field_key) in enumerate(diag_rows):
    lc = diag_inner.cell(i + 1, 0)
    vc = diag_inner.cell(i + 1, 1)
    set_cell_bg(lc, WHITE)
    set_cell_bg(vc, WHITE)
    set_cell_margins(lc, top=40, start=80, bottom=40, end=80)
    set_cell_margins(vc, top=40, start=80, bottom=40, end=80)
    pl = lc.paragraphs[0]
    pl.paragraph_format.space_before = Pt(0)
    pl.paragraph_format.space_after = Pt(0)
    rl = pl.add_run(field_label)
    rl.bold = False
    rl.font.size = Pt(8.5)
    rl.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    pv = vc.paragraphs[0]
    pv.paragraph_format.space_before = Pt(0)
    pv.paragraph_format.space_after = Pt(0)
    rv = pv.add_run(get_value(field_key))
    rv.bold = False
    rv.font.size = Pt(8.5)
    rv.font.color.rgb = RGBColor.from_string(BLACK)

small_spacer(doc, 3)

# ── SECTION 5: CONDITIONS D'ÉVALUATION ET DE PAIEMENT ─────────────────────────
add_section_header(doc, get_value("CONDITIONS D'ÉVALUATION ET DE PAIEMENT"), color=TEAL)

cond_tbl = doc.add_table(rows=1, cols=1)
set_table_border(cond_tbl, hex_color=BORDER_GRAY, sz=4)
cond_cell = cond_tbl.cell(0, 0)
set_cell_bg(cond_cell, WHITE)
set_cell_margins(cond_cell, top=80, start=100, bottom=80, end=100)

cond_texts = [
    get_value("Conditions - paragraphe 1"),
    get_value("Conditions - paragraphe 2"),
]

for idx, txt in enumerate(cond_texts):
    if idx == 0:
        pc = cond_cell.paragraphs[0]
    else:
        pc = cond_cell.add_paragraph()
    pc.paragraph_format.space_before = Pt(0)
    pc.paragraph_format.space_after = Pt(4)
    rc = pc.add_run(txt)
    rc.bold = False
    rc.font.size = Pt(8.5)
    rc.font.color.rgb = RGBColor.from_string(BLACK)

p_vat = cond_cell.add_paragraph()
p_vat.paragraph_format.space_before = Pt(0)
p_vat.paragraph_format.space_after = Pt(0)
r_vat_label = p_vat.add_run("Note TVA : ")
r_vat_label.bold = True
r_vat_label.font.size = Pt(8.5)
r_vat_label.font.color.rgb = RGBColor.from_string(BLACK)
r_vat_text = p_vat.add_run(get_value("Note TVA"))
r_vat_text.bold = False
r_vat_text.font.size = Pt(8.5)
r_vat_text.font.color.rgb = RGBColor.from_string(BLACK)

small_spacer(doc, 3)

# ── SECTION 6: BOTTOM TWO-COLUMN (Usage interne + Contrôle qualité) ───────────
bottom_tbl = doc.add_table(rows=1, cols=2)
set_table_border(bottom_tbl, hex_color=BORDER_GRAY, sz=4)

int_cell = bottom_tbl.cell(0, 0)
qc_cell = bottom_tbl.cell(0, 1)
set_col_width(int_cell, 9.5)
set_col_width(qc_cell, 9.5)
set_cell_bg(int_cell, WHITE)
set_cell_bg(qc_cell, WHITE)
set_cell_margins(int_cell, top=80, start=100, bottom=80, end=80)
set_cell_margins(qc_cell, top=80, start=100, bottom=80, end=80)

p_int_hdr = int_cell.paragraphs[0]
p_int_hdr.paragraph_format.space_before = Pt(0)
p_int_hdr.paragraph_format.space_after = Pt(4)
r_int_hdr = p_int_hdr.add_run(get_value("Utilisation interne - compagnie d'assurance"))
r_int_hdr.bold = True
r_int_hdr.font.size = Pt(9)
r_int_hdr.font.color.rgb = RGBColor.from_string(DARK_GRAY)

inner_int_tbl = int_cell.add_table(rows=3, cols=2)
no_table_border(inner_int_tbl)

int_fields = [
    ("Décision", get_value("Décision")),
    ("Nr. approbation", get_value("Nr. approbation")),
    ("Commentaires", get_value("Commentaires")),
]

for i, (lbl, val) in enumerate(int_fields):
    lc = inner_int_tbl.cell(i, 0)
    vc = inner_int_tbl.cell(i, 1)
    set_col_width(lc, 3.5)
    set_col_width(vc, 5.5)
    set_cell_margins(lc, top=30, start=0, bottom=30, end=40)
    set_cell_margins(vc, top=30, start=0, bottom=30, end=0)
    pl = lc.paragraphs[0]
    pl.paragraph_format.space_before = Pt(0)
    pl.paragraph_format.space_after = Pt(0)
    rl = pl.add_run(lbl)
    rl.bold = True
    rl.font.size = Pt(8.5)
    rl.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    pv = vc.paragraphs[0]
    pv.paragraph_format.space_before = Pt(0)
    pv.paragraph_format.space_after = Pt(0)
    rv = pv.add_run(val)
    rv.bold = False
    rv.font.size = Pt(8.5)
    rv.font.color.rgb = RGBColor.from_string(BLACK)

p_qc_hdr = qc_cell.paragraphs[0]
p_qc_hdr.paragraph_format.space_before = Pt(0)
p_qc_hdr.paragraph_format.space_after = Pt(4)
r_qc_hdr = p_qc_hdr.add_run(get_value("Contrôle qualité claim"))
r_qc_hdr.bold = True
r_qc_hdr.font.size = Pt(9)
r_qc_hdr.font.color.rgb = RGBColor.from_string(DARK_GRAY)

p_qc_sub = qc_cell.add_paragraph()
p_qc_sub.paragraph_format.space_before = Pt(0)
p_qc_sub.paragraph_format.space_after = Pt(20)
r_qc_sub = p_qc_sub.add_run(get_value("Contrôle qualité - description"))
r_qc_sub.bold = False
r_qc_sub.font.size = Pt(8)
r_qc_sub.font.color.rgb = RGBColor.from_string(MID_GRAY)

p_sig_line = qc_cell.add_paragraph()
p_sig_line.paragraph_format.space_before = Pt(0)
p_sig_line.paragraph_format.space_after = Pt(2)
r_sig_line = p_sig_line.add_run("_" * 45)
r_sig_line.font.size = Pt(9)
r_sig_line.font.color.rgb = RGBColor.from_string(MID_GRAY)

p_sig_lbl = qc_cell.add_paragraph()
p_sig_lbl.paragraph_format.space_before = Pt(0)
p_sig_lbl.paragraph_format.space_after = Pt(0)
r_sig_lbl = p_sig_lbl.add_run(get_value("Nom de l'évaluateur / signature / date"))
r_sig_lbl.bold = False
r_sig_lbl.font.size = Pt(8)
r_sig_lbl.font.color.rgb = RGBColor.from_string(MID_GRAY)

small_spacer(doc, 3)

# ── FOOTER ─────────────────────────────────────────────────────────────────────
footer_tbl = doc.add_table(rows=1, cols=2)
no_table_border(footer_tbl)
fl = footer_tbl.cell(0, 0)
fr = footer_tbl.cell(0, 1)
set_col_width(fl, 10.0)
set_col_width(fr, 9.0)
set_cell_margins(fl, top=40, start=0, bottom=40, end=0)
set_cell_margins(fr, top=40, start=0, bottom=40, end=0)

p_fl = fl.paragraphs[0]
p_fl.paragraph_format.space_before = Pt(0)
p_fl.paragraph_format.space_after = Pt(0)
r_fl = p_fl.add_run("Formulaire claim médical - RO | pag. 3/3")
r_fl.font.size = Pt(7.5)
r_fl.font.color.rgb = RGBColor.from_string(MID_GRAY)

p_fr = fr.paragraphs[0]
p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_fr.paragraph_format.space_before = Pt(0)
p_fr.paragraph_format.space_after = Pt(0)
r_fr = p_fr.add_run("Document modèle avec données anonymisées")
r_fr.font.size = Pt(7.5)
r_fr.font.color.rgb = RGBColor.from_string(MID_GRAY)

doc.save(OUTPUT_PATH)