import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p02.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_02.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para_in_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=3, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def cb(val):
    return "☑" if "checked" in str(val) else "☐"

# HEADER TABLE
ht = doc.add_table(rows=1, cols=2)
ht.width = Cm(18)
from docx.shared import Inches
ht.columns[0].width = Cm(12)
ht.columns[1].width = Cm(6)

lc = ht.rows[0].cells[0]
lc.width = Cm(12)
lp = lc.paragraphs[0]
lp.paragraph_format.space_after = Pt(0)
r1 = lp.add_run("MC  ")
r1.bold = True
r1.font.size = Pt(14)
r1.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
r2 = lp.add_run("Formular de cerere de decontare medicală")
r2.bold = True
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

lp2 = lc.add_paragraph()
lp2.paragraph_format.space_after = Pt(0)
r3 = lp2.add_run("Îngrijire primară - se completează de membru și de medicul curant")
r3.italic = True
r3.font.size = Pt(8)
r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

rc = ht.rows[0].cells[1]
rc.width = Cm(6)
set_table_border_outer = lambda cell: None

info_pairs = [
    ("Nr. cerere", lookup["N° de demande"]),
    ("Ref. aprobare", lookup["Réf. d'approbation"]),
    ("Data tratamentului", lookup["Date du traitement"]),
    ("Tip", lookup["Type"]),
]

it = doc.add_table(rows=0, cols=0)
it = rc.add_table(4, 2)
it.width = Cm(6)
set_table_border(it)
for i, (lbl, val) in enumerate(info_pairs):
    row = it.rows[i]
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(3)
    para_in_cell(row.cells[0], lbl, bold=False, size=8)
    para_in_cell(row.cells[1], val, bold=True, size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

def section_header(text):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(18)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "1F5C99")
    p = para_in_cell(c, text, bold=True, size=10, color=(255, 255, 255))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

section_header("SECȚIUNEA A - DATELE MEMBRULUI ASIGURAT")

ta = doc.add_table(rows=4, cols=2)
ta.width = Cm(18)
set_table_border(ta)

def fill_cell_label_value(cell, label, value, label_size=7, value_size=9):
    cell.paragraphs[0].clear()
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run(label)
    r.bold = True
    r.font.size = Pt(label_size)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p2 = cell.add_paragraph(value)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(value_size)

fill_cell_label_value(ta.rows[0].cells[0], "NUME MEMBRU / PACIENT", lookup["NOM DU MEMBRE / PATIENT"])
fill_cell_label_value(ta.rows[0].cells[1], "NUMĂR MEMBRU / CARD", lookup["NUMÉRO DE MEMBRE / CARTE"])

dob_cell = ta.rows[1].cells[0]
fill_cell_label_value(dob_cell, "DATA NAȘTERII", lookup["DATE DE NAISSANCE"])

sex_cell = ta.rows[1].cells[1]
sex_cell.paragraphs[0].clear()
sp = sex_cell.paragraphs[0]
sp.paragraph_format.space_after = Pt(0)
sr = sp.add_run("SEX")
sr.bold = True
sr.font.size = Pt(7)
sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
p_m = sex_cell.add_paragraph(f"{cb(lookup['SEXE - M'])} M")
p_m.paragraph_format.space_after = Pt(0)
p_m.runs[0].font.size = Pt(9)
p_f = sex_cell.add_paragraph(f"{cb(lookup['SEXE - F'])} F")
p_f.paragraph_format.space_after = Pt(0)
p_f.runs[0].font.size = Pt(9)

fill_cell_label_value(ta.rows[2].cells[0], "COMPANIE DE ASIGURĂRI / TPA", lookup["COMPAGNIE D'ASSURANCE / TPA"])
fill_cell_label_value(ta.rows[2].cells[1], "DEȚINATOR POLIȚĂ", lookup["TITULAIRE DE LA POLICE"])
fill_cell_label_value(ta.rows[3].cells[0], "CPR / PAȘAPORT", lookup["CPR / PASSEPORT"])
fill_cell_label_value(ta.rows[3].cells[1], "TELEFON MEMBRU", lookup["TÉLÉPHONE DU MEMBRE"])

doc.add_paragraph().paragraph_format.space_after = Pt(2)

section_header("SECȚIUNEA B - INFORMAȚII MEDICALE")

tb = doc.add_table(rows=1, cols=2)
tb.width = Cm(18)

left_cell = tb.rows[0].cells[0]
left_cell.width = Cm(9)
right_cell = tb.rows[0].cells[1]
right_cell.width = Cm(9)

lp = left_cell.paragraphs[0]
lp.clear()
lr = lp.add_run("Tip solicitare")
lr.bold = True
lr.font.size = Pt(10)

req_types = [
    ("Hospitalisation", lookup["Hospitalisation"]),
    ("Ambulatoriu", lookup["Ambulatoire"]),
    ("Urgență", lookup["Urgence"]),
]
for label, val in req_types:
    p = left_cell.add_paragraph(f"{cb(val)} {label}")
    p.paragraph_format.space_after = Pt(1)
    p.runs[0].font.size = Pt(9)

detail_rows = [
    ("Câmp", "Detalii"),
    ("Data tratamentului", lookup["Date du traitement"]),
    ("Furnizor / clinică", lookup["Fournisseur / clinique"]),
    ("Medic curant", lookup["Médecin traitant"]),
    ("Număr fișă medicală", lookup["Numéro de dossier médical"]),
]

dt = left_cell.add_table(len(detail_rows), 2)
dt.width = Cm(8.5)
set_table_border(dt)
for i, (k, v) in enumerate(detail_rows):
    r = dt.rows[i]
    para_in_cell(r.cells[0], k, bold=(i == 0), size=8 if i > 0 else 9)
    para_in_cell(r.cells[1], v, bold=(i == 0), size=8 if i > 0 else 9)

rp = right_cell.paragraphs[0]
rp.clear()
rr = rp.add_run("Condiții și semne vitale")
rr.bold = True
rr.font.size = Pt(10)

conditions = [
    ("Afecțiune preexistentă", "Affection préexistante - Oui", "Affection préexistante - Non"),
    ("Accident de muncă", "Accident du travail - Oui", "Accident du travail - Non"),
    ("Maternitate", "Maternité - Oui", "Maternité - Non"),
]

ct = right_cell.add_table(len(conditions) + 1, 2)
ct.width = Cm(8.5)
set_table_border(ct)
for i, (lbl, k_yes, k_no) in enumerate(conditions):
    para_in_cell(ct.rows[i].cells[0], lbl, size=8)
    cv = ct.rows[i].cells[1]
    cv.paragraphs[0].clear()
    p_y = cv.paragraphs[0]
    p_y.paragraph_format.space_after = Pt(0)
    p_y.add_run(f"{cb(lookup[k_yes])} Da").font.size = Pt(8)
    p_n = cv.add_paragraph(f"{cb(lookup[k_no])} Nu")
    p_n.paragraph_format.space_after = Pt(0)
    p_n.runs[0].font.size = Pt(8)

para_in_cell(ct.rows[3].cells[0], "Tensiune / puls / temperatură", size=8)
para_in_cell(ct.rows[3].cells[1], lookup["Tension / pouls / température"], size=8)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

motiv_t = doc.add_table(rows=1, cols=2)
motiv_t.width = Cm(18)
set_table_border(motiv_t)

mc = motiv_t.rows[0].cells[0]
mc.paragraphs[0].clear()
mp1 = mc.paragraphs[0]
mp1.paragraph_format.space_after = Pt(2)
mr1 = mp1.add_run("MOTIV CONSULTAȚIE / SIMPTOME PREZENTATE")
mr1.bold = True
mr1.font.size = Pt(8)
mp2 = mc.add_paragraph(lookup["MOTIF DE CONSULTATION / SYMPTÔMES PRÉSENTÉS"])
mp2.paragraph_format.space_after = Pt(0)
mp2.runs[0].font.size = Pt(8)

dc = motiv_t.rows[0].cells[1]
dc.paragraphs[0].clear()
dp1 = dc.paragraphs[0]
dp1.paragraph_format.space_after = Pt(2)
dr1 = dp1.add_run("DIAGNOSTIC FINAL")
dr1.bold = True
dr1.font.size = Pt(8)
diag_lines = lookup["DIAGNOSTIC FINAL"].split("\n")
for line in diag_lines:
    dp2 = dc.add_paragraph(line)
    dp2.paragraph_format.space_after = Pt(0)
    if dp2.runs:
        dp2.runs[0].font.size = Pt(8)
        if "Otite" in line or "Otită" in line:
            dp2.runs[0].bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(2)

section_header("SECȚIUNEA C - TRATAMENT, INVESTIGAȚII ȘI COST ESTIMAT")

tc = doc.add_table(rows=1, cols=5)
tc.width = Cm(18)
set_table_border(tc)
headers_c = ["Tratament / serviciu", "Categorie", "Cant.", "Cost brut", "Justificare medicală"]
col_widths = [Cm(5), Cm(4), Cm(1.5), Cm(2.5), Cm(5)]
for i, (h, w) in enumerate(zip(headers_c, col_widths)):
    tc.rows[0].cells[i].width = w
    para_in_cell(tc.rows[0].cells[i], h, bold=True, size=8)
    set_cell_bg(tc.rows[0].cells[i], "E8E8E8")

service_rows = [
    ("Traitement / service - Consultation ORL", "Consultație ORL"),
    ("Traitement / service - Injection Lincocin", "Injecție Lincocin"),
    ("Traitement / service - Injection Depo-Medrol", "Injecție Depo-Medrol"),
]

def parse_service(val):
    parts = {}
    for seg in val.split("|"):
        seg = seg.strip()
        if "Catégorie" in seg or "Categorie" in seg:
            parts["cat"] = seg.split(":", 1)[-1].strip()
        elif "Qté" in seg or "Cant." in seg:
            parts["qty"] = seg.split(":", 1)[-1].strip()
        elif "Coût brut" in seg or "Cost brut" in seg:
            parts["cost"] = seg.split(":", 1)[-1].strip()
        elif "Justification" in seg or "Justificare" in seg:
            parts["just"] = seg.split(":", 1)[-1].strip()
    return parts

for key, name in service_rows:
    val = lookup[key]
    p = parse_service(val)
    nr = tc.add_row()
    for i, w in enumerate(col_widths):
        nr.cells[i].width = w
    vals = [name, p.get("cat", ""), p.get("qty", ""), p.get("cost", ""), p.get("just", "")]
    for i, v in enumerate(vals):
        para_in_cell(nr.cells[i], v, size=8)

total_row = tc.add_row()
for i, w in enumerate(col_widths):
    total_row.cells[i].width = w
total_parts = lookup["Coût prévu / total brut"].split("|")
total_val = total_parts[0].strip()
total_note = total_parts[1].strip() if len(total_parts) > 1 else ""
merged = total_row.cells[0].merge(total_row.cells[1]).merge(total_row.cells[2])
para_in_cell(merged, "Cost anticipat / total brut", bold=True, size=8)
para_in_cell(total_row.cells[3], total_val, bold=True, size=8)
para_in_cell(total_row.cells[4], total_note, bold=True, size=8)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

td = doc.add_table(rows=1, cols=2)
td.width = Cm(18)
set_table_border(td)

decl_m = td.rows[0].cells[0]
decl_m.width = Cm(9)
decl_m.paragraphs[0].clear()
dm1 = decl_m.paragraphs[0]
dm1.paragraph_format.space_after = Pt(2)
dmr = dm1.add_run("Declarația membrului")
dmr.bold = True
dmr.font.size = Pt(9)
dm2 = decl_m.add_paragraph(lookup["Déclaration du membre"])
dm2.paragraph_format.space_after = Pt(0)
dm2.runs[0].font.size = Pt(7)
dm3 = decl_m.add_paragraph("\n\n_________________________________")
dm3.paragraph_format.space_after = Pt(0)
dm3.runs[0].font.size = Pt(8)
dm4 = decl_m.add_paragraph("Semnătură membru / dată")
dm4.paragraph_format.space_after = Pt(0)
dm4.runs[0].font.size = Pt(7)
dm4.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

decl_f = td.rows[0].cells[1]
decl_f.width = Cm(9)
decl_f.paragraphs[0].clear()
df1 = decl_f.paragraphs[0]
df1.paragraph_format.space_after = Pt(2)
dfr = df1.add_run("Declarația furnizorului medical")
dfr.bold = True
dfr.font.size = Pt(9)
df2 = decl_f.add_paragraph(lookup["Déclaration du prestataire médical"])
df2.paragraph_format.space_after = Pt(0)
df2.runs[0].font.size = Pt(7)
df3 = decl_f.add_paragraph("\n\n_________________________________")
df3.paragraph_format.space_after = Pt(0)
df3.runs[0].font.size = Pt(8)
df4 = decl_f.add_paragraph("Nume medic, semnătură și ștampilă")
df4.paragraph_format.space_after = Pt(0)
df4.runs[0].font.size = Pt(7)
df4.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(4)
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
fpr = fp.add_run("Formular claim medical - RO | pag. 2/3")
fpr.font.size = Pt(7)
fpr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)