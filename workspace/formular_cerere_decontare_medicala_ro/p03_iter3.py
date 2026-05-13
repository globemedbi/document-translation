import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p03.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_03.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

NAVY = "1F3864"
LIGHT_GREY = "F2F2F2"
WHITE = "FFFFFF"
DARK_GREY = "404040"
MID_GREY = "7F7F7F"
GREEN_APPROVED = "70AD47"

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_border(table, color="CCCCCC", sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def cell_para(cell, text, bold=False, size=9, color=DARK_GREY, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_section_header(doc, text, bg=NAVY, fg=WHITE, size=10):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(17)
    c = t.rows[0].cells[0]
    set_cell_bg(c, bg)
    p = cell_para(c, text, bold=True, size=size, color=fg)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return t

def merge_spacing(doc, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ── BAND 1: HEADER ──────────────────────────────────────────────────────────
t = doc.add_table(rows=1, cols=2)
t.width = Cm(17)
from docx.shared import Inches
t.columns[0].width = Cm(11)
t.columns[1].width = Cm(6)

left_cell = t.rows[0].cells[0]
right_cell = t.rows[0].cells[1]

set_cell_bg(left_cell, WHITE)
set_cell_bg(right_cell, LIGHT_GREY)

# Left: logo + title
inner = left_cell.add_table(rows=1, cols=2)
inner.width = Cm(11)
inner.columns[0].width = Cm(2.2)
inner.columns[1].width = Cm(8.8)

logo_cell = inner.rows[0].cells[0]
set_cell_bg(logo_cell, NAVY)
lp = logo_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lr = lp.add_run("MC")
lr.bold = True
lr.font.size = Pt(14)
lr.font.color.rgb = RGBColor.from_string(WHITE)

title_cell = inner.rows[0].cells[1]
set_cell_bg(title_cell, WHITE)
tp = title_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr2 = tp.add_run(lookup["Rapport d'approbation et détails de couverture"])
tr2.bold = True
tr2.font.size = Pt(14)
tr2.font.color.rgb = RGBColor.from_string(NAVY)

sub_p = title_cell.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
sr = sub_p.add_run("Résumé pour la compagnie d'assurance - éligibilité, documents et conditions de paiement")
sr.italic = True
sr.font.size = Pt(8)
sr.font.color.rgb = RGBColor.from_string(MID_GREY)

# Right: info box
info_fields = [
    ("Réf. approbation", lookup["Réf. approbation"]),
    ("Date d'approbation", lookup["Date d'approbation"]),
    ("Valable jusqu'au", lookup["Valable jusqu'au"]),
    ("Statut", lookup["Statut"]),
]
it = right_cell.add_table(rows=4, cols=2)
it.width = Cm(6)
for i, (k, v) in enumerate(info_fields):
    kc = it.rows[i].cells[0]
    vc = it.rows[i].cells[1]
    set_cell_bg(kc, LIGHT_GREY)
    set_cell_bg(vc, LIGHT_GREY)
    cell_para(kc, k, bold=False, size=8, color=DARK_GREY)
    cell_para(vc, v, bold=True, size=9, color=NAVY)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── BAND 2: DONNÉES DE RÉFÉRENCE ────────────────────────────────────────────
add_section_header(doc, lookup["DONNÉES DE RÉFÉRENCE"])

ref_t = doc.add_table(rows=3, cols=2)
ref_t.width = Cm(17)
set_table_border(ref_t, color="DDDDDD")

ref_pairs = [
    ("NUMÉRO INDIVIDUEL", lookup["NUMÉRO INDIVIDUEL"], "NOM DU TITULAIRE DE CARTE", lookup["NOM DU TITULAIRE DE CARTE"]),
    ("DATE DE NAISSANCE", lookup["DATE DE NAISSANCE"], "COMPAGNIE D'ASSURANCE", lookup["COMPAGNIE D'ASSURANCE"]),
    ("MÉDECIN / PRESTATAIRE", lookup["MÉDECIN / PRESTATAIRE"], "SPÉCIALITÉ", lookup["SPÉCIALITÉ"]),
]

for i, (k1, v1, k2, v2) in enumerate(ref_pairs):
    r = ref_t.rows[i]
    for ci, (k, v) in enumerate([(k1, v1), (k2, v2)]):
        c = r.cells[ci]
        set_cell_bg(c, WHITE)
        kp = c.paragraphs[0]
        kp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kr = kp.add_run(k)
        kr.font.size = Pt(7)
        kr.font.color.rgb = RGBColor.from_string(MID_GREY)
        kr.bold = False
        vp = c.add_paragraph()
        vr = vp.add_run(v)
        vr.bold = True
        vr.font.size = Pt(9)
        vr.font.color.rgb = RGBColor.from_string(DARK_GREY)

doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── BAND 3: SERVICES APPROUVÉS ──────────────────────────────────────────────
add_section_header(doc, lookup["SERVICES APPROUVÉS"])

svc_t = doc.add_table(rows=4, cols=6)
svc_t.width = Cm(17)
set_table_border(svc_t, color="CCCCCC")

headers_svc = ["Catégorie de service", "Élément approuvé", "Qté\ndemandée", "Qté\napprouvée", "Statut", "Motif / note"]
hr = svc_t.rows[0]
for i, h in enumerate(headers_svc):
    c = hr.cells[i]
    set_cell_bg(c, LIGHT_GREY)
    cell_para(c, h, bold=True, size=8, color=DARK_GREY)

svc_rows = [
    lookup["Catégorie de service - Ligne 1"],
    lookup["Catégorie de service - Ligne 2"],
    lookup["Catégorie de service - Ligne 3"],
]

for ri, row_val in enumerate(svc_rows):
    parts = [p.strip() for p in row_val.split("|")]
    r = svc_t.rows[ri + 1]
    for ci, part in enumerate(parts):
        c = r.cells[ci]
        set_cell_bg(c, WHITE)
        if ci == 4:
            cell_para(c, part, bold=True, size=8, color=GREEN_APPROVED)
        else:
            cell_para(c, part, bold=False, size=8, color=DARK_GREY)

doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── BAND 4: DOCUMENTS + DIAGNOSTIC (side by side) ───────────────────────────
main_t = doc.add_table(rows=1, cols=2)
main_t.width = Cm(17)
main_t.columns[0].width = Cm(8.5)
main_t.columns[1].width = Cm(8.5)

doc_cell = main_t.rows[0].cells[0]
diag_cell = main_t.rows[0].cells[1]

set_cell_bg(doc_cell, WHITE)
set_cell_bg(diag_cell, WHITE)

# Documents header
dp = doc_cell.paragraphs[0]
dp.clear()
dr = dp.add_run(lookup["DOCUMENTS REQUIS POUR LA SOUMISSION DE LA DEMANDE DE REMBOURSEMENT"])
dr.bold = True
dr.font.size = Pt(8)
dr.font.color.rgb = RGBColor.from_string(WHITE)
dp.paragraph_format.space_before = Pt(3)
dp.paragraph_format.space_after = Pt(3)
dp.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Redo doc cell with inner header
doc_inner = doc_cell.add_table(rows=1, cols=1)
doc_inner.width = Cm(8.5)
dh_cell = doc_inner.rows[0].cells[0]
set_cell_bg(dh_cell, NAVY)
cell_para(dh_cell, lookup["DOCUMENTS REQUIS POUR LA SOUMISSION DE LA DEMANDE DE REMBOURSEMENT"], bold=True, size=8, color=WHITE)

for line in lookup["Liste des documents requis"].split("\n"):
    lp2 = doc_cell.add_paragraph()
    lr2 = lp2.add_run(line)
    lr2.font.size = Pt(8)
    lr2.font.color.rgb = RGBColor.from_string(DARK_GREY)
    lp2.paragraph_format.space_before = Pt(1)
    lp2.paragraph_format.space_after = Pt(1)

# Clear default paragraph in doc_cell
for p in doc_cell.paragraphs:
    if p.text == "":
        p._element.getparent().remove(p._element)
        break

# Diagnostic header
diag_inner = diag_cell.add_table(rows=1, cols=1)
diag_inner.width = Cm(8.5)
dih_cell = diag_inner.rows[0].cells[0]
set_cell_bg(dih_cell, NAVY)
cell_para(dih_cell, lookup["DIAGNOSTIC ET INFORMATIONS MÉDICALES"], bold=True, size=8, color=WHITE)

# Diagnostic data table
diag_data_t = diag_cell.add_table(rows=6, cols=2)
diag_data_t.width = Cm(8.5)
set_table_border(diag_data_t, color="DDDDDD")

diag_col_header_r = diag_data_t.rows[0]
set_cell_bg(diag_col_header_r.cells[0], LIGHT_GREY)
set_cell_bg(diag_col_header_r.cells[1], LIGHT_GREY)
cell_para(diag_col_header_r.cells[0], "Champ", bold=True, size=8, color=DARK_GREY)
cell_para(diag_col_header_r.cells[1], "Valeur", bold=True, size=8, color=DARK_GREY)

diag_fields = [
    ("Diagnostic principal", lookup["Diagnostic principal"]),
    ("Date de la visite", lookup["Date de la visite"]),
    ("Tension / pouls / température", lookup["Tension / pouls / température"]),
    ("Durée de la maladie", lookup["Durée de la maladie"]),
    ("Fichiers téléchargés", lookup["Fichiers téléchargés"]),
]

for ri, (k, v) in enumerate(diag_fields):
    r = diag_data_t.rows[ri + 1]
    set_cell_bg(r.cells[0], WHITE)
    set_cell_bg(r.cells[1], WHITE)
    cell_para(r.cells[0], k, bold=False, size=8, color=DARK_GREY)
    cell_para(r.cells[1], v, bold=False, size=8, color=DARK_GREY)

doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── BAND 5: CONDITIONS D'ÉVALUATION ET DE PAIEMENT ──────────────────────────
add_section_header(doc, lookup["CONDITIONS D'ÉVALUATION ET DE PAIEMENT"])

cond_t = doc.add_table(rows=1, cols=1)
cond_t.width = Cm(17)
set_table_border(cond_t, color="DDDDDD")
cc = cond_t.rows[0].cells[0]
set_cell_bg(cc, WHITE)

for para_text in lookup["Texte des conditions d'évaluation"].split("\n"):
    pp = cc.add_paragraph()
    parts2 = para_text.split("Note TVA")
    if len(parts2) == 2:
        normal_run = pp.add_run(parts2[0])
        normal_run.font.size = Pt(8)
        normal_run.font.color.rgb = RGBColor.from_string(DARK_GREY)
        bold_run = pp.add_run("Note TVA")
        bold_run.bold = True
        bold_run.font.size = Pt(8)
        bold_run.font.color.rgb = RGBColor.from_string(DARK_GREY)
        rest_run = pp.add_run(parts2[1])
        rest_run.font.size = Pt(8)
        rest_run.font.color.rgb = RGBColor.from_string(DARK_GREY)
    else:
        r2 = pp.add_run(para_text)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor.from_string(DARK_GREY)
    pp.paragraph_format.space_before = Pt(2)
    pp.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── BAND 6: USAGE INTERNE + CONTRÔLE QUALITÉ ────────────────────────────────
bot_t = doc.add_table(rows=1, cols=2)
bot_t.width = Cm(17)
bot_t.columns[0].width = Cm(8.5)
bot_t.columns[1].width = Cm(8.5)
set_table_border(bot_t, color="CCCCCC")

left_bot = bot_t.rows[0].cells[0]
right_bot = bot_t.rows[0].cells[1]
set_cell_bg(left_bot, WHITE)
set_cell_bg(right_bot, WHITE)

# Left: internal use
lu_h = left_bot.add_table(rows=1, cols=1)
lu_h.width = Cm(8.5)
lu_hc = lu_h.rows[0].cells[0]
set_cell_bg(lu_hc, LIGHT_GREY)
cell_para(lu_hc, lookup["Usage interne - compagnie d'assurance"], bold=True, size=9, color=DARK_GREY)

dec_p = left_bot.add_paragraph()
dr3 = dec_p.add_run("Décision")
dr3.bold = True
dr3.font.size = Pt(8)
dr3.font.color.rgb = RGBColor.from_string(DARK_GREY)

chk1_p = left_bot.add_paragraph()
chk1_r = chk1_p.add_run("☑  Approuvé")
chk1_r.font.size = Pt(9)
chk1_r.font.color.rgb = RGBColor.from_string(DARK_GREY)

chk2_p = left_bot.add_paragraph()
chk2_r = chk2_p.add_run("☐  Refusé")
chk2_r.font.size = Pt(9)
chk2_r.font.color.rgb = RGBColor.from_string(DARK_GREY)

nr_p = left_bot.add_paragraph()
nr_lbl = nr_p.add_run("N° d'approbation     ")
nr_lbl.bold = True
nr_lbl.font.size = Pt(8)
nr_lbl.font.color.rgb = RGBColor.from_string(DARK_GREY)
nr_val = nr_p.add_run(lookup["N° d'approbation"])
nr_val.font.size = Pt(8)
nr_val.font.color.rgb = RGBColor.from_string(DARK_GREY)

com_p = left_bot.add_paragraph()
com_lbl = com_p.add_run("Commentaires     ")
com_lbl.bold = True
com_lbl.font.size = Pt(8)
com_lbl.font.color.rgb = RGBColor.from_string(DARK_GREY)
com_val = com_p.add_run(lookup["Commentaires"])
com_val.font.size = Pt(8)
com_val.font.color.rgb = RGBColor.from_string(DARK_GREY)

# Right: quality control
qc_h = right_bot.add_table(rows=1, cols=1)
qc_h.width = Cm(8.5)
qc_hc = qc_h.rows[0].cells[0]
set_cell_bg(qc_hc, LIGHT_GREY)
cell_para(qc_hc, lookup["Contrôle qualité de la demande"], bold=True, size=9, color=DARK_GREY)

qc_desc_p = right_bot.add_paragraph()
qc_desc_r = qc_desc_p.add_run(lookup["Contrôle qualité de la demande - description"])
qc_desc_r.font.size = Pt(8)
qc_desc_r.font.color.rgb = RGBColor.from_string(DARK_GREY)
qc_desc_p.paragraph_format.space_before = Pt(4)
qc_desc_p.paragraph_format.space_after = Pt(12)

sig_line_p = right_bot.add_paragraph()
sig_line_r = sig_line_p.add_run("_" * 45)
sig_line_r.font.size = Pt(9)
sig_line_r.font.color.rgb = RGBColor.from_string(DARK_GREY)

sig_label_p = right_bot.add_paragraph()
sig_label_r = sig_label_p.add_run(lookup["Nom de l'évaluateur / signature / date"])
sig_label_r.font.size = Pt(8)
sig_label_r.font.color.rgb = RGBColor.from_string(MID_GREY)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── FOOTER ──────────────────────────────────────────────────────────────────
ft = doc.add_table(rows=1, cols=2)
ft.width = Cm(17)
set_table_border(ft, color="CCCCCC", sz=2)
fl = ft.rows[0].cells[0]
fr = ft.rows[0].cells[1]
set_cell_bg(fl, LIGHT_GREY)
set_cell_bg(fr, LIGHT_GREY)
cell_para(fl, "Formulaire de demande médicale - RO | pag. 3/3", bold=False, size=7, color=MID_GREY)
cell_para(fr, "Document modèle avec données anonymisées", bold=False, size=7, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)