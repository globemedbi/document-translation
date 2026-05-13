import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p03.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_03.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

NAVY = RGBColor(0x1F, 0x2D, 0x5A)
GREY_BG = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xD9, 0xD9, 0xD9)
GREEN = RGBColor(0x70, 0xAD, 0x47)

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    hex_color = "{:02X}{:02X}{:02X}".format(rgb.red, rgb.green, rgb.blue)
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_border(table, color="AAAAAA"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def para(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_para(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def set_spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.5)

for p in doc.paragraphs:
    p._element.getparent().remove(p._element)

# BAND 1 - HEADER
t = doc.add_table(rows=1, cols=2)
t.width = Cm(18)
remove_borders(t)
t.rows[0].cells[0].width = Cm(11)
t.rows[0].cells[1].width = Cm(7)

left_cell = t.rows[0].cells[0]
right_cell = t.rows[0].cells[1]

# Left: MC logo + title
logo_title_t = left_cell.add_table(rows=1, cols=2)
remove_borders(logo_title_t)
logo_cell = logo_title_t.rows[0].cells[0]
logo_cell.width = Cm(2.2)
title_cell = logo_title_t.rows[0].cells[1]
title_cell.width = Cm(8.5)

set_cell_bg(logo_cell, NAVY)
p = logo_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MC")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = WHITE
set_spacing(p, 6, 2)

p2 = para(title_cell, lookup["Rapport d'approbation et détails de couverture"], bold=True, size=14, color=NAVY)
set_spacing(p2, 2, 2)

p3 = title_cell.add_paragraph()
run3 = p3.add_run("Résumé pour la compagnie d'assurance - éligibilité, documents et conditions de paiement")
run3.font.size = Pt(8)
run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run3.italic = True
set_spacing(p3, 0, 0)

# Right: info box
set_cell_bg(right_cell, GREY_BG)
info_keys = ["Réf. approbation", "Date d'approbation", "Valable jusqu'au", "Statut"]
for k in info_keys:
    v = lookup[k]
    p_info = right_cell.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(p_info, 1, 1)
    r1 = p_info.add_run(k + "  ")
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    r2 = p_info.add_run(v)
    r2.bold = True
    r2.font.size = Pt(9)
    if k == "Statut":
        r2.font.color.rgb = NAVY
    else:
        r2.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

sp = doc.add_paragraph()
set_spacing(sp, 4, 4)

# BAND 2 - DONNÉES DE RÉFÉRENCE
t2_header = doc.add_table(rows=1, cols=1)
t2_header.width = Cm(18)
remove_borders(t2_header)
hc = t2_header.rows[0].cells[0]
set_cell_bg(hc, NAVY)
ph = para(hc, lookup["DONNÉES DE RÉFÉRENCE"], bold=True, size=10, color=WHITE)
set_spacing(ph, 3, 3)

t2 = doc.add_table(rows=3, cols=2)
t2.width = Cm(18)
set_table_border(t2, "CCCCCC")

ref_left = [
    ("NUMÉRO INDIVIDUEL", lookup["NUMÉRO INDIVIDUEL"]),
    ("DATE DE NAISSANCE", lookup["DATE DE NAISSANCE"]),
    ("MÉDECIN / PRESTATAIRE", lookup["MÉDECIN / PRESTATAIRE"]),
]
ref_right = [
    ("NOM DU TITULAIRE DE CARTE", lookup["NOM DU TITULAIRE DE CARTE"]),
    ("COMPAGNIE D'ASSURANCE", lookup["COMPAGNIE D'ASSURANCE"]),
    ("SPÉCIALITÉ", lookup["SPÉCIALITÉ"]),
]

for i, ((lk, lv), (rk, rv)) in enumerate(zip(ref_left, ref_right)):
    lc = t2.rows[i].cells[0]
    rc = t2.rows[i].cells[1]
    lc.width = Cm(9)
    rc.width = Cm(9)
    set_cell_bg(lc, GREY_BG)
    set_cell_bg(rc, GREY_BG)
    p_l = para(lc, lk, bold=False, size=8, color=RGBColor(0x60, 0x60, 0x60))
    set_spacing(p_l, 2, 0)
    add_para(lc, lv, bold=True, size=9, color=NAVY)
    p_r = para(rc, rk, bold=False, size=8, color=RGBColor(0x60, 0x60, 0x60))
    set_spacing(p_r, 2, 0)
    add_para(rc, rv, bold=True, size=9, color=NAVY)

sp2 = doc.add_paragraph()
set_spacing(sp2, 4, 4)

# BAND 3 - SERVICES APPROUVÉS
t3_header = doc.add_table(rows=1, cols=1)
t3_header.width = Cm(18)
remove_borders(t3_header)
hc3 = t3_header.rows[0].cells[0]
set_cell_bg(hc3, NAVY)
ph3 = para(hc3, lookup["SERVICES APPROUVÉS"], bold=True, size=10, color=WHITE)
set_spacing(ph3, 3, 3)

t3 = doc.add_table(rows=4, cols=6)
t3.width = Cm(18)
set_table_border(t3, "CCCCCC")

col_widths = [3.5, 5.5, 1.5, 1.5, 2.0, 3.5]
headers_row = ["Catégorie de service", "Élément approuvé", "Qté\ndem.", "Qté\nappr.", "Statut", "Motif / note"]
for j, (hw, ht) in enumerate(zip(col_widths, headers_row)):
    c = t3.rows[0].cells[j]
    c.width = Cm(hw)
    set_cell_bg(c, LIGHT_GREY)
    p_h = para(c, ht, bold=True, size=8)
    set_spacing(p_h, 2, 2)

svc_rows = [
    lookup["Catégorie de service - Ligne 1"],
    lookup["Catégorie de service - Ligne 2"],
    lookup["Catégorie de service - Ligne 3"],
]

for i, row_data in enumerate(svc_rows):
    parts = [p.strip() for p in row_data.split("|")]
    for j, val in enumerate(parts[:6]):
        c = t3.rows[i+1].cells[j]
        c.width = Cm(col_widths[j])
        if j == 4:
            set_cell_bg(c, RGBColor(0xE2, 0xEF, 0xDA))
            p_v = para(c, val, bold=True, size=8, color=RGBColor(0x37, 0x6A, 0x23), align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            p_v = para(c, val, size=9)
        set_spacing(p_v, 2, 2)

sp3 = doc.add_paragraph()
set_spacing(sp3, 4, 4)

# BAND 4 - DOCUMENTS + DIAGNOSTIC (side by side)
t4 = doc.add_table(rows=1, cols=2)
t4.width = Cm(18)
remove_borders(t4)
left_col = t4.rows[0].cells[0]
right_col = t4.rows[0].cells[1]
left_col.width = Cm(9)
right_col.width = Cm(9)

# Left: Documents
doc_inner = left_col.add_table(rows=2, cols=1)
remove_borders(doc_inner)
doc_inner.width = Cm(8.8)

dh = doc_inner.rows[0].cells[0]
set_cell_bg(dh, NAVY)
p_dh = para(dh, lookup["DOCUMENTS REQUIS POUR LA SOUMISSION DE LA DEMANDE DE REMBOURSEMENT"], bold=True, size=8, color=WHITE)
set_spacing(p_dh, 3, 3)

db = doc_inner.rows[1].cells[0]
set_cell_bg(db, GREY_BG)
doc_list = lookup["Liste des documents requis"].split("\n")
for item in doc_list:
    p_item = db.add_paragraph()
    run_item = p_item.add_run(item)
    run_item.font.size = Pt(8)
    set_spacing(p_item, 1, 1)

# Right: Diagnostic
diag_inner = right_col.add_table(rows=2, cols=1)
remove_borders(diag_inner)
diag_inner.width = Cm(8.8)

dgh = diag_inner.rows[0].cells[0]
set_cell_bg(dgh, NAVY)
p_dgh = para(dgh, lookup["DIAGNOSTIC ET INFORMATIONS MÉDICALES"], bold=True, size=8, color=WHITE)
set_spacing(p_dgh, 3, 3)

dgb = diag_inner.rows[1].cells[0]
set_cell_bg(dgb, GREY_BG)

diag_table = dgb.add_table(rows=5, cols=2)
remove_borders(diag_table)
diag_table.width = Cm(8.6)

diag_fields = [
    ("Diagnostic principal", lookup["Diagnostic principal"]),
    ("Date de la visite", lookup["Date de la visite"]),
    ("Tension / pouls / température", lookup["Tension / pouls / température"]),
    ("Durée de la maladie", lookup["Durée de la maladie"]),
    ("Fichiers téléchargés", lookup["Fichiers téléchargés"]),
]

for i, (fk, fv) in enumerate(diag_fields):
    fc = diag_table.rows[i].cells[0]
    vc = diag_table.rows[i].cells[1]
    fc.width = Cm(4.0)
    vc.width = Cm(4.6)
    set_cell_border(fc, bottom={"val": "single", "sz": "2", "color": "CCCCCC"})
    set_cell_border(vc, bottom={"val": "single", "sz": "2", "color": "CCCCCC"})
    pf = para(fc, fk, bold=False, size=8, color=RGBColor(0x60, 0x60, 0x60))
    set_spacing(pf, 2, 2)
    pv = para(vc, fv, bold=False, size=8)
    set_spacing(pv, 2, 2)

sp4 = doc.add_paragraph()
set_spacing(sp4, 4, 4)

# BAND 5 - CONDITIONS D'ÉVALUATION ET DE PAIEMENT
t5_header = doc.add_table(rows=1, cols=1)
t5_header.width = Cm(18)
remove_borders(t5_header)
hc5 = t5_header.rows[0].cells[0]
set_cell_bg(hc5, NAVY)
ph5 = para(hc5, lookup["CONDITIONS D'ÉVALUATION ET DE PAIEMENT"], bold=True, size=10, color=WHITE)
set_spacing(ph5, 3, 3)

t5b = doc.add_table(rows=1, cols=1)
t5b.width = Cm(18)
set_table_border(t5b, "CCCCCC")
cond_cell = t5b.rows[0].cells[0]
set_cell_bg(cond_cell, GREY_BG)

cond_text = lookup["Texte des conditions d'évaluation"].split("\n")
for i, line in enumerate(cond_text):
    if i == 0:
        p_c = para(cond_cell, line, size=8)
    else:
        p_c = add_para(cond_cell, line, size=8)
    if "Note TVA" in line or "TVA" in line:
        p_c.clear()
        r_bold = p_c.add_run("Note TVA : ")
        r_bold.bold = True
        r_bold.font.size = Pt(8)
        rest = line.replace("Note TVA : ", "")
        r_rest = p_c.add_run(rest)
        r_rest.font.size = Pt(8)
    set_spacing(p_c, 2, 2)

sp5 = doc.add_paragraph()
set_spacing(sp5, 4, 4)

# BAND 6 - Internal use + Quality control (side by side)
t6 = doc.add_table(rows=1, cols=2)
t6.width = Cm(18)
set_table_border(t6, "CCCCCC")
int_cell = t6.rows[0].cells[0]
qc_cell = t6.rows[0].cells[1]
int_cell.width = Cm(9)
qc_cell.width = Cm(9)

set_cell_bg(int_cell, GREY_BG)
set_cell_bg(qc_cell, GREY_BG)

# Internal use
p_int_h = para(int_cell, lookup["Usage interne - compagnie d'assurance"], bold=True, size=9, color=NAVY)
set_spacing(p_int_h, 3, 4)

p_dec = add_para(int_cell, "Décision", bold=True, size=8)
set_spacing(p_dec, 2, 2)

p_app = add_para(int_cell, "☑ " + "Approuvé", bold=False, size=9)
set_spacing(p_app, 2, 2)

p_ref = add_para(int_cell, "☐ " + "Refusé", bold=False, size=9)
set_spacing(p_ref, 2, 4)

p_nr_l = add_para(int_cell, "N° d'approbation", bold=True, size=8)
set_spacing(p_nr_l, 2, 0)
p_nr_v = add_para(int_cell, lookup["N° d'approbation"], bold=False, size=9)
set_spacing(p_nr_v, 0, 2)

p_com_l = add_para(int_cell, "Commentaires", bold=True, size=8)
set_spacing(p_com_l, 2, 0)
p_com_v = add_para(int_cell, lookup["Commentaires"], bold=False, size=9)
set_spacing(p_com_v, 0, 4)

# Quality control
p_qc_h = para(qc_cell, lookup["Contrôle qualité de la demande"], bold=True, size=9, color=NAVY)
set_spacing(p_qc_h, 3, 4)

p_qc_desc = add_para(qc_cell, lookup["Contrôle qualité de la demande - description"], size=8)
set_spacing(p_qc_desc, 2, 10)

p_sig_line = add_para(qc_cell, "_" * 45, size=9)
set_spacing(p_sig_line, 8, 2)

p_sig_label = add_para(qc_cell, "Nom de l'évaluateur / signature / date", size=8, color=RGBColor(0x60, 0x60, 0x60))
set_spacing(p_sig_label, 0, 4)

sp6 = doc.add_paragraph()
set_spacing(sp6, 6, 6)

# FOOTER
t_footer = doc.add_table(rows=1, cols=2)
t_footer.width = Cm(18)
remove_borders(t_footer)
fc_left = t_footer.rows[0].cells[0]
fc_right = t_footer.rows[0].cells[1]
fc_left.width = Cm(9)
fc_right.width = Cm(9)

pfl = para(fc_left, "Formulaire demande médicale - FR | pag. 3/3", size=7, color=RGBColor(0x80, 0x80, 0x80))
set_spacing(pfl, 0, 0)
pfr = para(fc_right, "Document modèle avec données anonymisées", size=7, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.RIGHT)
set_spacing(pfr, 0, 0)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)