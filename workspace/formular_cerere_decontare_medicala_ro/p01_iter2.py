import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/app/workspace/formular_cerere_decontare_medicala_ro/data_p01.json"
OUTPUT_PATH = "/app/workspace/formular_cerere_decontare_medicala_ro/page_01.docx"
DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_font(run, size, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def no_space(par):
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)

def add_par(text, size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, container=None):
    if container is None:
        p = doc.add_paragraph()
    else:
        p = container.add_paragraph()
    no_space(p)
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p

# BAND 1 - Header bar table (logo + title | info box)
t1 = doc.add_table(rows=1, cols=2)
t1.width = Cm(18)
from docx.shared import Inches
t1.columns[0].width = Cm(12)
t1.columns[1].width = Cm(6)

# Left cell: logo + title
lc = t1.rows[0].cells[0]
set_cell_bg(lc, "FFFFFF")
# MC logo simulation with text
p_logo = lc.paragraphs[0]
no_space(p_logo)
r_mc = p_logo.add_run("MC  ")
set_font(r_mc, 14, True, (255, 255, 255))
p_logo.paragraph_format.space_before = Pt(2)

# Use a nested table for logo area
lc.paragraphs[0].clear()
logo_t = lc.add_table(rows=1, cols=2)
logo_t.columns[0].width = Cm(1.5)
logo_t.columns[1].width = Cm(10)
mc_cell = logo_t.rows[0].cells[0]
set_cell_bg(mc_cell, "2E6E6E")
mc_p = mc_cell.paragraphs[0]
no_space(mc_p)
mc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mc_r = mc_p.add_run("MC")
set_font(mc_r, 14, True, (255, 255, 255))

title_cell = logo_t.rows[0].cells[1]
set_cell_bg(title_cell, "FFFFFF")
tp = title_cell.paragraphs[0]
no_space(tp)
tr = tp.add_run(lookup["Reçu médical et récapitulatif des coûts"])
set_font(tr, 14, True, (30, 30, 30))
sp = title_cell.add_paragraph()
no_space(sp)
sr = sp.add_run(lookup["Dossier de remboursement - services ORL, soins primaires / ambulatoire"])
set_font(sr, 8, False, (100, 100, 100))

# Right cell: info box
rc = t1.rows[0].cells[1]
set_cell_bg(rc, "F5F5F5")
info_labels = ["N° de dossier", "Date", "Devise", "Statut"]
info_t = rc.add_table(rows=4, cols=2)
info_t.width = Cm(5.8)
for i, lbl in enumerate(info_labels):
    lc2 = info_t.rows[i].cells[0]
    vc2 = info_t.rows[i].cells[1]
    set_cell_bg(lc2, "F5F5F5")
    set_cell_bg(vc2, "F5F5F5")
    lp = lc2.paragraphs[0]
    no_space(lp)
    lr2 = lp.add_run(lbl)
    set_font(lr2, 8, True, (60, 60, 60))
    vp = vc2.paragraphs[0]
    no_space(vp)
    vr = vp.add_run(lookup[lbl])
    set_font(vr, 8, True, (30, 30, 30))
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# BAND 2 - Document purpose banner
banner_t = doc.add_table(rows=1, cols=1)
banner_t.width = Cm(18)
bc = banner_t.rows[0].cells[0]
set_cell_bg(bc, "EEF4F4")
bp = bc.paragraphs[0]
no_space(bp)
br1 = bp.add_run("Objet du document: ")
set_font(br1, 9, True, (30, 30, 30))
br2 = bp.add_run(lookup["Objet du document"])
set_font(br2, 9, False, (50, 50, 50))
bp.paragraph_format.space_before = Pt(4)
bp.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# BAND 3 - Provider / Insurance two-column panel
pi_t = doc.add_table(rows=1, cols=2)
pi_t.width = Cm(18)
pi_t.columns[0].width = Cm(9)
pi_t.columns[1].width = Cm(9)

# Left: medical provider
prov_cell = pi_t.rows[0].cells[0]
set_cell_bg(prov_cell, "FAFAFA")
set_cell_borders(prov_cell, top="single", bottom="single", left="single", right="single")
pp = prov_cell.paragraphs[0]
no_space(pp)
pr = pp.add_run(lookup["Prestataire médical"])
set_font(pr, 10, True, (30, 30, 30))

prov_fields = [
    ("Clinique", lookup["Clinique"]),
    ("Spécialité", lookup["Spécialité"]),
    ("Adresse", lookup["Adresse"]),
    ("Téléphone", lookup["Téléphone"]),
]
for lbl, val in prov_fields:
    row_t = prov_cell.add_table(rows=1, cols=2)
    row_t.columns[0].width = Cm(3)
    row_t.columns[1].width = Cm(5.5)
    lc3 = row_t.rows[0].cells[0]
    vc3 = row_t.rows[0].cells[1]
    set_cell_bg(lc3, "FAFAFA")
    set_cell_bg(vc3, "FAFAFA")
    lp3 = lc3.paragraphs[0]
    no_space(lp3)
    lr3 = lp3.add_run(lbl)
    set_font(lr3, 9, True, (70, 100, 120))
    vp3 = vc3.paragraphs[0]
    no_space(vp3)
    vr3 = vp3.add_run(val)
    set_font(vr3, 9, False, (30, 30, 30))

# Right: insurance
ins_cell = pi_t.rows[0].cells[1]
set_cell_bg(ins_cell, "FAFAFA")
set_cell_borders(ins_cell, top="single", bottom="single", left="single", right="single")
ip = ins_cell.paragraphs[0]
no_space(ip)
ir = ip.add_run(lookup["Assurance / gestionnaire de sinistre"])
set_font(ir, 10, True, (30, 30, 30))

ins_fields = [
    ("Compagnie / TPA", lookup["Compagnie / TPA"]),
    ("Titulaire de la police", lookup["Titulaire de la police"]),
    ("Membre", lookup["Membre"]),
    ("Type de demande", lookup["Type de demande"]),
]
for lbl, val in ins_fields:
    row_t2 = ins_cell.add_table(rows=1, cols=2)
    row_t2.columns[0].width = Cm(3.5)
    row_t2.columns[1].width = Cm(5)
    lc4 = row_t2.rows[0].cells[0]
    vc4 = row_t2.rows[0].cells[1]
    set_cell_bg(lc4, "FAFAFA")
    set_cell_bg(vc4, "FAFAFA")
    lp4 = lc4.paragraphs[0]
    no_space(lp4)
    lr4 = lp4.add_run(lbl)
    set_font(lr4, 9, True, (70, 100, 120))
    vp4 = vc4.paragraphs[0]
    no_space(vp4)
    vr4 = vp4.add_run(val)
    set_font(vr4, 9, False, (30, 30, 30))

doc.add_paragraph()

# BAND 4 - Financial summary section header
fs_t = doc.add_table(rows=1, cols=1)
fs_t.width = Cm(18)
fsc = fs_t.rows[0].cells[0]
set_cell_bg(fsc, "2E6E6E")
fsp = fsc.paragraphs[0]
no_space(fsp)
fsp.paragraph_format.space_before = Pt(4)
fsp.paragraph_format.space_after = Pt(4)
fsr = fsp.add_run(lookup["SYNTHÈSE FINANCIÈRE"])
set_font(fsr, 11, True, (255, 255, 255))

# Financial summary 3-column totals
sum_t = doc.add_table(rows=1, cols=3)
sum_t.width = Cm(18)
for i in range(3):
    sum_t.columns[i].width = Cm(6)

total_items = [
    ("TOTAL BRUT", lookup["TOTAL BRUT"], "Servicii facturate de clinică"),
    ("RÉDUCTIONS CONTRACTUELLES", lookup["RÉDUCTIONS CONTRACTUELLES"], "Appliquées sur consultation et injections"),
    ("MONTANT NET DEMANDÉ", lookup["MONTANT NET DEMANDÉ"], "Après réduction et déduction/quote-part"),
]
for i, (lbl, val, sub) in enumerate(total_items):
    c = sum_t.rows[0].cells[i]
    set_cell_bg(c, "F0F0F0")
    set_cell_borders(c, top="single", bottom="single", left="single", right="single")
    p1 = c.paragraphs[0]
    no_space(p1)
    r1 = p1.add_run(lbl)
    set_font(r1, 8, True, (80, 80, 80))
    p2 = c.add_paragraph()
    no_space(p2)
    r2 = p2.add_run(val)
    set_font(r2, 14, True, (30, 30, 30))
    p3 = c.add_paragraph()
    no_space(p3)
    r3 = p3.add_run(sub)
    set_font(r3, 7, False, (100, 100, 100))

# Services table
svc_data = lookup["Tableau des services"]
svc_rows = [row.strip() for row in svc_data.strip().split("\n")]

headers = ["Serviciu / tratament", "Cant.", "Tarif brut", "Reducere", "Valoare eligibilă", "Observații"]
col_widths = [Cm(5.5), Cm(1.5), Cm(2.5), Cm(2.5), Cm(3), Cm(3)]

svc_t = doc.add_table(rows=1 + len(svc_rows), cols=6)
svc_t.width = Cm(18)
for i, w in enumerate(col_widths):
    for r in svc_t.rows:
        r.cells[i].width = w

# Header row
hr = svc_t.rows[0]
for i, h in enumerate(headers):
    c = hr.cells[i]
    set_cell_bg(c, "E8E8E8")
    p = c.paragraphs[0]
    no_space(p)
    run = p.add_run(h)
    set_font(run, 9, True, (30, 30, 30))

# Data rows
for ri, row_str in enumerate(svc_rows):
    parts = [p.strip() for p in row_str.split("|")]
    while len(parts) < 6:
        parts.append("")
    tr2 = svc_t.rows[ri + 1]
    is_total = parts[0].strip().lower().startswith("total")
    bg = "F5F5F5" if is_total else "FFFFFF"
    for ci, val in enumerate(parts[:6]):
        c = tr2.cells[ci]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        no_space(p)
        run = p.add_run(val)
        set_font(run, 9, is_total, (30, 30, 30))

# Deduction row
ded_t = doc.add_table(rows=1, cols=3)
ded_t.width = Cm(18)
ded_t.columns[0].width = Cm(9)
ded_t.columns[1].width = Cm(5.5)
ded_t.columns[2].width = Cm(3.5)

ded_lc = ded_t.rows[0].cells[0]
set_cell_bg(ded_lc, "FAFAFA")
ded_label_cell = ded_t.rows[0].cells[1]
set_cell_bg(ded_label_cell, "FAFAFA")
dp = ded_label_cell.paragraphs[0]
no_space(dp)
dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
dr = dp.add_run("Déduction / quote-part du membre")
set_font(dr, 9, True, (30, 30, 30))

ded_val_cell = ded_t.rows[0].cells[2]
set_cell_bg(ded_val_cell, "FAFAFA")
dvp = ded_val_cell.paragraphs[0]
no_space(dvp)
dvp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
dvr = dvp.add_run(lookup["Déduction / quote-part du membre"])
set_font(dvr, 9, True, (30, 30, 30))

# Net amount row
net_t = doc.add_table(rows=1, cols=3)
net_t.width = Cm(18)
net_t.columns[0].width = Cm(9)
net_t.columns[1].width = Cm(5.5)
net_t.columns[2].width = Cm(3.5)

net_lc = net_t.rows[0].cells[0]
set_cell_bg(net_lc, "FAFAFA")
net_label_cell = net_t.rows[0].cells[1]
set_cell_bg(net_label_cell, "FAFAFA")
np2 = net_label_cell.paragraphs[0]
no_space(np2)
np2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
nr2 = np2.add_run("Montant net à rembourser / à percevoir")
set_font(nr2, 9, True, (30, 30, 30))

net_val_cell = net_t.rows[0].cells[2]
set_cell_bg(net_val_cell, "E8F4E8")
nvp = net_val_cell.paragraphs[0]
no_space(nvp)
nvp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
nvr = nvp.add_run(lookup["Montant net à rembourser / à percevoir"])
set_font(nvr, 10, True, (30, 100, 30))

doc.add_paragraph()

# BAND 5 - Signature block (two columns)
sig_t = doc.add_table(rows=1, cols=2)
sig_t.width = Cm(18)
sig_t.columns[0].width = Cm(9)
sig_t.columns[1].width = Cm(9)

# Left signature: provider
lsig = sig_t.rows[0].cells[0]
set_cell_bg(lsig, "FAFAFA")
set_cell_borders(lsig, top="single", bottom="single", left="single", right="single")
lsp = lsig.paragraphs[0]
no_space(lsp)
lsr = lsp.add_run(lookup["Confirmation du prestataire médical"])
set_font(lsr, 9, True, (30, 30, 30))
lsp2 = lsig.add_paragraph()
no_space(lsp2)
lsr2 = lsp2.add_run(lookup["Confirmation du prestataire médical - texte"])
set_font(lsr2, 8, False, (60, 60, 60))
for _ in range(2):
    lsig.add_paragraph()
lsp3 = lsig.add_paragraph()
no_space(lsp3)
lsr3 = lsp3.add_run(lookup["Nom, signature et cachet"] if lookup["Nom, signature et cachet"] else "Nom, signature et cachet ________________________")
set_font(lsr3, 8, False, (100, 100, 100))

# Right signature: member/patient
rsig = sig_t.rows[0].cells[1]
set_cell_bg(rsig, "FAFAFA")
set_cell_borders(rsig, top="single", bottom="single", left="single", right="single")
rsp = rsig.paragraphs[0]
no_space(rsp)
rsr = rsp.add_run(lookup["Confirmation du membre / patient"])
set_font(rsr, 9, True, (30, 30, 30))
rsp2 = rsig.add_paragraph()
no_space(rsp2)
rsr2 = rsp2.add_run(lookup["Confirmation du membre / patient - texte"])
set_font(rsr2, 8, False, (60, 60, 60))
for _ in range(2):
    rsig.add_paragraph()
rsp3 = rsig.add_paragraph()
no_space(rsp3)
rsr3 = rsp3.add_run(lookup["Nom et signature du membre"] if lookup["Nom et signature du membre"] else "Nom et signature du membre ________________________")
set_font(rsr3, 8, False, (100, 100, 100))

# Footer
doc.add_paragraph()
fp = doc.add_paragraph()
no_space(fp)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Formular claim medical - RO | pag. 1/3                                        Document model cu date anonimizate")
set_font(fr, 7, False, (150, 150, 150))

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)