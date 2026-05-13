import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_01.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p01.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key, default=""):
    return lookup.get(key, default)

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement("w:" + side)
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_borders(table, color="CCCCCC", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_run_with_font(para, text, bold=False, size_pt=9, color_hex=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run

def set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)

def set_paragraph_space(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def merge_row_cells(table, row_idx, start_col, end_col):
    row = table.rows[row_idx]
    merged = row.cells[start_col].merge(row.cells[end_col])
    return merged

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

for para in doc.paragraphs:
    para.clear()

def clear_default_paragraph():
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

clear_default_paragraph()

TEAL = "1A7A7A"
LIGHT_TEAL_BG = "E8F4F4"
LIGHT_GRAY_BG = "F5F5F5"
DARK_GRAY_TEXT = "2D2D2D"
MID_GRAY = "666666"
BORDER_GRAY = "CCCCCC"
SECTION_HEADER_BG = "2B7A7A"
WHITE = "FFFFFF"
BLUE_GRAY_BG = "EEF3F7"

# ─── BAND 1: HEADER BAR ────────────────────────────────────────────────────────
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
header_table.allow_autofit = False
set_no_table_borders(header_table)

total_width = Cm(18)
left_col_w = Cm(12)
right_col_w = Cm(6)

header_table.columns[0].width = left_col_w
header_table.columns[1].width = right_col_w

row = header_table.rows[0]
row.cells[0].width = left_col_w
row.cells[1].width = right_col_w

# Left cell: logo + title
left_cell = row.cells[0]
remove_cell_borders(left_cell)
set_cell_vertical_align(left_cell, "center")

# Inner table for logo + text side by side
inner_tbl = left_cell.add_table(rows=1, cols=2)
inner_tbl.autofit = False
set_no_table_borders(inner_tbl)
inner_tbl.columns[0].width = Cm(1.6)
inner_tbl.columns[1].width = Cm(10.2)

logo_cell = inner_tbl.rows[0].cells[0]
logo_cell.width = Cm(1.6)
remove_cell_borders(logo_cell)
set_cell_background(logo_cell, TEAL)
set_cell_vertical_align(logo_cell, "center")
logo_para = logo_cell.paragraphs[0]
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_space(logo_para, 0, 0)
add_run_with_font(logo_para, "MC", bold=True, size_pt=14, color_hex=WHITE)

title_cell = inner_tbl.rows[0].cells[1]
title_cell.width = Cm(10.2)
remove_cell_borders(title_cell)
set_cell_vertical_align(title_cell, "center")

title_para = title_cell.paragraphs[0]
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(title_para, 0, 0)
add_run_with_font(title_para, get_value("Medical Receipt and Cost Summary"), bold=True, size_pt=14, color_hex=DARK_GRAY_TEXT)

sub_para = title_cell.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(sub_para, 1, 0)
add_run_with_font(sub_para, get_value("Reimbursement File - ENT Services, Primary Care / Outpatient"), bold=False, size_pt=8, color_hex=MID_GRAY, italic=True)

# Right cell: info box
right_cell = row.cells[1]
remove_cell_borders(right_cell)
set_cell_vertical_align(right_cell, "center")

info_table = right_cell.add_table(rows=4, cols=2)
info_table.autofit = False
set_table_borders(info_table, color=BORDER_GRAY, sz=4)
info_table.columns[0].width = Cm(2.2)
info_table.columns[1].width = Cm(3.6)

info_data = [
    (get_value("File No.", "File No."), get_value("File No.")),
    (get_value("Date", "Date"), get_value("Date")),
    (get_value("Currency", "Currency"), get_value("Currency")),
    (get_value("Status", "Status"), get_value("Status")),
]

info_labels = ["File No.", "Date", "Currency", "Status"]
info_values = [get_value("File No."), get_value("Date"), get_value("Currency"), get_value("Status")]

for i, (label, value) in enumerate(zip(info_labels, info_values)):
    lc = info_table.rows[i].cells[0]
    vc = info_table.rows[i].cells[1]
    lc.width = Cm(2.2)
    vc.width = Cm(3.6)
    set_cell_vertical_align(lc, "center")
    set_cell_vertical_align(vc, "center")
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_space(lp, 1, 1)
    add_run_with_font(lp, label, bold=False, size_pt=8, color_hex=MID_GRAY)
    vp = vc.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_space(vp, 1, 1)
    add_run_with_font(vp, value, bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

# spacer
spacer1 = doc.add_paragraph()
set_paragraph_space(spacer1, 4, 2)

# ─── BAND 2: PURPOSE BANNER ────────────────────────────────────────────────────
purpose_table = doc.add_table(rows=1, cols=1)
purpose_table.autofit = False
set_table_borders(purpose_table, color=BORDER_GRAY, sz=4)
purpose_table.columns[0].width = Cm(18)

pc = purpose_table.rows[0].cells[0]
set_cell_background(pc, LIGHT_GRAY_BG)
remove_cell_borders(pc)
set_cell_borders(pc,
    top={"val": "single", "sz": 4, "color": BORDER_GRAY},
    bottom={"val": "single", "sz": 4, "color": BORDER_GRAY},
    left={"val": "single", "sz": 4, "color": BORDER_GRAY},
    right={"val": "single", "sz": 4, "color": BORDER_GRAY})

purpose_para = pc.paragraphs[0]
purpose_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(purpose_para, 3, 3)

purpose_text = get_value(
    "Document purpose",
    "Document purpose: professional form in Romanian for a medical insurance file, built on the structure of a clinic/hospital claim: receipt, request form, medical details, costs, approved services and list of supporting documents. Identification data is anonymised."
)

if not purpose_text or purpose_text == "":
    purpose_text = "Document purpose: professional form in Romanian for a medical insurance file, built on the structure of a clinic/hospital claim: receipt, request form, medical details, costs, approved services and list of supporting documents. Identification data is anonymised."

add_run_with_font(purpose_para, "Document purpose: ", bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)
add_run_with_font(purpose_para, purpose_text.replace("Document purpose: ", "").replace("Document purpose:", "").strip(), bold=False, size_pt=9, color_hex=DARK_GRAY_TEXT)

spacer2 = doc.add_paragraph()
set_paragraph_space(spacer2, 4, 2)

# ─── BAND 3: PROVIDER / INSURANCE SPLIT ────────────────────────────────────────
prov_table = doc.add_table(rows=1, cols=2)
prov_table.autofit = False
set_table_borders(prov_table, color=BORDER_GRAY, sz=4)

left_w = Cm(8.8)
right_w = Cm(9.2)
prov_table.columns[0].width = left_w
prov_table.columns[1].width = right_w

prov_left = prov_table.rows[0].cells[0]
prov_right = prov_table.rows[0].cells[1]
prov_left.width = left_w
prov_right.width = right_w
set_cell_vertical_align(prov_left, "top")
set_cell_vertical_align(prov_right, "top")

# Left: Medical Provider
lp0 = prov_left.paragraphs[0]
lp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(lp0, 2, 2)
add_run_with_font(lp0, get_value("Medical Provider"), bold=True, size_pt=10, color_hex=DARK_GRAY_TEXT)

prov_fields = [
    ("Clinic", get_value("Clinic")),
    ("Specialty", get_value("Specialty")),
    ("Address", get_value("Address")),
    ("Phone", get_value("Phone")),
]

for label, value in prov_fields:
    row_para = prov_left.add_paragraph()
    row_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_space(row_para, 1, 1)
    add_run_with_font(row_para, label + "  ", bold=True, size_pt=9, color_hex=TEAL)
    add_run_with_font(row_para, value, bold=False, size_pt=9, color_hex=DARK_GRAY_TEXT)

# Right: Insurance / Claim Administrator
rp0 = prov_right.paragraphs[0]
rp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(rp0, 2, 2)
add_run_with_font(rp0, get_value("Insurance / Claim Administrator"), bold=True, size_pt=10, color_hex=DARK_GRAY_TEXT)

ins_fields = [
    ("Company / TPA", get_value("Company / TPA")),
    ("Policy Holder", get_value("Policy Holder")),
    ("Member", get_value("Member")),
    ("Claim Type", get_value("Claim Type")),
]

for label, value in ins_fields:
    row_para = prov_right.add_paragraph()
    row_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_space(row_para, 1, 1)
    add_run_with_font(row_para, label + "  ", bold=True, size_pt=9, color_hex=TEAL)
    add_run_with_font(row_para, value, bold=False, size_pt=9, color_hex=DARK_GRAY_TEXT)

spacer3 = doc.add_paragraph()
set_paragraph_space(spacer3, 4, 2)

# ─── BAND 4: FINANCIAL SUMMARY ─────────────────────────────────────────────────

# Section header
fin_header_table = doc.add_table(rows=1, cols=1)
fin_header_table.autofit = False
fin_header_table.columns[0].width = Cm(18)
fhc = fin_header_table.rows[0].cells[0]
set_cell_background(fhc, SECTION_HEADER_BG)
remove_cell_borders(fhc)
fhp = fhc.paragraphs[0]
fhp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(fhp, 3, 3)
add_run_with_font(fhp, " " + get_value("FINANCIAL SUMMARY"), bold=True, size_pt=11, color_hex=WHITE)

# Summary boxes row
sum_table = doc.add_table(rows=1, cols=3)
sum_table.autofit = False
set_table_borders(sum_table, color=BORDER_GRAY, sz=4)

box_w = Cm(6)
for i in range(3):
    sum_table.columns[i].width = box_w

sum_cells = sum_table.rows[0].cells
for i in range(3):
    sum_cells[i].width = box_w
    set_cell_vertical_align(sum_cells[i], "top")

# Box 1: GROSS TOTAL
set_cell_background(sum_cells[0], LIGHT_GRAY_BG)
b1p1 = sum_cells[0].paragraphs[0]
b1p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b1p1, 3, 0)
add_run_with_font(b1p1, "GROSS TOTAL", bold=True, size_pt=8, color_hex=MID_GRAY)
b1p2 = sum_cells[0].add_paragraph()
b1p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b1p2, 0, 0)
add_run_with_font(b1p2, get_value("GROSS TOTAL"), bold=True, size_pt=18, color_hex=DARK_GRAY_TEXT)
b1p3 = sum_cells[0].add_paragraph()
b1p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b1p3, 0, 3)

gross_subtitle = get_value("Clinic-billed services", "Clinic-billed services")
add_run_with_font(b1p3, gross_subtitle, bold=False, size_pt=8, color_hex=MID_GRAY)

# Box 2: CONTRACTUAL DISCOUNTS
set_cell_background(sum_cells[1], LIGHT_GRAY_BG)
b2p1 = sum_cells[1].paragraphs[0]
b2p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b2p1, 3, 0)
add_run_with_font(b2p1, "CONTRACTUAL DISCOUNTS", bold=True, size_pt=8, color_hex=MID_GRAY)
b2p2 = sum_cells[1].add_paragraph()
b2p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b2p2, 0, 0)
add_run_with_font(b2p2, get_value("CONTRACTUAL DISCOUNTS"), bold=True, size_pt=18, color_hex=DARK_GRAY_TEXT)
b2p3 = sum_cells[1].add_paragraph()
b2p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b2p3, 0, 3)
discount_subtitle = get_value("Applied per consultation and injections", "Applied per consultation and injections")
add_run_with_font(b2p3, discount_subtitle, bold=False, size_pt=8, color_hex=MID_GRAY)

# Box 3: NET AMOUNT REQUESTED
set_cell_background(sum_cells[2], LIGHT_TEAL_BG)
b3p1 = sum_cells[2].paragraphs[0]
b3p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b3p1, 3, 0)
add_run_with_font(b3p1, "NET AMOUNT REQUESTED", bold=True, size_pt=8, color_hex=MID_GRAY)
b3p2 = sum_cells[2].add_paragraph()
b3p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b3p2, 0, 0)
add_run_with_font(b3p2, get_value("NET AMOUNT REQUESTED"), bold=True, size_pt=18, color_hex=TEAL)
b3p3 = sum_cells[2].add_paragraph()
b3p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(b3p3, 0, 3)
net_subtitle = get_value("After discount and deductible/copay", "After discount and deductible/copay")
add_run_with_font(b3p3, net_subtitle, bold=False, size_pt=8, color_hex=MID_GRAY)

# Services table
svc_table = doc.add_table(rows=5, cols=6)
svc_table.autofit = False
set_table_borders(svc_table, color=BORDER_GRAY, sz=4)

col_widths = [Cm(5.5), Cm(1.5), Cm(2.5), Cm(2.5), Cm(2.8), Cm(3.2)]
for i, w in enumerate(col_widths):
    svc_table.columns[i].width = w

headers = ["Service / Treatment", "Qty.", "Gross Rate", "Reduction", "Eligible Amount", "Observations"]

for j, h in enumerate(headers):
    hc = svc_table.rows[0].cells[j]
    hc.width = col_widths[j]
    set_cell_background(hc, LIGHT_GRAY_BG)
    set_cell_vertical_align(hc, "center")
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_space(hp, 2, 2)
    add_run_with_font(hp, h, bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

# Service rows data
svc_rows = []

svc1_name = get_value("ENT Consultation - Medical Specialist", "ENT Consultation - Medical Specialist")
svc1_qty = get_value("Quantity_1", "1")
svc1_rate = get_value("Gross Rate_1", "15.00 BD")
svc1_red = get_value("Reduction_1", "3.00 BD")
svc1_elig = get_value("Eligible Amount_1", "12.00 BD")
svc1_obs = get_value("Observations_1", "Approved")
svc_rows.append((svc1_name, svc1_qty, svc1_rate, svc1_red, svc1_elig, svc1_obs))

svc2_name = get_value("Lincocin Injection", "Lincocin Injection")
svc2_qty = get_value("Quantity_2", "1")
svc2_rate = get_value("Gross Rate_2", "10.00 BD")
svc2_red = get_value("Reduction_2", "2.00 BD")
svc2_elig = get_value("Eligible Amount_2", "8.00 BD")
svc2_obs = get_value("Observations_2", "Medical indication")
svc_rows.append((svc2_name, svc2_qty, svc2_rate, svc2_red, svc2_elig, svc2_obs))

svc3_name = get_value("Depo-Medrol Injection", "Depo-Medrol Injection")
svc3_qty = get_value("Quantity_3", "1")
svc3_rate = get_value("Gross Rate_3", "10.00 BD")
svc3_red = get_value("Reduction_3", "2.00 BD")
svc3_elig = get_value("Eligible Amount_3", "8.00 BD")
svc3_obs = get_value("Observations_3", "Medical indication")
svc_rows.append((svc3_name, svc3_qty, svc3_rate, svc3_red, svc3_elig, svc3_obs))

for ri, (name, qty, rate, red, elig, obs) in enumerate(svc_rows):
    row_idx = ri + 1
    row_data = [name, qty, rate, red, elig, obs]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    for j, (val, align) in enumerate(zip(row_data, aligns)):
        cell = svc_table.rows[row_idx].cells[j]
        cell.width = col_widths[j]
        set_cell_vertical_align(cell, "center")
        if ri % 2 == 1:
            set_cell_background(cell, LIGHT_GRAY_BG)
        cp = cell.paragraphs[0]
        cp.alignment = align
        set_paragraph_space(cp, 2, 2)
        add_run_with_font(cp, val, bold=False, size_pt=9, color_hex=DARK_GRAY_TEXT)

# Total row (row 4)
total_row = svc_table.rows[4]
tc0 = total_row.cells[0]
tc0.width = col_widths[0]
set_cell_background(tc0, LIGHT_GRAY_BG)
tp0 = tc0.paragraphs[0]
tp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(tp0, 2, 2)
add_run_with_font(tp0, get_value("Total", "Total"), bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

tc1 = total_row.cells[1]
tc1.width = col_widths[1]
set_cell_background(tc1, LIGHT_GRAY_BG)
tp1 = tc1.paragraphs[0]
tp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_space(tp1, 2, 2)
add_run_with_font(tp1, "", bold=False, size_pt=9, color_hex=DARK_GRAY_TEXT)

tc2 = total_row.cells[2]
tc2.width = col_widths[2]
set_cell_background(tc2, LIGHT_GRAY_BG)
tp2 = tc2.paragraphs[0]
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_space(tp2, 2, 2)
gross_total_val = get_value("GROSS TOTAL", "35.00 BD")
add_run_with_font(tp2, gross_total_val, bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

tc3 = total_row.cells[3]
tc3.width = col_widths[3]
set_cell_background(tc3, LIGHT_GRAY_BG)
tp3 = tc3.paragraphs[0]
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_space(tp3, 2, 2)
contract_disc = get_value("CONTRACTUAL DISCOUNTS", "7.00 BD")
add_run_with_font(tp3, contract_disc, bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

tc4 = total_row.cells[4]
tc4.width = col_widths[4]
set_cell_background(tc4, LIGHT_GRAY_BG)
tp4 = tc4.paragraphs[0]
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_space(tp4, 2, 2)
add_run_with_font(tp4, get_value("Total Eligible Amount", "28.00 BD"), bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

tc5 = total_row.cells[5]
tc5.width = col_widths[5]
set_cell_background(tc5, LIGHT_GRAY_BG)
tp5 = tc5.paragraphs[0]
tp5.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_space(tp5, 2, 2)
add_run_with_font(tp5, get_value("Before deductible", "Before deductible"), bold=True, size_pt=9, color_hex=DARK_GRAY_TEXT)

# Deductible row
ded_table = doc.add_table(rows=2, cols=2)
ded_table.autofit = False
set_table_borders(ded_table, color=BORDER_GRAY, sz=4)
ded_table.columns[0].width = Cm(13.0)
ded_table.columns[1].width = Cm(5.0)

ded_r0c