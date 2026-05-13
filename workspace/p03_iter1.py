import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_03.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p03.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key):
    return lookup.get(key, "")

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement("w:" + side)
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_borders(table, color="D0D0D0", sz="4"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_run_bold(para, text, size=10, bold=True, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run

def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def cell_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, italic=False):
    para = cell.paragraphs[0]
    para.alignment = align
    set_para_spacing(para, 1, 1)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para

def add_section_header(doc, text, bg="1A3C6B", fg="FFFFFF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, 2, 2)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(fg)
    set_table_borders(tbl, color="1A3C6B")
    return tbl

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

# ── HEADER TABLE ──────────────────────────────────────────────────────────────
header_tbl = doc.add_table(rows=1, cols=2)
header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(header_tbl, color="CCCCCC", sz="4")

left_w = Cm(11.5)
right_w = Cm(6.5)
header_tbl.columns[0].width = left_w
header_tbl.columns[1].width = right_w

left_cell = header_tbl.rows[0].cells[0]
right_cell = header_tbl.rows[0].cells[1]

left_cell.width = left_w
right_cell.width = right_w

set_cell_bg(left_cell, "FFFFFF")
set_cell_bg(right_cell, "F0F4FA")

# Left cell: logo + title
logo_tbl = left_cell.add_table(rows=1, cols=2)
logo_tbl.columns[0].width = Cm(1.5)
logo_tbl.columns[1].width = Cm(10.0)
logo_cell = logo_tbl.rows[0].cells[0]
title_cell = logo_tbl.rows[0].cells[1]
logo_cell.width = Cm(1.5)
title_cell.width = Cm(10.0)
set_cell_bg(logo_cell, "2C3E50")
logo_para = logo_cell.paragraphs[0]
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(logo_para, 8, 8)
logo_run = logo_para.add_run("MC")
logo_run.bold = True
logo_run.font.size = Pt(14)
logo_run.font.color.rgb = RGBColor.from_string("FFFFFF")

title_para = title_cell.paragraphs[0]
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(title_para, 2, 0)
t_run = title_para.add_run(get_value("Approval Report and Coverage Details"))
t_run.bold = True
t_run.font.size = Pt(14)
t_run.font.color.rgb = RGBColor.from_string("1A3C6B")

sub_para = title_cell.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(sub_para, 0, 4)
sub_run = sub_para.add_run("Summary for the insurance company - eligibility, documents and payment conditions")
sub_run.bold = False
sub_run.font.size = Pt(8)
sub_run.font.color.rgb = RGBColor.from_string("555555")

# Right cell: 4-row info table
info_data = [
    ("Approval Ref.", get_value("Approval Ref.")),
    ("Approval Date", get_value("Approval Date")),
    ("Valid Until", get_value("Valid Until")),
    ("Status", get_value("Status")),
]
info_tbl = right_cell.add_table(rows=4, cols=2)
set_table_borders(info_tbl, color="CCCCCC", sz="2")
for i, (lbl, val) in enumerate(info_data):
    lc = info_tbl.rows[i].cells[0]
    vc = info_tbl.rows[i].cells[1]
    lc.width = Cm(2.8)
    vc.width = Cm(3.7)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(lp, 1, 1)
    lr = lp.add_run(lbl)
    lr.bold = True
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor.from_string("333333")
    vp = vc.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(vp, 1, 1)
    vr = vp.add_run(val)
    vr.bold = True
    vr.font.size = Pt(8)
    if val == "Approved":
        vr.font.color.rgb = RGBColor.from_string("1A3C6B")
    else:
        vr.font.color.rgb = RGBColor.from_string("111111")

sp = doc.add_paragraph()
set_para_spacing(sp, 2, 2)

# ── SECTION 2: REFERENCE DATA ─────────────────────────────────────────────────
add_section_header(doc, get_value("REFERENCE DATA"))

ref_tbl = doc.add_table(rows=3, cols=2)
ref_tbl.style = "Table Grid"
set_table_borders(ref_tbl, color="D0D0D0", sz="4")
ref_tbl.columns[0].width = Cm(9)
ref_tbl.columns[1].width = Cm(9)

ref_fields = [
    ("INDIVIDUAL NUMBER", "CARD HOLDER NAME"),
    ("DATE OF BIRTH", "INSURANCE COMPANY"),
    ("PHYSICIAN / PROVIDER", "SPECIALTY"),
]
for r_idx, (left_key, right_key) in enumerate(ref_fields):
    lc = ref_tbl.rows[r_idx].cells[0]
    rc = ref_tbl.rows[r_idx].cells[1]
    lc.width = Cm(9)
    rc.width = Cm(9)
    set_cell_bg(lc, "FFFFFF")
    set_cell_bg(rc, "FFFFFF")
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(lp, 2, 0)
    lr = lp.add_run(left_key)
    lr.bold = True
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor.from_string("555555")
    lp2 = lc.add_paragraph()
    set_para_spacing(lp2, 0, 2)
    lr2 = lp2.add_run(get_value(left_key))
    lr2.bold = True
    lr2.font.size = Pt(9)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(rp, 2, 0)
    rr = rp.add_run(right_key)
    rr.bold = True
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor.from_string("555555")
    rp2 = rc.add_paragraph()
    set_para_spacing(rp2, 0, 2)
    rr2 = rp2.add_run(get_value(right_key))
    rr2.bold = True
    rr2.font.size = Pt(9)

sp2 = doc.add_paragraph()
set_para_spacing(sp2, 2, 2)

# ── SECTION 3: APPROVED SERVICES ──────────────────────────────────────────────
add_section_header(doc, get_value("APPROVED SERVICES"))

svc_tbl = doc.add_table(rows=4, cols=6)
svc_tbl.style = "Table Grid"
set_table_borders(svc_tbl, color="D0D0D0", sz="4")

col_widths = [Cm(3.2), Cm(5.5), Cm(1.5), Cm(1.5), Cm(2.0), Cm(4.3)]
for i, w in enumerate(col_widths):
    svc_tbl.columns[i].width = w

headers_svc = ["Service Category", "Approved Item", "Req. Qty", "App. Qty", "Status", "Reason / Notes"]
header_row = svc_tbl.rows[0]
for i, h in enumerate(headers_svc):
    c = header_row.cells[i]
    c.width = col_widths[i]
    set_cell_bg(c, "E8ECF2")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, 2, 2)
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("1A3C6B")

svc_rows_data = [
    get_value("Service Category / Approved Item / Requested Qty / Approved Qty / Status / Reason-Notes - Row 1"),
    get_value("Service Category / Approved Item / Requested Qty / Approved Qty / Status / Reason-Notes - Row 2"),
    get_value("Service Category / Approved Item / Requested Qty / Approved Qty / Status / Reason-Notes - Row 3"),
]

for r_idx, row_val in enumerate(svc_rows_data):
    parts = [p.strip() for p in row_val.split("|")]
    while len(parts) < 6:
        parts.append("")
    row = svc_tbl.rows[r_idx + 1]
    for c_idx, part in enumerate(parts):
        cell = row.cells[c_idx]
        cell.width = col_widths[c_idx]
        set_cell_bg(cell, "FFFFFF")
        p = cell.paragraphs[0]
        if c_idx in [2, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, 2, 2)
        if c_idx == 4:
            r = p.add_run(part)
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string("1A6B3C")
        else:
            r = p.add_run(part)
            r.bold = False
            r.font.size = Pt(8.5)

sp3 = doc.add_paragraph()
set_para_spacing(sp3, 2, 2)

# ── SECTION 4 & 5: DOCUMENTS + DIAGNOSTIC (side by side) ─────────────────────
two_col_tbl = doc.add_table(rows=1, cols=2)
two_col_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(two_col_tbl, color="D0D0D0", sz="4")
two_col_tbl.columns[0].width = Cm(9.0)
two_col_tbl.columns[1].width = Cm(9.0)

left_docs_cell = two_col_tbl.rows[0].cells[0]
right_diag_cell = two_col_tbl.rows[0].cells[1]
left_docs_cell.width = Cm(9.0)
right_diag_cell.width = Cm(9.0)

set_cell_bg(left_docs_cell, "FFFFFF")
set_cell_bg(right_diag_cell, "FFFFFF")

# Left: Documents header + list
ldp = left_docs_cell.paragraphs[0]
set_para_spacing(ldp, 2, 2)
ldr = ldp.add_run(get_value("DOCUMENTS REQUIRED FOR CLAIM SUBMISSION"))
ldr.bold = True
ldr.font.size = Pt(9)
ldr.font.color.rgb = RGBColor.from_string("FFFFFF")
ldp_bg_para = left_docs_cell.paragraphs[0]

# Rebuild left cell with header styling
left_docs_cell.paragraphs[0].clear()
hdr_p = left_docs_cell.paragraphs[0]
set_para_spacing(hdr_p, 3, 3)
set_cell_bg(left_docs_cell, "F8F9FA")
hdr_r = hdr_p.add_run(get_value("DOCUMENTS REQUIRED FOR CLAIM SUBMISSION"))
hdr_r.bold = True
hdr_r.font.size = Pt(9)
hdr_r.font.color.rgb = RGBColor.from_string("1A3C6B")

doc_items_keys = [
    "Required Documents - item 1",
    "Required Documents - item 2",
    "Required Documents - item 3",
    "Required Documents - item 4",
    "Required Documents - item 5",
    "Required Documents - item 6",
]
for i, key in enumerate(doc_items_keys):
    val = get_value(key)
    if val:
        dp = left_docs_cell.add_paragraph()
        set_para_spacing(dp, 1, 1)
        dp.paragraph_format.left_indent = Cm(0.3)
        dr = dp.add_run(str(i + 1) + ". " + val)
        dr.font.size = Pt(8.5)
        dr.bold = False

# Right: Diagnostic header + table
rdp = right_diag_cell.paragraphs[0]
set_para_spacing(rdp, 3, 3)
set_cell_bg(right_diag_cell, "F8F9FA")
rdr = rdp.add_run(get_value("DIAGNOSTIC AND MEDICAL INFORMATION"))
rdr.bold = True
rdr.font.size = Pt(9)
rdr.font.color.rgb = RGBColor.from_string("1A3C6B")

diag_inner_tbl = right_diag_cell.add_table(rows=1, cols=2)
set_table_borders(diag_inner_tbl, color="D0D0D0", sz="4")
diag_inner_tbl.columns[0].width = Cm(3.8)
diag_inner_tbl.columns[1].width = Cm(5.0)

dih = diag_inner_tbl.rows[0]
set_cell_bg(dih.cells[0], "E8ECF2")
set_cell_bg(dih.cells[1], "E8ECF2")
dih.cells[0].width = Cm(3.8)
dih.cells[1].width = Cm(5.0)
hp0 = dih.cells[0].paragraphs[0]
hp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(hp0, 2, 2)
hr0 = hp0.add_run("Field")
hr0.bold = True
hr0.font.size = Pt(8.5)
hr0.font.color.rgb = RGBColor.from_string("1A3C6B")
hp1 = dih.cells[1].paragraphs[0]
hp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(hp1, 2, 2)
hr1 = hp1.add_run("Value")
hr1.bold = True
hr1.font.size = Pt(8.5)
hr1.font.color.rgb = RGBColor.from_string("1A3C6B")

diag_fields = [
    ("Primary Diagnosis", get_value("Primary Diagnosis")),
    ("Visit Date", get_value("Visit Date")),
    ("Blood Pressure / Pulse / Temperature", get_value("Blood Pressure / Pulse / Temperature")),
    ("Duration of Illness", get_value("Duration of Illness")),
    ("Uploaded Files", get_value("Uploaded Files")),
]
for field_lbl, field_val in diag_fields:
    new_row = diag_inner_tbl.add_row()
    new_row.cells[0].width = Cm(3.8)
    new_row.cells[1].width = Cm(5.0)
    set_cell_bg(new_row.cells[0], "FFFFFF")
    set_cell_bg(new_row.cells[1], "FFFFFF")
    fp = new_row.cells[0].paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(fp, 2, 2)
    fr = fp.add_run(field_lbl)
    fr.font.size = Pt(8.5)
    fr.bold = False
    vp = new_row.cells[1].paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(vp, 2, 2)
    vr = vp.add_run(field_val)
    vr.font.size = Pt(8.5)
    vr.bold = False

sp4 = doc.add_paragraph()
set_para_spacing(sp4, 2, 2)

# ── SECTION 6: CONDITIONS ─────────────────────────────────────────────────────
add_section_header(doc, get_value("ASSESSMENT AND PAYMENT CONDITIONS"))

cond_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(cond_tbl, color="D0D0D0", sz="4")
cond_cell = cond_tbl.rows[0].cells[0]
set_cell_bg(cond_cell, "FFFFFF")

cond_texts = [
    get_value("Payment Conditions - paragraph 1"),
    get_value("Payment Conditions - paragraph 2"),
]
for i, txt in enumerate(cond_texts):
    if i == 0:
        cp = cond_cell.paragraphs[0]
    else:
        cp = cond_cell.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(cp, 2, 3)
    cr = cp.add_run(txt)
    cr.font.size = Pt(8.5)
    cr.bold = False

vat_key = "VAT Note"
vat_val = get_value(vat_key)
if vat_val:
    vat_p = cond_cell.add_paragraph()
    set_para_spacing(vat_p, 2, 2)
    vat_label = vat_p.add_run("VAT Note: ")
    vat_label.bold = True
    vat_label.font.size = Pt(8.5)
    vat_rest = vat_p.add_run(vat_val)
    vat_rest.bold = False
    vat_rest.font.size = Pt(8.5)

sp5 = doc.add_paragraph()
set_para_spacing(sp5, 2, 2)

# ── SECTION 7: INTERNAL USE + QUALITY CONTROL (side by side) ─────────────────
bottom_tbl = doc.add_table(rows=1, cols=2)
bottom_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(bottom_tbl, color="D0D0D0", sz="4")
bottom_tbl.columns[0].width = Cm(9.0)
bottom_tbl.columns[1].width = Cm(9.0)

int_cell = bottom_tbl.rows[0].cells[0]
qc_cell = bottom_tbl.rows[0].cells[1]
int_cell.width = Cm(9.0)
qc_cell.width = Cm(9.0)
set_cell_bg(int_cell, "FAFAFA")
set_cell_bg(qc_cell, "FAFAFA")

int_hdr_p = int_cell.paragraphs[0]
set_para_spacing(int_hdr_p, 2, 3)
int_hdr_r = int_hdr_p.add_run(get_value("Internal Use - Insurance Company"))
int_hdr_r.bold = True
int_hdr_r.font.size = Pt(9)
int_hdr_r.font.color.rgb = RGBColor.from_string("1A3C6B")

decision_lbl_p = int_cell.add_paragraph()
set_para_spacing(decision_lbl_p, 2, 1)
decision_lbl_r = decision_lbl_p.add_run(get_value("Decision Label"))
decision_lbl_r.bold = True
decision_lbl_r.font.size = Pt(9)

approved_p = int_cell.add_paragraph()
set_para_spacing(approved_p, 1, 1)
approved_p.paragraph_format.left_indent = Cm(0.5)
approved_r = approved_p.add_run("\u2611 " + get_value("Decision - Approved option"))
approved_r.font.size = Pt(9)
approved_r.bold = False

respins_p = int_cell.add_paragraph()
set_para_spacing(respins_p, 1, 3)
respins_p.paragraph_format.left_indent = Cm(0.5)
respins_r = respins_p.add_run("\u2610 " + get_value("Decision - Rejected option"))
respins_r.font.size = Pt(9)
respins_r.bold = False

nr_p = int_cell.add_paragraph()
set_para_spacing(nr_p, 2, 1)
nr_lbl = nr_p.add_run(get_value("Approval Number Label") + "  ")
nr_lbl.bold = True
nr_lbl.font.size = Pt(9)
nr_val = nr_p.add_run(get_value("Approval Number Value"))
nr_val.bold = False
nr_val.font.size = Pt(9)

comm_p = int_cell.add_paragraph()
set_para_spacing(comm_p, 2, 1)
comm_lbl = comm_p.add_run(get_value("Comments Label") + "  ")
comm_lbl.bold = True
comm_lbl.font.size = Pt(9)
comm_val = comm_p.add_run(get_value("Comments Value"))
comm_val.bold = False
comm_val.font.size = Pt(9)

qc_hdr_p = qc_cell.paragraphs[0]
set_para_spacing(qc_hdr_p, 2, 3)
qc_hdr_r = qc_hdr_p.add_run(get_value("Quality Control Claim"))
qc_hdr_r.bold = True
qc_hdr_r.font.size = Pt(9)
qc_hdr_r.font.color.rgb = RGBColor.from_string("1A3C6B")

qc_desc_p = qc_cell.add_paragraph()
set_para_spacing(qc_desc_p, 1, 6)
qc_desc_r = qc_desc_p.add_run(get_value("Quality Control Description"))
qc_desc_r.font.size = Pt(8)
qc_desc_r.bold = False
qc_desc_r.font.color.rgb = RGBColor.from_string("555555")

line_p = qc_cell.add_paragraph()
set_para_spacing(line_p, 8, 2)
line_r = line_p.add_run("_" * 45)
line_r.font.size = Pt(8)
line_r.font.color.rgb = RGBColor.from_string("999999")

sig_p = qc_cell.add_paragraph()
set_para_spacing(sig_p, 1, 2)
sig_r = sig_p.add_run(get_value("Evaluator Name / Signature / Date"))
sig_r.font.size = Pt(8)
sig_r.bold = False
sig_r.font.color.rgb = RGBColor.from_string("555555")

sp6 = doc.add_paragraph()
set_para_spacing(sp6, 2, 2)

# ── FOOTER ────────────────────────────────────────────────────────────────────
footer_tbl = doc.add_table(rows=1, cols=2)
set_table_borders(footer_tbl, color="DDDDDD", sz="2")
footer_tbl.columns[0].width = Cm(9.0)
footer_tbl.columns[1].width = Cm(9.0)
fl_cell = footer_tbl.rows[0].cells[0]
fr_cell = footer_tbl.rows[0].cells[1]
set_cell_bg(fl_cell, "F5F5F5")
set_cell_bg(fr_cell, "F5F5F5")
fl_p = fl_cell.paragraphs[0]
fl_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(fl_p, 2, 2)
fl_r = fl_p.add_run("Medical claim form - EN | p. 3/3")
fl_r.font.size = Pt(7.5)
fl_r.font.color.rgb = RGBColor.from_string("888888")
fr_p = fr_cell.paragraphs[0]
fr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_spacing(fr_p, 2, 2)
fr_r = fr_p.add_run("Model document with anonymized data")
fr_r.font.size = Pt(7.5)
fr_r.font.color.rgb = RGBColor.from_string("888888")

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)