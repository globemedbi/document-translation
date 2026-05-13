import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_01.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p01.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key):
    return lookup.get(key, "")

NAVY = RGBColor(0x1e, 0x3a, 0x5f)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x44, 0x44, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

def set_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)

def add_run_fmt(para, text, bold=False, size=10, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return run

def add_horizontal_rule(doc, color_hex="1e3a5f", thickness=12):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=2, after=4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def add_section_header(doc, text):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=8, after=2)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    # Add bottom border to simulate underline
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1e3a5f")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "none"))
            el.set(qn("w:sz"), str(val.get("sz", 0)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

# ─── HEADER BLOCK ───────────────────────────────────────────────────────────

# Full Name
para_name = doc.add_paragraph()
set_paragraph_spacing(para_name, before=0, after=2)
run_name = para_name.add_run(get_value("Nom complet"))
run_name.bold = True
run_name.font.size = Pt(28)
run_name.font.color.rgb = DARK
run_name.font.name = "Calibri"

# Title / Role
para_title = doc.add_paragraph()
set_paragraph_spacing(para_title, before=0, after=2)
run_title = para_title.add_run(get_value("Titre / Rôle"))
run_title.font.size = Pt(11)
run_title.font.color.rgb = NAVY
run_title.font.name = "Calibri"

# Contact line
para_contact = doc.add_paragraph()
set_paragraph_spacing(para_contact, before=0, after=4)
contact_parts = [
    get_value("Localisation"),
    get_value("Téléphone"),
    get_value("E-mail"),
    get_value("LinkedIn"),
]
contact_text = "  \u2022  ".join([p for p in contact_parts if p])
run_contact = para_contact.add_run(contact_text)
run_contact.font.size = Pt(9)
run_contact.font.color.rgb = GRAY
run_contact.font.name = "Calibri"

# Horizontal rule
add_horizontal_rule(doc, color_hex="1e3a5f", thickness=12)

# ─── PROFESSIONAL SUMMARY ───────────────────────────────────────────────────

add_section_header(doc, get_value("RÉSUMÉ PROFESSIONNEL"))

para_summary = doc.add_paragraph()
set_paragraph_spacing(para_summary, before=4, after=4)
run_summary = para_summary.add_run(get_value("Texte du résumé professionnel"))
run_summary.font.size = Pt(10)
run_summary.font.color.rgb = BLACK
run_summary.font.name = "Calibri"

# ─── CORE SKILLS ────────────────────────────────────────────────────────────

add_section_header(doc, get_value("COMPÉTENCES CLÉS"))

# 2-column table for skills
skills_table = doc.add_table(rows=1, cols=2)
skills_table.alignment = WD_TABLE_ALIGNMENT.LEFT
skills_table.style = "Table Grid"

# Set column widths
col_width = Cm(8.5)
for col in skills_table.columns:
    for cell in col.cells:
        cell.width = col_width

# Remove all borders
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

remove_table_borders(skills_table)

left_cell = skills_table.rows[0].cells[0]
right_cell = skills_table.rows[0].cells[1]

# Left column: AI/ML, Data & Infrastructure, Architecture & Leadership
def add_skill_line(cell, label, value, first=False):
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before=1 if not first else 3, after=1)
    run_label = para.add_run(label + ": ")
    run_label.bold = True
    run_label.font.size = Pt(9.5)
    run_label.font.name = "Calibri"
    run_val = para.add_run(value)
    run_val.font.size = Pt(9.5)
    run_val.font.name = "Calibri"

# Clear default paragraph in cells
for cell in [left_cell, right_cell]:
    for para in cell.paragraphs:
        p = para._p
        p.getparent().remove(p)

add_skill_line(left_cell, "IA / ML", get_value("IA / ML"), first=True)
add_skill_line(left_cell, "Données et infrastructure", get_value("Données et infrastructure"))
add_skill_line(left_cell, "Architecture et leadership", get_value("Architecture et leadership"))

add_skill_line(right_cell, "Programmation", get_value("Programmation"), first=True)
add_skill_line(right_cell, "Frameworks et outils", get_value("Frameworks et outils"))

# ─── LANGUAGES ──────────────────────────────────────────────────────────────

add_section_header(doc, get_value("LANGUES"))

para_lang = doc.add_paragraph()
set_paragraph_spacing(para_lang, before=4, after=4)

# Language entries: French, English, Arabic with their levels
french_name = get_value("Français")
french_level = get_value("Niveau Français")
english_name = get_value("Anglais")
english_level = get_value("Niveau Anglais")
arabic_name = get_value("Arabe")
arabic_level = get_value("Niveau Arabe")

# Fallback to known keys if empty
if not french_name:
    french_name = "Français"
if not french_level:
    french_level = get_value("Niveau - Français")
if not english_name:
    english_name = "Anglais"
if not english_level:
    english_level = get_value("Niveau - Anglais")
if not arabic_name:
    arabic_name = "Arabe"
if not arabic_level:
    arabic_level = get_value("Niveau - Arabe")

# Try alternative key patterns
lang_keys = [d["key"] for d in DATA]
for k in lang_keys:
    val = get_value(k)
    k_lower = k.lower()
    if "français" in k_lower or "french" in k_lower:
        if "niveau" in k_lower or "level" in k_lower:
            french_level = val
        else:
            french_name = val
    elif "anglais" in k_lower or "english" in k_lower:
        if "niveau" in k_lower or "level" in k_lower:
            english_level = val
        else:
            english_name = val
    elif "arabe" in k_lower or "arabic" in k_lower:
        if "niveau" in k_lower or "level" in k_lower:
            arabic_level = val
        else:
            arabic_name = val

# Build language line
def add_lang_entry(para, name, level, separator=False):
    if separator:
        run_sep = para.add_run("  \u2022  ")
        run_sep.font.size = Pt(10)
        run_sep.font.name = "Calibri"
    run_name = para.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(10)
    run_name.font.name = "Calibri"
    if level:
        run_dash = para.add_run(" \u2014 ")
        run_dash.font.size = Pt(10)
        run_dash.font.name = "Calibri"
        run_level = para.add_run(level)
        run_level.italic = True
        run_level.font.size = Pt(10)
        run_level.font.name = "Calibri"

# Check DATA for language-related entries directly
lang_data = {}
for d in DATA:
    k = d["key"]
    v = d["value"]
    ok = d.get("original_key", "")
    if ok in ["French", "English", "Arabic"]:
        lang_data[ok + "_name"] = v
    if ok in ["French Level", "English Level", "Arabic Level"]:
        lang_data[ok.replace(" Level", "") + "_level"] = v
    # Also check original_value patterns
    ov = d.get("original_value", "")
    if ov in ["Native", "Professional", "Fluent"] or "Native" in str(ov) or "Professional" in str(ov):
        if "french" in ok.lower() or "french" in k.lower():
            lang_data["French_level"] = v
        elif "english" in ok.lower() or "english" in k.lower():
            lang_data["English_level"] = v
        elif "arabic" in ok.lower() or "arabic" in k.lower():
            lang_data["Arabic_level"] = v

# Use direct key lookup for language names
f_name = lang_data.get("French_name", french_name)
f_level = lang_data.get("French_level", french_level)
e_name = lang_data.get("English_name", english_name)
e_level = lang_data.get("English_level", english_level)
a_name = lang_data.get("Arabic_name", arabic_name)
a_level = lang_data.get("Arabic_level", arabic_level)

# Final fallback: scan all data for language entries
for d in DATA:
    k = d["key"]
    v = d["value"]
    ov = d.get("original_value", "")
    ok = d.get("original_key", "")
    if ov == "French" or ok == "French":
        f_name = v
    elif ov == "English" or ok == "English":
        e_name = v
    elif ov == "Arabic" or ok == "Arabic":
        a_name = v
    elif ov == "Native" or ok == "French Level":
        f_level = v
    elif ov == "Professional" and ("english" in ok.lower()):
        e_level = v
    elif ov == "Professional" and ("arabic" in ok.lower()):
        a_level = v

add_lang_entry(para_lang, f_name, f_level)
add_lang_entry(para_lang, e_name, e_level, separator=True)
add_lang_entry(para_lang, a_name, a_level, separator=True)

# ─── EDUCATION ──────────────────────────────────────────────────────────────

add_section_header(doc, get_value("FORMATION"))

# Education entries - find them from DATA
edu_entries = []
for d in DATA:
    ok = d.get("original_key", "")
    k = d["key"]
    v = d["value"]
    ov = d.get("original_value", "")
    # Look for degree entries
    if "degree" in ok.lower() or "master" in ok.lower() or "bachelor" in ok.lower() or "diplôme" in k.lower() or "master" in k.lower() or "bachelor" in k.lower() or "licence" in k.lower():
        pass

# Direct key-based lookup for education
edu1_degree = get_value("Diplôme 1")
edu1_institution = get_value("Établissement 1")
edu1_location = get_value("Lieu formation 1")
edu1_date = get_value("Date formation 1")

edu2_degree = get_value("Diplôme 2")
edu2_institution = get_value("Établissement 2")
edu2_location = get_value("Lieu formation 2")
edu2_date = get_value("Date formation 2")

# Try alternative keys
edu_key_map = {}
for d in DATA:
    k = d["key"]
    v = d["value"]
    ok = d.get("original_key", "")
    ov = d.get("original_value", "")
    ok_lower = ok.lower()
    k_lower = k.lower()
    if "master" in ov.lower() or "master" in ok_lower:
        if "degree" in ok_lower or "diplôme" in k_lower or "master" in ok_lower:
            edu_key_map["edu1_degree"] = v
    if "bachelor" in ov.lower() or "bachelor" in ok_lower:
        if "degree" in ok_lower or "diplôme" in k_lower or "bachelor" in ok_lower:
            edu_key_map["edu2_degree"] = v
    if "epita" in ov.lower() or "epita" in str(v).lower():
        edu_key_map["edu1_institution"] = v
    if "usj" in ov.lower() or "saint joseph" in ov.lower() or "usj" in str(v).lower():
        edu_key_map["edu2_institution"] = v
    if ok == "Education Date 1" or (ok == "Degree 1 Date"):
        edu_key_map["edu1_date"] = v
    if ok == "Education Date 2" or (ok == "Degree 2 Date"):
        edu_key_map["edu2_date"] = v
    if ok == "Education Location 1":
        edu_key_map["edu1_location"] = v
    if ok == "Education Location 2":
        edu_key_map["edu2_location"] = v

# Scan all data for education-related original values
for d in DATA:
    ok = d.get("original_key", "")
    ov = d.get("original_value", "")
    v = d["value"]
    k = d["key"]
    if "Master of Science" in ov or "Master" in ov and "Artificial" in ov:
        edu_key_map["edu1_degree"] = v
    if "Bachelor of Science" in ov or "Bachelor" in ov and "Management" in ov:
        edu_key_map["edu2_degree"] = v
    if "EPITA" in ov:
        edu_key_map["edu1_institution"] = v
    if "USJ" in ov or "Universit" in ov and "Saint" in ov:
        edu_key_map["edu2_institution"] = v
    if "Jun 2025" in ov:
        edu_key_map["edu1_date"] = v
    if "Jun 2020" in ov:
        edu_key_map["edu2_date"] = v
    if "Paris, France" in ov and "edu1_location" not in edu_key_map:
        # Only assign if this looks like education location
        if "formation" in k.lower() or "education" in ok.lower() or "location" in ok.lower():
            edu_key_map["edu1_location"] = v
    if "Beirut, Lebanon" in ov and "edu2_location" not in edu_key_map:
        if "formation" in k.lower() or "education" in ok.lower() or "location" in ok.lower():
            edu_key_map["edu2_location"] = v

if not edu1_degree:
    edu1_degree = edu_key_map.get("edu1_degree", "Master of Science en systèmes d'intelligence artificielle")
if not edu1_institution:
    edu1_institution = edu_key_map.get("edu1_institution", "EPITA : École pour l'informatique et les Technologies Avancées")
if not edu1_location:
    edu1_location = edu_key_map.get("edu1_location", "Paris, France")
if not edu1_date:
    edu1_date = edu_key_map.get("edu1_date", "Juin 2025")

if not edu2_degree:
    edu2_degree = edu_key_map.get("edu2_degree", "Licence en systèmes d'information de gestion")
if not edu2_institution:
    edu2_institution = edu_key_map.get("edu2_institution", "USJ : Université Saint Joseph de Beyrouth")
if not edu2_location:
    edu2_location = edu_key_map.get("edu2_location", "Beyrouth, Liban")
if not edu2_date:
    edu2_date = edu_key_map.get("edu2_date", "Juin 2020")

def add_education_entry(doc, degree, institution, location, date):
    # Table with 2 cols: left=degree+institution+location, right=date
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(tbl)

    left_w = Cm(13)
    right_w = Cm(4)

    left_c = tbl.rows[0].cells[0]
    right_c = tbl.rows[0].cells[1]
    left_c.width = left_w
    right_c.width = right_w

    # Clear default paragraphs
    for cell in [left_c, right_c]:
        for para in cell.paragraphs:
            p = para._p
            p.getparent().remove(p)

    # Degree line
    para_deg = left_c.add_paragraph()
    set_paragraph_spacing(para_deg, before=2, after=0)
    run_deg = para_deg.add_run(degree)
    run_deg.bold = True
    run_deg.font.size = Pt(10)
    run_deg.font.name = "Calibri"

    # Location line
    if location:
        para_loc = left_c.add_paragraph()
        set_paragraph_spacing(para_loc, before=0, after=0)
        run_loc = para_loc.add_run(location)
        run_loc.italic = True
        run_loc.font.size = Pt(9.5)
        run_loc.font.name = "Calibri"

    # Institution line
    if institution:
        para_inst = left_c.add_paragraph()
        set_paragraph_spacing(para_inst, before=0, after=2)
        run_inst = para_inst.add_run(institution)
        run_inst.font.size = Pt(9.5)
        run_inst.font.name = "Calibri"

    # Date (right aligned)
    para_date = right_c.add_paragraph()
    set_paragraph_spacing(para_date, before=2, after=0)
    para_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = para_date.add_run(date)
    run_date.bold = True
    run_date.font.size = Pt(10)
    run_date.font.name = "Calibri"

add_education_entry(doc, edu1_degree, edu1_institution, edu1_location, edu1_date)
add_education_entry(doc, edu2_degree, edu2_institution, edu2_location, edu2_date)

# ─── PROFESSIONAL EXPERIENCE ────────────────────────────────────────────────

add_section_header(doc, get_value("EXPÉRIENCE PROFESSIONNELLE"))

# Helper to add job header with title | company and date on right
def add_job_header(doc, title, company, date, location):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(tbl)

    left_c = tbl.rows[0].cells[0]
    right_c = tbl.rows[0].cells[1]
    left_c.width = Cm(12)
    right_c.width = Cm(5)

    for cell in [left_c, right_c]:
        for para in cell.paragraphs:
            p = para._p
            p.getparent().remove(p)

    # Title | Company
    para_title = left_c.add_paragraph()
    set_paragraph_spacing(para_title, before=4, after=0)
    # Add left border to simulate the style
    pPr = para_title._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_bdr = OxmlElement("w:left")
    left_bdr.set(qn("w:val"), "single")
    left_bdr.set(qn("w:sz"), "12")
    left_bdr.set(qn("w:space"), "4")
    left_bdr.set(qn("w:color"), "1e3a5f")
    pBdr.append(left_bdr)
    pPr.append(pBdr)

    run_t = para_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(10.5)
    run_t.font.name = "Calibri"
    if company:
        run_sep = para_title.add_run("  |  ")
        run_sep.font.size = Pt(10.5)
        run_sep.font.name = "Calibri"
        run_comp = para_title.add_run(company)
        run_comp.font.size = Pt(10.5)
        run_comp.font.name = "Calibri"

    # Location (italic, below title)
    if location:
        para_loc = left_c.add_paragraph()
        set_paragraph_spacing(para_loc, before=0, after=2)
        run_loc = para_loc.add_run(location)
        run_loc.italic = True
        run_loc.font.size = Pt(9.5)
        run_loc.font.name = "Calibri"

    # Date right
    para_date = right_c.add_paragraph()
    set_paragraph_spacing(para_date, before=4, after=0)
    para_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = para_date.add_run(date)
    run_date.bold = True
    run_date.font.size = Pt(10)
    run_date.font.name = "Calibri"

def add_bullet_point(doc, text):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=1, after=1)
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.first_line_indent = Cm(-0.3)
    run_bullet = para.add_run("\u2022  ")
    run_bullet.font.size = Pt(10)
    run_bullet.font.name = "Calibri"
    run_text = para.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.name = "Calibri"
    return para

# ── Job 1: AI Engineer at Zaka/MadeBy ──────────────────────────────────────
job1_title = get_value("Poste 1")
job1_company = get_value("Entreprise 1")
job1_date = get_value("Dates poste 1")
job1_location = get_value("Lieu poste 1")

# Try alternative keys
job_key_map = {}
for d in DATA:
    ok = d.get("original_key", "")
    ov = d.get("original_value", "")
    v = d["value"]
    k = d["key"]
    if ok == "Job Title 1" or ov == "AI Engineer":
        job_key_map["job1_title"] = v
    if ok == "Company 1" or "zaka" in ov.lower() or "madeby" in ov.lower():
        job_key_map["job1_company"] = v
    if ok == "Job Date 1" or "Dec 2025" in ov:
        job_key_map["job1_date"] = v
    if ok == "Job Location 1" and "beirut" in ov.lower():
        job_key_map["job1_location"] = v

if not job1_title:
    job1_title = job_key_map.get("job1_title", "Ingénieur IA")
if not job1_company:
    job1_company = job_key_map.get("job1_company", "Zaka / MadeBy")
if not job1_date:
    job1_date = job_key_map.get("job1_date", "Déc 2025 – Présent")
if not job1_location:
    job1_location = job_key_map.get("job1_location", "Beyrouth, Liban - Télétravail")

add_job_header(doc, job1_title, job1_company, job1_date, job1_location)

# Bullet points for job 1
bullet_keys_1 = [
    "Responsabilité 1.1",
    "Responsabilité 1.2",
    "Responsabilité 1.3",
    "Responsabilité 1.4",
]

job1_bullets = []
for k in bullet_keys_1:
    v = get_value(k)
    if v:
        job1_bullets.append(v)

# Try to find bullets from DATA if direct keys fail
if not job1_bullets:
    for d in DATA:
        ok = d.get("original_key", "")
        ov = d.get("original_value", "")
        v = d["value"]
        if ok in ["Job 1 Bullet 1", "Job 1 Bullet 2", "Job 1 Bullet 3", "Job 1 Bullet 4"]:
            job1_bullets.append(v)
        elif ok in ["Bullet 1.1", "Bullet 1.2", "Bullet 1.3", "Bullet 1.4"]:
            job1_bullets.append(v)

# Fallback: scan for bullet-like content related to job 1
if not job1_bullets:
    bullet_candidates = []
    for d in DATA:
        ok = d.get("original_key", "")
        ov = d.get("original_value", "")
        v = d["value"]
        if "workedplace" in ov.lower() or "zoom" in ov.lower() or "certified ai" in ov.lower() or "visa in morocco" in ov.lower():
            bullet_candidates.append(v)
    job1_bullets = bullet_candidates

for bullet in job1_bullets:
    add_