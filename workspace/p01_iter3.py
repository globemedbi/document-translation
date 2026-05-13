import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_01.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p01.json"

with open(DATA_FILE, "r", encoding="utf-8") as f:
    data = json.load(f)

def get_value(key):
    for item in data:
        if item["key"] == key:
            return item["value"]
    return ""

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), val.get("val", "single"))
            border.set(qn("w:sz"), str(val.get("sz", 4)))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_border(table, hex_color="CCCCCC", sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hex_color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

def add_run(para, text, bold=False, italic=False, size=9, color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_para_spacing(para, before=0, after=60):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.4)
section.bottom_margin = Inches(0.4)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)

TEAL = (0, 128, 128)
DARK_GRAY = (50, 50, 50)
GRAY = (120, 120, 120)
LIGHT_GRAY = (220, 220, 220)
TEAL_HEX = "007B7B"
LIGHT_TEAL_HEX = "E8F4F4"
BORDER_GRAY = "CCCCCC"
DARK_TEAL_HEX = "005F5F"

# ─── BAND 1: HEADER BAR ───
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
header_table.width = Inches(7.27)
remove_table_borders(header_table)

left_cell = header_table.cell(0, 0)
left_cell.width = Inches(5.2)
right_cell = header_table.cell(0, 1)
right_cell.width = Inches(2.07)

# Left cell: icon + title
left_inner = left_cell.add_table(1, 2)
left_inner.autofit = False
remove_table_borders(left_inner)
icon_cell = left_inner.cell(0, 0)
icon_cell.width = Inches(0.6)
title_cell = left_inner.cell(0, 1)
title_cell.width = Inches(4.5)

# Icon
icon_para = icon_cell.paragraphs[0]
no_space_para(icon_para)
icon_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_bg(icon_cell, TEAL_HEX)
icon_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
add_run(icon_para, "MC", bold=True, size=14, color=(255, 255, 255))

# Title
title_para = title_cell.paragraphs[0]
no_space_para(title_para)
title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
add_run(title_para, get_value("Medical Receipt and Cost Summary"), bold=True, size=14, color=DARK_GRAY)

subtitle_para = title_cell.add_paragraph()
no_space_para(subtitle_para)
add_run(subtitle_para, get_value("Reimbursement File - ENT Services, Primary Care / Outpatient"), italic=True, size=8, color=GRAY)

# Remove inner table borders cleanly
for r in left_inner.rows:
    for c in r.cells:
        for side in ["top","bottom","left","right"]:
            set_cell_border(c, **{side: {"val":"none","sz":0,"color":"auto"}})

# Right cell: info box
right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
info_table = right_cell.add_table(4, 2)
info_table.autofit = False
set_table_border(info_table, BORDER_GRAY, 4)

info_data = [
    ("File No.", get_value("File No.")),
    ("Date", get_value("Date")),
    ("Currency", get_value("Currency")),
    ("Status", get_value("Status")),
]

for i, (lbl, val) in enumerate(info_data):
    lc = info_table.cell(i, 0)
    vc = info_table.cell(i, 1)
    lc.width = Inches(0.8)
    vc.width = Inches(1.0)
    lp = lc.paragraphs[0]
    no_space_para(lp)
    add_run(lp, lbl, size=8, color=GRAY)
    vp = vc.paragraphs[0]
    no_space_para(vp)
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(vp, val, bold=True, size=8, color=DARK_GRAY)

doc.add_paragraph()

# ─── BAND 2: PURPOSE BANNER ───
purpose_table = doc.add_table(rows=1, cols=1)
purpose_table.autofit = False
purpose_table.width = Inches(7.27)
set_table_border(purpose_table, BORDER_GRAY, 4)
pc = purpose_table.cell(0, 0)
set_cell_bg(pc, "F0F8F8")
pp = pc.paragraphs[0]
no_space_para(pp)
add_run(pp, "Document Purpose: ", bold=True, size=8.5, color=DARK_GRAY)
add_run(pp, "professional form in Romanian for a medical insurance file, built on the structure of a clinic/hospital claim: receipt, request form, medical details, costs, approved services and list of supporting documents. Identification data is anonymised.", size=8.5, color=DARK_GRAY)

doc.add_paragraph()

# ─── BAND 3: PROVIDER / INSURANCE SPLIT ROW ───
split_table = doc.add_table(rows=1, cols=2)
split_table.autofit = False
split_table.width = Inches(7.27)
remove_table_borders(split_table)

prov_cell = split_table.cell(0, 0)
prov_cell.width = Inches(3.5)
ins_cell = split_table.cell(0, 1)
ins_cell.width = Inches(3.77)

set_table_border(split_table, BORDER_GRAY, 4)

# Provider sub-table
prov_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
p0 = prov_cell.paragraphs[0]
no_space_para(p0)
add_run(p0, get_value("Medical Provider"), bold=True, size=9, color=DARK_GRAY)

prov_rows = [
    ("Clinic", get_value("Clinic")),
    ("Specialty", get_value("Specialty")),
    ("Address", get_value("Address")),
    ("Phone", get_value("Phone")),
]

for lbl, val in prov_rows:
    pt = prov_cell.add_table(1, 2)
    pt.autofit = False
    remove_table_borders(pt)
    lc2 = pt.cell(0, 0)
    lc2.width = Inches(1.0)
    vc2 = pt.cell(0, 1)
    vc2.width = Inches(2.4)
    lp2 = lc2.paragraphs[0]
    no_space_para(lp2)
    add_run(lp2, lbl, bold=True, size=8.5, color=(0, 120, 120))
    vp2 = vc2.paragraphs[0]
    no_space_para(vp2)
    add_run(vp2, val, size=8.5, color=DARK_GRAY)

# Insurance sub-section
ins_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
i0 = ins_cell.paragraphs[0]
no_space_para(i0)
add_run(i0, get_value("Insurance / Claim Administrator"), bold=True, size=9, color=DARK_GRAY)

ins_rows = [
    ("Company / TPA", get_value("Company / TPA")),
    ("Policy Holder", get_value("Policy Holder")),
    ("Member", get_value("Member")),
    ("Claim Type", get_value("Claim Type")),
]

for lbl, val in ins_rows:
    it = ins_cell.add_table(1, 2)
    it.autofit = False
    remove_table_borders(it)
    lc3 = it.cell(0, 0)
    lc3.width = Inches(1.1)
    vc3 = it.cell(0, 1)
    vc3.width = Inches(2.57)
    lp3 = lc3.paragraphs[0]
    no_space_para(lp3)
    add_run(lp3, lbl, bold=True, size=8.5, color=(0, 120, 120))
    vp3 = vc3.paragraphs[0]
    no_space_para(vp3)
    add_run(vp3, val, size=8.5, color=DARK_GRAY)

doc.add_paragraph()

# ─── BAND 4: FINANCIAL SUMMARY ───
fin_header = doc.add_paragraph()
no_space_para(fin_header)
set_cell_bg_para = fin_header
fh_table = doc.add_table(rows=1, cols=1)
fh_table.autofit = False
fh_table.width = Inches(7.27)
fhc = fh_table.cell(0, 0)
set_cell_bg(fhc, TEAL_HEX)
fhp = fhc.paragraphs[0]
no_space_para(fhp)
add_run(fhp, get_value("FINANCIAL SUMMARY"), bold=True, size=10, color=(255, 255, 255))

# Three summary boxes
summary_table = doc.add_table(rows=1, cols=3)
summary_table.autofit = False
summary_table.width = Inches(7.27)
set_table_border(summary_table, BORDER_GRAY, 4)

sum_data = [
    ("GROSS TOTAL", get_value("GROSS TOTAL"), "Clinic-billed services"),
    ("CONTRACTUAL DISCOUNTS", get_value("CONTRACTUAL DISCOUNTS"), "Applied on consultation and injections"),
    ("NET AMOUNT REQUESTED", get_value("NET AMOUNT REQUESTED"), "After discount and deductible/copay"),
]

col_w = [2.42, 2.42, 2.43]
for i, (lbl, val, sub) in enumerate(sum_data):
    sc = summary_table.cell(0, i)
    sc.width = Inches(col_w[i])
    set_cell_bg(sc, "F5FAFA")
    sp1 = sc.paragraphs[0]
    no_space_para(sp1)
    add_run(sp1, lbl, bold=True, size=7.5, color=GRAY)
    sp2 = sc.add_paragraph()
    no_space_para(sp2)
    add_run(sp2, val, bold=True, size=16, color=DARK_GRAY)
    sp3 = sc.add_paragraph()
    no_space_para(sp3)
    add_run(sp3, sub, size=7.5, color=GRAY)

# Services table
svc_table = doc.add_table(rows=6, cols=6)
svc_table.autofit = False
svc_table.width = Inches(7.27)
set_table_border(svc_table, BORDER_GRAY, 4)

svc_col_w = [2.5, 0.5, 1.0, 1.0, 1.0, 1.27]
headers = ["Service / Treatment", "Qty.", "Gross Tariff", "Discount", "Eligible Value", "Notes"]

for j, (h, w) in enumerate(zip(headers, svc_col_w)):
    hc = svc_table.cell(0, j)
    hc.width = Inches(w)
    set_cell_bg(hc, "E8F4F4")
    hp = hc.paragraphs[0]
    no_space_para(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(hp, h, bold=True, size=8.5, color=DARK_GRAY)

svc_data = [
    ("ORL Consultation - Specialist Doctor", "1", "15.00 BD", "3.00 BD", "12.00 BD", "Approved"),
    ("Lincocin Injection", "1", "10.00 BD", "2.00 BD", "8.00 BD", "Medically indicated"),
    ("Depo-Medrol Injection", "1", "10.00 BD", "2.00 BD", "8.00 BD", "Medically indicated"),
    ("Total", "", "35.00 BD", "7.00 BD", "28.00 BD", "Before\ndeduction"),
]

for r_idx, row_data in enumerate(svc_data):
    row_num = r_idx + 1
    for j, (val, w) in enumerate(zip(row_data, svc_col_w)):
        c = svc_table.cell(row_num, j)
        c.width = Inches(w)
        if row_num == 4:
            set_cell_bg(c, "F0F0F0")
        p = c.paragraphs[0]
        no_space_para(p)
        bold_row = (row_num == 4)
        add_run(p, val, bold=bold_row, size=8.5, color=DARK_GRAY)

# Deduction row
ded_row = svc_table.rows[5]
# Merge cells 0-4 for label
merged = svc_table.cell(5, 0)
for j in range(1, 4):
    merged = merged.merge(svc_table.cell(5, j))
ded_lbl_p = merged.paragraphs[0]
no_space_para(ded_lbl_p)
ded_lbl_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(ded_lbl_p, "Member Deductible / Copay", bold=True, size=8.5, color=DARK_GRAY)

ded_val_c = svc_table.cell(5, 4)
ded_val_p = ded_val_c.paragraphs[0]
no_space_para(ded_val_p)
add_run(ded_val_p, get_value("Member Deductible / Copay"), bold=True, size=8.5, color=DARK_GRAY)

# Net row
net_table = doc.add_table(rows=1, cols=2)
net_table.autofit = False
net_table.width = Inches(7.27)
set_table_border(net_table, BORDER_GRAY, 4)

net_lbl_c = net_table.cell(0, 0)
net_lbl_c.width = Inches(5.77)
net_val_c = net_table.cell(0, 1)
net_val_c.width = Inches(1.5)

net_lbl_p = net_lbl_c.paragraphs[0]
no_space_para(net_lbl_p)
net_lbl_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(net_lbl_p, "Net Amount to Reimburse / Collect", bold=True, size=9, color=DARK_GRAY)

net_val_p = net_val_c.paragraphs[0]
no_space_para(net_val_p)
set_cell_bg(net_val_c, "E8F4F4")
add_run(net_val_p, get_value("Net Amount to Reimburse / Collect"), bold=True, size=11, color=(0, 100, 100))

doc.add_paragraph()

# ─── BAND 5: CONFIRMATION SIGNATURES ───
conf_table = doc.add_table(rows=1, cols=2)
conf_table.autofit = False
conf_table.width = Inches(7.27)
set_table_border(conf_table, BORDER_GRAY, 4)

conf_data = [
    (
        "Medical Provider Confirmation",
        "I confirm that the above services were provided to the patient on the stated date.",
        "Name, signature and stamp"
    ),
    (
        "Member / Patient Confirmation",
        "I confirm receipt of services and the accuracy of the declared information.",
        "Name and member signature"
    ),
]

for i, (title, body, sig) in enumerate(conf_data):
    cc = conf_table.cell(0, i)
    cc.width = Inches(3.6) if i == 0 else Inches(3.67)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cp1 = cc.paragraphs[0]
    no_space_para(cp1)
    add_run(cp1, title, bold=True, size=9, color=DARK_GRAY)
    cp2 = cc.add_paragraph()
    set_para_spacing(cp2, 0, 60)
    add_run(cp2, body, size=8, color=GRAY)
    cp3 = cc.add_paragraph()
    set_para_spacing(cp3, 60, 0)
    add_run(cp3, "_" * 45, size=8, color=GRAY)
    cp4 = cc.add_paragraph()
    no_space_para(cp4)
    add_run(cp4, sig, size=8, color=GRAY)

doc.add_paragraph()

# ─── FOOTER ───
footer_para = doc.add_paragraph()
no_space_para(footer_para)
footer_table = doc.add_table(rows=1, cols=2)
footer_table.autofit = False
footer_table.width = Inches(7.27)
remove_table_borders(footer_table)
fl = footer_table.cell(0, 0)
fr = footer_table.cell(0, 1)
fl.width = Inches(3.6)
fr.width = Inches(3.67)
flp = fl.paragraphs[0]
no_space_para(flp)
add_run(flp, "Medical claim form - EN | pg. 1/3", size=7.5, color=GRAY)
frp = fr.paragraphs[0]
no_space_para(frp)
frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(frp, "Model document with anonymised data", size=7.5, color=GRAY)

doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")