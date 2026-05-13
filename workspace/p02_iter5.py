import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_02.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p02.json"

with open(DATA_FILE, "r", encoding="utf-8") as f:
    data = json.load(f)

def get_value(key):
    for item in data:
        if item["key"] == key:
            return item["value"]
    return ""

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), val)
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_outer_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

def add_run(para, text, bold=False, size=8, italic=False, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    return run

def cell_para(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.italic = italic
    return para

def merge_row_cells(table, row_idx, start_col, end_col):
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.4)
section.bottom_margin = Inches(0.4)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(8)

# ─── HEADER ───
header_table = doc.add_table(rows=1, cols=3)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_table.autofit = False
total_w = Inches(7.5)
header_table.columns[0].width = Inches(1.5)
header_table.columns[1].width = Inches(4.5)
header_table.columns[2].width = Inches(1.5)

row = header_table.rows[0]
row.cells[0]._tc.get_or_add_tcPr()
set_no_border(row.cells[0])
set_no_border(row.cells[1])
set_no_border(row.cells[2])

p0 = row.cells[0].paragraphs[0]
p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p0, "BIA", bold=True, size=16)

p1 = row.cells[1].paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p1, "NATIONAL CLAIM FORM - PRIMARY MEDICAL CARE\n", bold=True, size=11)
add_run(p1, get_value("FORMULAIRE NATIONAL DE DEMANDE DE REMBOURSEMENT - SOINS MÉDICAUX PRIMAIRES"), bold=False, size=7, italic=True)

p2 = row.cells[2].paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p2, "GlobeMed\n", bold=True, size=10)
add_run(p2, f"Sr. No.: {get_value('N° de série')}", bold=False, size=8)

doc.add_paragraph()

# ─── MAIN FORM TABLE ───
main = doc.add_table(rows=0, cols=2)
main.alignment = WD_TABLE_ALIGNMENT.CENTER
main.autofit = False
main.columns[0].width = Inches(3.9)
main.columns[1].width = Inches(3.6)
set_table_outer_border(main)

def add_section_header(table, text):
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    set_cell_bg(cell, "D3D3D3")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text, bold=True, size=8)

def add_two_col_row(table, left_text, right_text, height_pt=18):
    row = table.add_row()
    row.height = Pt(height_pt)
    lc = row.cells[0]
    rc = row.cells[1]
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(1)
    lp.paragraph_format.space_after = Pt(1)
    add_run(lp, left_text, size=8)
    rp = rc.paragraphs[0]
    rp.paragraph_format.space_before = Pt(1)
    rp.paragraph_format.space_after = Pt(1)
    add_run(rp, right_text, size=8)
    return row

# SECTION A header
add_section_header(main, get_value("SECTION A : INFORMATIONS DU MEMBRE (À REMPLIR PAR LE MEMBRE ASSURÉ)"))

# Member Name | Insurance Company
r1 = main.add_row()
r1.height = Pt(18)
r1.cells[0].paragraphs[0].add_run(f"Nom du membre: {get_value('Nom du membre')}").font.size = Pt(8)
r1.cells[1].paragraphs[0].add_run(f"Nom de la compagnie d'assurance / TPA: {get_value('Nom de la compagnie d\u2019assurance / TPA')}").font.size = Pt(8)

# Member ID | Policy Holder
r2 = main.add_row()
r2.height = Pt(18)
r2.cells[0].paragraphs[0].add_run(f"Membre (ID/Numéro): {get_value('Membre (ID/Numéro)')}").font.size = Pt(8)
r2.cells[1].paragraphs[0].add_run(f"Titulaire de la police: {get_value('Titulaire de la police')}").font.size = Pt(8)

# DOB + Gender | CPR/Passport
r3 = main.add_row()
r3.height = Pt(18)
gender_m = get_value("Sexe : M")
gender_f = get_value("Sexe : F")
r3.cells[0].paragraphs[0].add_run(
    f"Date de naissance: {get_value('Date de naissance')}   Sexe: M {gender_m}  F {gender_f}"
).font.size = Pt(8)
r3.cells[1].paragraphs[0].add_run(f"CPR/Passeport: {get_value('CPR/Passeport')}").font.size = Pt(8)

# Single/Married | Phone
r4 = main.add_row()
r4.height = Pt(16)
single = get_value("Célibataire")
married = get_value("Marié(e)")
r4.cells[0].paragraphs[0].add_run(
    f"Célibataire: {single}   Marié(e): {married}"
).font.size = Pt(8)
r4.cells[1].paragraphs[0].add_run(f"Numéro de téléphone du membre: {get_value('Numéro de téléphone du membre')}").font.size = Pt(8)

# SECTION B header
add_section_header(main, get_value("SECTION B : SECTION MÉDICALE (À REMPLIR UNIQUEMENT PAR LE MÉDECIN TRAITANT)"))

# Tick boxes | Provider Name
r5 = main.add_row()
r5.height = Pt(16)
inp = get_value("Veuillez cocher : Hospitalisation")
outp = get_value("Veuillez cocher : Consultation externe")
emer = get_value("Veuillez cocher : Cas d'urgence")
r5.cells[0].paragraphs[0].add_run(
    f"Cocher: Hosp. {inp}  Cons. ext. {outp}  Urgence {emer}"
).font.size = Pt(8)
r5.cells[1].paragraphs[0].add_run(f"Nom du prestataire: {get_value('Nom du prestataire')}").font.size = Pt(8)

# Date of Treatment | Medical Record No
r6 = main.add_row()
r6.height = Pt(16)
r6.cells[0].paragraphs[0].add_run(f"Date du traitement: {get_value('Date du traitement')}").font.size = Pt(8)
r6.cells[1].paragraphs[0].add_run(f"Numéro de dossier médical: {get_value('Numéro de dossier médical')}").font.size = Pt(8)

# Pre Existing | Vital Signs
r7 = main.add_row()
r7.height = Pt(16)
rta = get_value("Accident de la route")
r7.cells[0].paragraphs[0].add_run(f"Condition préexistante: {get_value('Condition préexistante')}  RTA: {rta}").font.size = Pt(8)
r7.cells[1].paragraphs[0].add_run(f"Signes vitaux: {get_value('Signes vitaux')}").font.size = Pt(8)

# Chronic | Blood Pressure
r8 = main.add_row()
r8.height = Pt(16)
chron = get_value("Maladie chronique")
work = get_value("Accident du travail")
r8.cells[0].paragraphs[0].add_run(f"Maladie chronique: {chron}  Accident du travail: {work}").font.size = Pt(8)
r8.cells[1].paragraphs[0].add_run(f"Pression artérielle: {get_value('Pression artérielle')}").font.size = Pt(8)

# Maternity | Pulse
r9 = main.add_row()
r9.height = Pt(16)
mat = get_value("Maternité")
edd = get_value("Date prévue d'accouchement")
r9.cells[0].paragraphs[0].add_run(f"Maternité: {mat}  DPA: {edd}").font.size = Pt(8)
r9.cells[1].paragraphs[0].add_run(f"Pouls: {get_value('Pouls')}").font.size = Pt(8)

# Others | Temp
r10 = main.add_row()
r10.height = Pt(16)
r10.cells[0].paragraphs[0].add_run(f"Autres: {get_value('Autres (veuillez préciser)')}").font.size = Pt(8)
r10.cells[1].paragraphs[0].add_run(f"Température: {get_value('Température')}").font.size = Pt(8)

# Main Complaint | Duration
r11 = main.add_row()
r11.height = Pt(16)
r11.cells[0].paragraphs[0].add_run(f"Plainte principale: {get_value('Plainte principale et symptômes présentés')}").font.size = Pt(8)
r11.cells[1].paragraphs[0].add_run(f"Durée/Historique de la maladie: {get_value('Durée/Historique de la maladie')}").font.size = Pt(8)

# Main Complaint large area
r11b = main.add_row()
r11b.height = Pt(50)
r11b.cells[0].merge(r11b.cells[1])
p_mc = r11b.cells[0].paragraphs[0]
p_mc.add_run(get_value("Plainte principale et symptômes présentés")).font.size = Pt(10)

# Clinical Findings
r12 = main.add_row()
r12.height = Pt(16)
r12.cells[0].merge(r12.cells[1])
p12 = r12.cells[0].paragraphs[0]
add_run(p12, "Constatations cliniques et diagnostic final (utiliser les codes CIM si approprié):", bold=True, size=8)

r12b = main.add_row()
r12b.height = Pt(50)
r12b.cells[0].paragraphs[0].add_run(get_value("Constatations cliniques et diagnostic final (utiliser les codes CIM si approprié)")).font.size = Pt(10)
r12b.cells[1].paragraphs[0].add_run(get_value("Médicaments / Traitement prescrit")).font.size = Pt(8)

# PRE-AUTHORIZATION SECTION header
add_section_header(main, "SECTION DE PRÉ-AUTORISATION (LE RAPPORT MÉDICAL ET D'INVESTIGATION DOIT ÊTRE JOINT SI APPLICABLE)")

# Anticipated Length of Stay | Anticipated Cost
r13 = main.add_row()
r13.height = Pt(18)
r13.cells[0].paragraphs[0].add_run(f"Durée de séjour prévue: {get_value('Durée de séjour prévue')}").font.size = Pt(8)
r13.cells[1].paragraphs[0].add_run(f"COÛT PRÉVU: {get_value('COÛT PRÉVU')}").font.size = Pt(8)

# Signature area
r14 = main.add_row()
r14.height = Pt(40)
r14.cells[0].paragraphs[0].add_run(get_value("Signature du membre assuré")).font.size = Pt(8)
# Medical Service Provider Declaration
decl_text = get_value("Déclaration du prestataire de services médicaux")
r14.cells[1].paragraphs[0].add_run(f"Déclaration du prestataire:\n{decl_text}").font.size = Pt(7)

# Package Deal Code
r15 = main.add_row()
r15.height = Pt(18)
r15.cells[0].paragraphs[0].add_run(f"Code forfait: {get_value('Code forfait')}").font.size = Pt(8)
r15.cells[1].paragraphs[0].add_run("").font.size = Pt(8)

# Name of Doctor / Signature / Stamp
r16 = main.add_row()
r16.height = Pt(18)
r16.cells[0].paragraphs[0].add_run(f"Nom du médecin: {get_value('Nom du médecin')}").font.size = Pt(8)
r16.cells[1].paragraphs[0].add_run(f"Signature: {get_value('Signature du médecin')}").font.size = Pt(8)

r17 = main.add_row()
r17.height = Pt(18)
r17.cells[0].paragraphs[0].add_run(f"Cachet et numéro de licence: {get_value('Cachet et numéro de licence')}").font.size = Pt(8)
r17.cells[1].paragraphs[0].add_run(f"Date: {get_value('Date (médecin)')}").font.size = Pt(8)

# Member Signature
r18 = main.add_row()
r18.height = Pt(18)
r18.cells[0].paragraphs[0].add_run(f"Signature du membre: {get_value('Signature du membre')}").font.size = Pt(8)
r18.cells[1].paragraphs[0].add_run(f"Date: {get_value('Date (membre)')}").font.size = Pt(8)

# FOR INSURANCE COMPANY USE ONLY
add_section_header(main, "RÉSERVÉ À LA COMPAGNIE D'ASSURANCE UNIQUEMENT")

r19 = main.add_row()
r19.height = Pt(16)
appr = get_value("Approuvé")
not_appr = get_value("Non approuvé")
r19.cells[0].paragraphs[0].add_run(f"Approuvé: {appr}   Non approuvé: {not_appr}").font.size = Pt(8)
r19.cells[1].paragraphs[0].add_run(f"Commentaires: {get_value('Commentaires')}").font.size = Pt(8)

r20 = main.add_row()
r20.height = Pt(16)
r20.cells[0].paragraphs[0].add_run(f"N° d'approbation: {get_value('N° d\u2019approbation')}   Validité: {get_value('Validité de l\u2019approbation')}").font.size = Pt(8)
r20.cells[1].paragraphs[0].add_run(f"N° de réclamation: {get_value('N° de réclamation')}").font.size = Pt(8)

r21 = main.add_row()
r21.height = Pt(16)
r21.cells[0].paragraphs[0].add_run(f"Agent d'assurance: {get_value('Agent d\u2019assurance')}   Signature: {get_value('Signature de l\u2019agent')}").font.size = Pt(8)
r21.cells[1].paragraphs[0].add_run(f"Date: {get_value('Date (assurance)')}").font.size = Pt(8)

# Footer
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.paragraph_format.space_before = Pt(2)
add_run(footer_para,
    "Hotline et contact de pré-autorisation: +973 17530886  Fax: +973 17530602\n"
    "Pour les approbations ambulatoires: op.approvals@globemedbahrain.com\n"
    "Pour les approbations hospitalisées: Ip.approvals@globemedbahrain.com\n"
    "Pour le remboursement des réclamations: cr@globemedbahrain.com",
    size=6)

doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")