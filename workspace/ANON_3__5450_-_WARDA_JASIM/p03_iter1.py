import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/data_p03.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/page_03.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()

s = doc.sections[0]
s.top_margin = Cm(2)
s.bottom_margin = Cm(2)
s.left_margin = Cm(2)
s.right_margin = Cm(2)

def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        pgBorders.append(border)
    sectPr.append(pgBorders)

add_page_border(s)

for _ in range(10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

h = doc.add_paragraph()
h.paragraph_format.space_before = Pt(0)
h.paragraph_format.space_after = Pt(6)
h.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = h.add_run(lookup["PAGE D'ID EXCLUE"])
run.bold = True
run.font.size = Pt(22)

p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
p1.paragraph_format.space_before = Pt(6)
p1.paragraph_format.space_after = Pt(0)
r1 = p1.add_run(lookup["Nom du patient correspondant aux autres pages"])
r1.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
r2 = p2.add_run("Nom correspondant: " + lookup["Nom correspondant"])
r2.font.size = Pt(10)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(0)
r3 = p3.add_run("Type de document: " + lookup["Type de document"])
r3.font.size = Pt(10)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(0)
r4 = p4.add_run("Numéro de page original: " + lookup["Numéro de page original"])
r4.font.size = Pt(10)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)