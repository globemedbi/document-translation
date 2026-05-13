import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/data_p05.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/page_05.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()

s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def set_font(run, size, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold

def add_par(text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size, bold)
    return p

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        val = kwargs.get(edge, "none")
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_border(t, border="single"):
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), border)
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def cell_par(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_font(r, size, bold)
    return p

def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p)
r = p.add_run(lookup["Nom de la clinique"])
set_font(r, 14, True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p2)
r2 = p2.add_run(lookup["Type de formulaire"])
set_font(r2, 10, False)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p3)
r3 = p3.add_run("Référence#:  " + lookup["Référence#"])
set_font(r3, 10, False)

patient_labels = [
    ("Numéro individuel", "Individual Number"),
    ("Nom du titulaire de la carte", "Card Holder Name"),
    ("Date de naissance", "Date Of Birth"),
    ("Compagnie d'assurance", "Insurance Co."),
    ("Médecin", "Physician"),
    ("Date d'approbation (début)", "Approval Date"),
    ("Date d'approbation (fin)", ""),
]

t1 = doc.add_table(rows=7, cols=2)
t1.width = Cm(17)
set_table_border(t1, "single")
from docx.shared import Cm as C
col_widths = [C(6.5), C(10.5)]
for i, row_data in enumerate([
    ("Individual Number", lookup["Numéro individuel"]),
    ("Card Holder Name", lookup["Nom du titulaire de la carte"]),
    ("Date Of Birth", lookup["Date de naissance"]),
    ("Insurance Co.", lookup["Compagnie d'assurance"]),
    ("Physician", lookup["Médecin"]),
    ("Approval Date", lookup["Date d'approbation (début)"]),
    ("", lookup["Date d'approbation (fin)"]),
]):
    r = t1.rows[i]
    r.cells[0].width = col_widths[0]
    r.cells[1].width = col_widths[1]
    cell_par(r.cells[0], row_data[0], size=9, bold=True)
    cell_par(r.cells[1], row_data[1], size=9, bold=False)

svc1_parts = lookup["Service / Article 1"].split(" / ")
svc1_service = svc1_parts[0] if len(svc1_parts) > 0 else ""
svc1_item = svc1_parts[1] if len(svc1_parts) > 1 else ""

t2 = doc.add_table(rows=3, cols=7)
t2.width = Cm(17)
set_table_border(t2, "single")
hdr_widths2 = [C(3.5), C(3.5), C(1.5), C(1.5), C(2.5), C(2.5), C(2.0)]
hdr_labels = ["Service", "Item", "Qty Clm", "Qty App", "Status", "Notes", "Reason"]
for j, lbl in enumerate(hdr_labels):
    t2.rows[0].cells[j].width = hdr_widths2[j]
    cell_par(t2.rows[0].cells[j], lbl, size=9, bold=True)

t2.rows[1].cells[0].width = hdr_widths2[0]
t2.rows[1].cells[1].width = hdr_widths2[1]
cell_par(t2.rows[1].cells[0], svc1_service, size=9, bold=False)
cell_par(t2.rows[1].cells[1], svc1_item, size=9, bold=False)
cell_par(t2.rows[1].cells[2], lookup["Qté réclamée (Consultation)"], size=9)
cell_par(t2.rows[1].cells[3], lookup["Qté approuvée (Consultation)"], size=9)
cell_par(t2.rows[1].cells[4], lookup["Statut (Consultation)"], size=10, bold=True)
cell_par(t2.rows[1].cells[5], lookup["Notes (Consultation)"], size=9)
cell_par(t2.rows[1].cells[6], lookup["Motif (Consultation)"], size=9)

for j in range(7):
    t2.rows[2].cells[j].width = hdr_widths2[j]
cell_par(t2.rows[2].cells[0], "", size=9)

svc2_parts = lookup["Service / Article 2"].split(" / ")
svc2_service = svc2_parts[0] if len(svc2_parts) > 0 else ""
svc2_item = svc2_parts[1] if len(svc2_parts) > 1 else ""

svc3_parts = lookup["Service / Article 3"].split(" / ")
svc3_service = svc3_parts[0] if len(svc3_parts) > 0 else ""
svc3_item = svc3_parts[1] if len(svc3_parts) > 1 else ""

t3 = doc.add_table(rows=4, cols=7)
t3.width = Cm(17)
set_table_border(t3, "single")
hdr_widths3 = [C(3.5), C(3.5), C(1.5), C(1.5), C(2.5), C(2.5), C(2.0)]
for j, lbl in enumerate(hdr_labels):
    t3.rows[0].cells[j].width = hdr_widths3[j]
    cell_par(t3.rows[0].cells[j], lbl, size=9, bold=True)

for j in range(7):
    t3.rows[1].cells[j].width = hdr_widths3[j]
cell_par(t3.rows[1].cells[0], svc2_service, size=9)
cell_par(t3.rows[1].cells[1], svc2_item, size=9)
cell_par(t3.rows[1].cells[2], lookup["Qté réclamée (LINCOCIN)"], size=9)
cell_par(t3.rows[1].cells[3], lookup["Qté approuvée (LINCOCIN)"], size=9)
cell_par(t3.rows[1].cells[4], lookup["Statut (LINCOCIN)"], size=10, bold=True)
cell_par(t3.rows[1].cells[5], lookup["Notes (LINCOCIN)"], size=9)
cell_par(t3.rows[1].cells[6], lookup["Raison (LINCOCIN)"], size=9)

for j in range(7):
    t3.rows[2].cells[j].width = hdr_widths3[j]
cell_par(t3.rows[2].cells[0], svc3_service, size=9)
cell_par(t3.rows[2].cells[1], svc3_item, size=9)
cell_par(t3.rows[2].cells[2], lookup["Qté Réclamée (Depomedrol)"], size=9)
cell_par(t3.rows[2].cells[3], lookup["Qté Approuvée (Depomedrol)"], size=9)
cell_par(t3.rows[2].cells[4], lookup["Statut (Depomedrol)"], size=10, bold=True)
cell_par(t3.rows[2].cells[5], lookup["Notes (Depomedrol)"], size=9)
cell_par(t3.rows[2].cells[6], lookup["Raison (Depomedrol)"], size=9)

for j in range(7):
    t3.rows[3].cells[j].width = hdr_widths3[j]
t3.rows[3].cells[0]._tc.merge(t3.rows[3].cells[6]._tc)
cell_par(t3.rows[3].cells[0], lookup["Détails de la couverture"], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

t4 = doc.add_table(rows=1, cols=1)
t4.width = Cm(17)
set_table_border(t4, "single")
t4.rows[0].cells[0].width = Cm(17)
cell_par(t4.rows[0].cells[0], lookup["Notes importantes"], size=9, bold=True)

t5 = doc.add_table(rows=1, cols=1)
t5.width = Cm(17)
set_table_border(t5, "single")
t5.rows[0].cells[0].width = Cm(17)
notes_content = lookup["Contenu des notes importantes"]
c = t5.rows[0].cells[0]
c.paragraphs[0].clear()
for line in notes_content.split("\n"):
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    rr = p.add_run(line)
    set_font(rr, 9, False)
    if line.startswith("Veuillez"):
        rr.font.italic = True

t6 = doc.add_table(rows=3, cols=1)
t6.width = Cm(17)
set_table_border(t6, "single")
cell_par(t6.rows[0].cells[0], lookup["Diagnostic"], size=9, bold=True)
cell_par(t6.rows[1].cells[0], lookup["Diagnostic 1"], size=9)
cell_par(t6.rows[2].cells[0], lookup["Diagnostic 2"], size=9)

t7 = doc.add_table(rows=3, cols=8)
t7.width = Cm(17)
set_table_border(t7, "single")
t7.rows[0].cells[0]._tc.merge(t7.rows[0].cells[7]._tc)
cell_par(t7.rows[0].cells[0], lookup["Informations médicales"], size=9, bold=True)

med_headers = ["Date", "TA", "Pouls", "Température", "Durée de la maladie", "Plaintes", "Signes", "Autres"]
med_keys = ["Date", "TA", "Pouls", "Température", "Durée de la maladie", "Plaintes", "Signes", "Autres"]
col_w = Cm(17/8)
for j, lbl in enumerate(med_headers):
    t7.rows[1].cells[j].width = col_w
    cell_par(t7.rows[1].cells[j], lbl, size=9, bold=True)
for j, k in enumerate(med_keys):
    t7.rows[2].cells[j].width = col_w
    cell_par(t7.rows[2].cells[j], lookup[k], size=9)

t8 = doc.add_table(rows=2, cols=1)
t8.width = Cm(17)
set_table_border(t8, "single")
cell_par(t8.rows[0].cells[0], "Fichiers téléchargés", size=9, bold=True)
cell_par(t8.rows[1].cells[0], lookup["Fichiers téléchargés"], size=9)

t9 = doc.add_table(rows=2, cols=1)
t9.width = Cm(17)
set_table_border(t9, "single")
cell_par(t9.rows[0].cells[0], "Informations supplémentaires", size=9, bold=True)
cell_par(t9.rows[1].cells[0], lookup["Informations supplémentaires"], size=9)

p_code = doc.add_paragraph()
no_space(p_code)
r_code = p_code.add_run(lookup["Code"])
set_font(r_code, 9, False)

t10 = doc.add_table(rows=1, cols=1)
t10.width = Cm(17)
set_table_border(t10, "single")
cell_par(t10.rows[0].cells[0], lookup["Avertissement"], size=9)

p_vat = doc.add_paragraph()
no_space(p_vat)
r_vat = p_vat.add_run(lookup["Avis de TVA"])
set_font(r_vat, 10, True)
p_vat.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)