import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/data_p04.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/page_04.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()

s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.5)

def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        pgBorders.append(border)
    sectPr.append(pgBorders)

add_page_border(s)

def add_spacer(d, lines=1):
    for _ in range(lines):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

for _ in range(10):
    add_spacer(doc)

h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.LEFT
h.paragraph_format.space_before = Pt(0)
h.paragraph_format.space_after = Pt(12)
run = h.add_run(lookup["PAGE D'IDENTITÉ EXCLUE"])
run.font.size = Pt(22)
run.font.bold = False
run.font.name = "Calibri"

add_spacer(doc)

text_fields = [
    "Nom du patient correspondant aux autres pages",
    ("Nom correspondant", "Nom correspondant"),
    ("Type de document", "Type de document"),
    ("Numéro de page original", "Numéro de page original"),
]

p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
p1.paragraph_format.space_before = Pt(0)
p1.paragraph_format.space_after = Pt(0)
r1 = p1.add_run(lookup["Nom du patient correspondant aux autres pages"])
r1.font.size = Pt(10)
r1.font.name = "Calibri"

for key, label in [
    ("Nom correspondant", "Nom correspondant"),
    ("Type de document", "Type de document"),
    ("Numéro de page original", "Numéro de page original"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: {lookup[key]}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)