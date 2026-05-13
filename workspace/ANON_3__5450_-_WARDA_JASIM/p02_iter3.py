import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/data_p02.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/page_02.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()

s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.2)

def set_cell_bg(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, active in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if active:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def para(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    if len(cell.paragraphs) > 0 and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def section_header(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.width = Cm(18)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "BFBFBF")
    para(c, text, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

def add_table_bordered(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.width = Cm(18)
    from docx.oxml import OxmlElement
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)
    return t

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

hp = doc.add_paragraph()
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = hp.add_run("BIA                    ")
r1.font.size = Pt(20)
r1.font.bold = True
r2 = hp.add_run("GlobeMed")
r2.font.size = Pt(14)
r2.font.bold = True

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run(lookup["FORMULAIRE NATIONAL DE DEMANDE DE REMBOURSEMENT - SOINS MÉDICAUX PRIMAIRES"])
tr.font.size = Pt(12)
tr.font.bold = True

sr_p = doc.add_paragraph()
sr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sr_r = sr_p.add_run("N° de série: " + lookup["N° de série"])
sr_r.font.size = Pt(9)

section_header(doc, lookup["SECTION A : INFORMATIONS SUR LE MEMBRE (À REMPLIR PAR LE MEMBRE ASSURÉ)"])

t = add_table_bordered(doc, 4, 2)
t.width = Cm(18)

r0 = t.rows[0]
r0.cells[0].width = Cm(9.9)
r0.cells[1].width = Cm(8.1)
para(r0.cells[0], "Nom du membre: " + lookup["Nom du membre"], size=9)
ins_text = "Nom de la compagnie d'assurance/TPA: " + lookup["Nom de la compagnie d'assurance/TPA"]
para(r0.cells[1], ins_text, size=9)

r1 = t.rows[1]
para(r1.cells[0], "Membre (N° ID/Carte): " + lookup["Membre (N° ID/Carte)"], size=9)
para(r1.cells[1], "Titulaire de la police: " + lookup["Titulaire de la police"], size=9)

r2 = t.rows[2]
dob = lookup["Date de naissance"]
gm = lookup["Sexe : M"]
gf = lookup["Sexe : F"]
para(r2.cells[0], "Date de naissance: " + dob + "    Sexe : M " + gm + "  F " + gf, size=9)
para(r2.cells[1], "CPR/Passeport: " + lookup["CPR/Passeport"], size=9)

r3 = t.rows[3]
single = lookup["Célibataire"]
married = lookup["Marié(e)"]
para(r3.cells[0], "Célibataire " + single + "    Marié(e) " + married, size=9)
para(r3.cells[1], "Numéro de téléphone du membre: " + lookup["Numéro de téléphone du membre"], size=9)

section_header(doc, lookup["SECTION B : SECTION MÉDICALE (À REMPLIR UNIQUEMENT PAR LE MÉDECIN TRAITANT)"])

t2 = add_table_bordered(doc, 2, 2)
t2.width = Cm(18)

r0 = t2.rows[0]
inp = lookup["Veuillez cocher : Patient hospitalisé"]
outp = lookup["Veuillez cocher : Patient ambulatoire"]
emer = lookup["Veuillez cocher : Cas d'urgence"]
tick_text = "Veuillez cocher : Patient hospitalisé " + inp + "  Patient ambulatoire " + outp + "  Cas d'urgence " + emer
para(r0.cells[0], tick_text, size=9)
para(r0.cells[1], "Nom du prestataire: " + lookup["Nom du prestataire"], size=9)

r1 = t2.rows[1]
para(r1.cells[0], "Date du traitement: " + lookup["Date du traitement"], size=9)
para(r1.cells[1], "Numéro de dossier médical: " + lookup["Numéro de dossier médical"], size=9)

t3 = add_table_bordered(doc, 5, 2)
t3.width = Cm(18)

r0 = t3.rows[0]
pre = lookup["Condition préexistante : Accident de la route"]
para(r0.cells[0], "Condition préexistante : Accident de la route " + pre, size=9)
para(r0.cells[1], "Signes vitaux: " + lookup["Signes vitaux"], size=9)

r1 = t3.rows[1]
chron = lookup["Condition chronique : Accident du travail"]
para(r1.cells[0], "Condition chronique : Accident du travail " + chron, size=9)
para(r1.cells[1], "Tension artérielle: " + lookup["Tension artérielle"], size=9)

r2 = t3.rows[2]
para(r2.cells[0], "Maternité : Date prévue d'accouchement: " + lookup["Maternité : Date prévue d'accouchement"], size=9)
para(r2.cells[1], "Pouls: " + lookup["Pouls"], size=9)

r3 = t3.rows[3]
para(r3.cells[0], "Autres (veuillez préciser): " + lookup["Autres (veuillez préciser)"], size=9)
para(r3.cells[1], "Température: " + lookup["Température"], size=9)

r4 = t3.rows[4]
para(r4.cells[0], " ", size=9)
para(r4.cells[1], "Durée / Antécédents de la maladie: " + lookup["Durée / Antécédents de la maladie"], size=9)

t4 = add_table_bordered(doc, 2, 1)
t4.width = Cm(18)
para(t4.rows[0].cells[0], "Plainte principale et symptômes présentés: " + lookup["Plainte principale et symptômes présentés"], size=9)
para(t4.rows[1].cells[0], "Constatations cliniques et diagnostic final (utiliser les codes CIM le cas échéant): " + lookup["Constatations cliniques et diagnostic final (utiliser les codes CIM le cas échéant)"], size=9)

section_header(doc, lookup["SECTION DE PRÉ-AUTORISATION"])

t5 = add_table_bordered(doc, 3, 2)
t5.width = Cm(18)

r0 = t5.rows[0]
para(r0.cells[0], "Autorisation / Prescription: " + lookup["Autorisation / Prescription"], size=9)
para(r0.cells[1], "Durée de séjour prévue: " + lookup["Durée de séjour prévue"], size=9)

r1 = t5.rows[1]
para(r1.cells[0], "Code de forfait: " + lookup["Code de forfait"], size=9)
para(r1.cells[1], "COÛT ANTICIPÉ: " + lookup["COÛT ANTICIPÉ"], size=9)

r2 = t5.rows[2]
r2.cells[0].merge(r2.cells[1])
decl_text = (
    "Déclaration du prestataire de services médicaux\n"
    "Nous certifions que TOUTES les informations mentionnées ici sont correctes et que les services médicaux "
    "présentés sur ce formulaire étaient médicalement indiqués et nécessaires pour la gestion de ce cas."
)
para(r2.cells[0], decl_text, size=8)

t6 = add_table_bordered(doc, 3, 2)
t6.width = Cm(18)

r0 = t6.rows[0]
para(r0.cells[0], "Signature (Membre): " + lookup["Signature (Membre)"], size=9)
para(r0.cells[1], "Nom du médecin: " + lookup["Nom du médecin"], size=9)

r1 = t6.rows[1]
para(r1.cells[0], "Date (Membre): " + lookup["Date (Membre)"], size=9)
para(r1.cells[1], "Signature (Médecin): " + lookup["Signature (Médecin)"], size=9)

r2 = t6.rows[2]
para(r2.cells[0], " ", size=9)
para(r2.cells[1], "N° L. / Cachet & N° de licence: " + lookup["N° L. / Cachet & N° de licence"] + "    Date (Médecin): " + lookup["Date (Médecin)"], size=9)

section_header(doc, lookup["RÉSERVÉ À L'USAGE DE LA COMPAGNIE D'ASSURANCE"])

t7 = add_table_bordered(doc, 3, 2)
t7.width = Cm(18)

r0 = t7.rows[0]
r0.cells[0].merge(r0.cells[1])
para(r0.cells[0], "Commentaires: " + lookup["Commentaires"], size=9)

r1 = t7.rows[1]
appr = lookup["Approuvé"]
not_appr = lookup["Non approuvé"]
para(r1.cells[0], "Approuvé " + appr + "    Non approuvé " + not_appr, size=9)
para(r1.cells[1], "N° de SINISTRE: " + lookup["N° de SINISTRE"], size=9)

r2 = t7.rows[2]
para(r2.cells[0], "N° d'approbation: " + lookup["N° d'approbation"] + "    Validité de l'approbation: " + lookup["Validité de l'approbation"], size=9)
para(r2.cells[1], "Date (Assurance): " + lookup["Date (Assurance)"], size=9)

t8 = add_table_bordered(doc, 1, 2)
t8.width = Cm(18)
r0 = t8.rows[0]
para(r0.cells[0], "Agent d'assurance: " + lookup["Agent d'assurance"], size=9)
para(r0.cells[1], "Signature (Agent d'assurance): " + lookup["Signature (Agent d'assurance)"], size=9)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    "Hotline et contact pré-autorisation: +973 17530886  Fax: +973 17530602\n"
    "Pour les approbations ambulatoires: op.approvals@globemedbahrain.com\n"
    "Pour les remboursements: cr@globemedbahrain.com"
)
fr.font.size = Pt(7)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)