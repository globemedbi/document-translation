import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/data_p01.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/page_01.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(15)
s.top_margin = Cm(1)
s.bottom_margin = Cm(1)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)

def set_rtl(par):
    par.paragraph_format.bidi = True
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_font(run, size, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_par(container, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, rtl=False):
    if hasattr(container, "add_paragraph"):
        p = container.add_paragraph()
    else:
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.alignment = align
    if rtl:
        p.paragraph_format.bidi = True
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p

def cell_par(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, rtl=False):
    p = cell.paragraphs[0]
    p.alignment = align
    if rtl:
        p.paragraph_format.bidi = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p

def add_cell_line(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, rtl=False):
    p = cell.add_paragraph()
    p.alignment = align
    if rtl:
        p.paragraph_format.bidi = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width.cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def no_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def redacted_box(cell, width_cm=5):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "000000")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("  ")
    r.font.size = Pt(16)

t0 = doc.add_table(rows=2, cols=2)
t0.width = Cm(18)
no_border(t0)

r0 = t0.rows[0]
r0.cells[0].width = Cm(9)
r0.cells[1].width = Cm(9)

redacted_box(r0.cells[0], 5)
r0.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

redacted_box(r0.cells[1], 5)
r0.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

r1 = t0.rows[1]
lc = r1.cells[0]
rc = r1.cells[1]

cell_par(lc, f"Mob.: {lookup['Tél. mobile']}", size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
add_cell_line(lc, f"Tel: {lookup['Tél.']}", size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
add_cell_line(lc, lookup["Adresse"], size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

cell_par(rc, lookup["Spécialité ORL / للأنف والأذن والحنجرة"], size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
add_cell_line(rc, f"نقال: {lookup['Mobile']}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
add_cell_line(rc, f"هاتف: {lookup['Téléphone']}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
add_cell_line(rc, lookup["الرفاع الغربي، مملكة البحرين"], size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)

doc.add_paragraph()

t1 = doc.add_table(rows=1, cols=2)
t1.width = Cm(18)
no_border(t1)

lc1 = t1.rows[0].cells[0]
rc1 = t1.rows[0].cells[1]

p_no = lc1.paragraphs[0]
p_no.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_no.paragraph_format.space_before = Pt(0)
p_no.paragraph_format.space_after = Pt(0)
rn1 = p_no.add_run("No.:  ")
set_font(rn1, 12, bold=True)
rn2 = p_no.add_run(lookup["N°"])
set_font(rn2, 18, bold=True, color=(200, 0, 0))

p_date = rc1.paragraphs[0]
p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_date.paragraph_format.space_before = Pt(0)
p_date.paragraph_format.space_after = Pt(0)
rd1 = p_date.add_run(f"Date / التاريخ:  {lookup['Date / التاريخ']}")
set_font(rd1, 11, bold=False)

doc.add_paragraph()

t2 = doc.add_table(rows=1, cols=2)
t2.width = Cm(18)
no_border(t2)

lc2 = t2.rows[0].cells[0]
rc2 = t2.rows[0].cells[1]

p_rec = lc2.paragraphs[0]
p_rec.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_rec.paragraph_format.space_before = Pt(0)
p_rec.paragraph_format.space_after = Pt(0)
rr1 = p_rec.add_run(f"Reçu de:  {lookup['Reçu de / إستلمنا من']}")
set_font(rr1, 10)

p_rec_ar = rc2.paragraphs[0]
p_rec_ar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_rec_ar.paragraph_format.bidi = True
p_rec_ar.paragraph_format.space_before = Pt(0)
p_rec_ar.paragraph_format.space_after = Pt(0)
rr2 = p_rec_ar.add_run("إستلمنا من")
set_font(rr2, 10)

doc.add_paragraph()

t3 = doc.add_table(rows=1, cols=2)
t3.width = Cm(18)
no_border(t3)

lc3 = t3.rows[0].cells[0]
rc3 = t3.rows[0].cells[1]

p_sum = lc3.paragraphs[0]
p_sum.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sum.paragraph_format.space_before = Pt(0)
p_sum.paragraph_format.space_after = Pt(0)
rs1 = p_sum.add_run(f"The sum of BD:  {lookup['La somme de BD / مبلغ وقدره']}")
set_font(rs1, 10)

p_sum_ar = rc3.paragraphs[0]
p_sum_ar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_sum_ar.paragraph_format.bidi = True
p_sum_ar.paragraph_format.space_before = Pt(0)
p_sum_ar.paragraph_format.space_after = Pt(0)
rs2 = p_sum_ar.add_run("مبلغ وقدره")
set_font(rs2, 10)

doc.add_paragraph()

t4 = doc.add_table(rows=1, cols=2)
t4.width = Cm(18)
no_border(t4)

lc4 = t4.rows[0].cells[0]
rc4 = t4.rows[0].cells[1]

p_being = lc4.paragraphs[0]
p_being.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_being.paragraph_format.space_before = Pt(0)
p_being.paragraph_format.space_after = Pt(0)
rb1 = p_being.add_run(f"Being of / Au titre de:  {lookup['Au titre de / وذلك عن']}")
set_font(rb1, 9)

p_being_ar = rc4.paragraphs[0]
p_being_ar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_being_ar.paragraph_format.bidi = True
p_being_ar.paragraph_format.space_before = Pt(0)
p_being_ar.paragraph_format.space_after = Pt(0)
rb2 = p_being_ar.add_run("وذلك عن")
set_font(rb2, 9)

doc.add_paragraph()

t5 = doc.add_table(rows=1, cols=2)
t5.width = Cm(18)
no_border(t5)

lc5 = t5.rows[0].cells[0]
rc5 = t5.rows[0].cells[1]

p_sig = lc5.paragraphs[0]
p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sig.paragraph_format.space_before = Pt(0)
p_sig.paragraph_format.space_after = Pt(0)
rsig = p_sig.add_run(f"Signature / التوقيع:  {lookup['Signature / التوقيع']}")
set_font(rsig, 10)

p_net = rc5.paragraphs[0]
p_net.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_net.paragraph_format.space_before = Pt(0)
p_net.paragraph_format.space_after = Pt(0)
rnet1 = p_net.add_run("Net BD / المبلغ:  ")
set_font(rnet1, 10, bold=True)
rnet2 = p_net.add_run(lookup["Net BD / المبلغ"])
set_font(rnet2, 12, bold=True)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)