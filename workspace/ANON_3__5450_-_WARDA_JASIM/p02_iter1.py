import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/data_p02.json"
OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/ANON_3__5450_-_WARDA_JASIM/page_02.docx"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.left_margin = Cm(1.5)
s.right_margin = Cm(1.5)
s.top_margin = Cm(1.2)
s.bottom_margin = Cm(1.2)

def set_font(run, size=9, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold

def shade_cell(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size, bold)
    return p

def add_para(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size, bold)
    return p

def section_header(t, text, cols):
    r = t.add_row()
    c = r.cells[0]
    if cols > 1:
        c.merge(r.cells[cols - 1])
    shade_cell(c)
    para(c, text, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

def dotline(n=30):
    return "." * n

header_t = doc.add_table(rows=1, cols=3)
header_t.width = Cm(18)
header_t.autofit = False
from docx.shared import Inches
widths = [Cm(4), Cm(9), Cm(5)]
for i, w in enumerate(widths):
    header_t.columns[i].width = w

bia = header_t.rows[0].cells[0]
p = bia.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("BIA")
r.font.size = Pt(24)
r.font.bold = True

mid = header_t.rows[0].cells[1]
p = mid.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(lookup["FORMULAIRE NATIONAL DE DEMANDE DE REMBOURSEMENT - SOINS MÉDICAUX PRIMAIRES"])
r.font.size = Pt(11)
r.font.bold = True

gmed = header_t.rows[0].cells[2]
p = gmed.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("GlobeMed")
r.font.size = Pt(12)
r.font.bold = True
p2 = gmed.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run("bahrain's ppa approval")
r2.font.size = Pt(7)
r2.font.italic = True
p3 = gmed.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r3 = p3.add_run(f"N° de série: {lookup['N° de série']}")
r3.font.size = Pt(8)

doc.add_paragraph()

t = doc.add_table(rows=0, cols=2)
t.width = Cm(18)
t.autofit = False
t.columns[0].width = Cm(9.9)
t.columns[1].width = Cm(8.1)

section_header(t, lookup["SECTION A : INFORMATIONS SUR LE MEMBRE (À REMPLIR PAR LE MEMBRE ASSURÉ)"], 2)

r1 = t.add_row()
para(r1.cells[0], f"Nom du membre: {lookup['Nom du membre']}", size=9)
para(r1.cells[1], f"Nom de la compagnie d'assurance/TPA: {lookup['Nom de la compagnie d\u2019assurance/TPA']}", size=9)

r2 = t.add_row()
para(r2.cells[0], f"Membre (N° ID/Carte): {lookup['Membre (N° ID/Carte)']}", size=9)
para(r2.cells[1], f"Titulaire de la police: {lookup['Titulaire de la police']}", size=9)

r3 = t.add_row()
c_left = r3.cells[0]
p = c_left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run(f"Date de naissance: {lookup['Date de naissance']}    Sexe : M {lookup['Sexe : M']}  F {lookup['Sexe : F']}")
r.font.size = Pt(9)
para(r3.cells[1], f"CPR/Passeport: {lookup['CPR/Passeport']}", size=9)

r4 = t.add_row()
c_left4 = r4.cells[0]
p = c_left4.paragraphs[0]
r = p.add_run(f"Célibataire {lookup['Célibataire']}    Marié(e) {lookup['Marié(e)']}")
r.font.size = Pt(9)
para(r4.cells[1], f"Numéro de téléphone du membre: {lookup['Numéro de téléphone du membre']}", size=9)

section_header(t, lookup["SECTION B : SECTION MÉDICALE (À REMPLIR UNIQUEMENT PAR LE MÉDECIN TRAITANT)"], 2)

r5 = t.add_row()
c_left5 = r5.cells[0]
p = c_left5.paragraphs[0]
r = p.add_run(
    f"Veuillez cocher : Hospitalisé {lookup['Veuillez cocher : Patient hospitalisé']}  "
    f"Ambulatoire {lookup['Veuillez cocher : Patient ambulatoire']}  "
    f"Urgence {lookup['Veuillez cocher : Cas d\u2019urgence']}"
)
r.font.size = Pt(9)
para(r5.cells[1], f"Nom du prestataire: {lookup['Nom du prestataire']}", size=9)

r6 = t.add_row()
para(r6.cells[0], f"Date du traitement: {lookup['Date du traitement']}", size=9)
para(r6.cells[1], f"Numéro de dossier médical: {lookup['Numéro de dossier médical']}", size=9)

r7 = t.add_row()
para(r7.cells[0], f"Condition préexistante: Acc. route {lookup['Condition préexistante : Accident de la route']}", size=9)
para(r7.cells[1], f"Signes vitaux: {lookup['Signes vitaux']}", size=9)

r8 = t.add_row()
para(r8.cells[0], f"Condition chronique: Acc. travail {lookup['Condition chronique : Accident du travail']}", size=9)
para(r8.cells[1], f"Tension artérielle: {lookup['Tension artérielle']}", size=9)

r9 = t.add_row()
para(r9.cells[0], f"Maternité / Date prévue d'accouchement: {lookup['Maternité : Date prévue d\u2019accouchement']}", size=9)
para(r9.cells[1], f"Pouls: {lookup['Pouls']}", size=9)

r10 = t.add_row()
para(r10.cells[0], f"Autres (veuillez préciser): {lookup['Autres (veuillez préciser)']}", size=9)
para(r10.cells[1], f"Température: {lookup['Température']}", size=9)

r11 = t.add_row()
c11_l = r11.cells[0]
c11_r = r11.cells[1]
para(c11_l, f"Plainte principale et symptômes présentés:\n{lookup['Plainte principale et symptômes présentés']}", size=9)
para(c11_r, f"Durée / Antécédents de la maladie: {lookup['Durée / Antécédents de la maladie']}", size=9)

r12 = t.add_row()
c12 = r12.cells[0]
c12.merge(r12.cells[1])
para(c12, f"Constatations cliniques et diagnostic final (utiliser les codes CIM le cas échéant):\n{lookup['Constatations cliniques et diagnostic final (utiliser les codes CIM le cas échéant)']}", size=9)

pr_hdr_row = t.add_row()
c_pr = pr_hdr_row.cells[0]
c_pr.merge(pr_hdr_row.cells[1])
shade_cell(c_pr)
para(c_pr, lookup["SECTION DE PRÉ-AUTORISATION"], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

r13 = t.add_row()
para(r13.cells[0], f"Autorisation / Prescription: {lookup['Autorisation / Prescription']}", size=9)
para(r13.cells[1], f"Durée de séjour prévue: {lookup['Durée de séjour prévue']}", size=9)

r14 = t.add_row()
para(r14.cells[0], f"Code de forfait: {lookup['Code de forfait']}", size=9)
para(r14.cells[1], f"COÛT ANTICIPÉ: {lookup['COÛT ANTICIPÉ']}", size=9, bold=True)

decl_hdr = t.add_row()
c_decl = decl_hdr.cells[0]
c_decl.merge(decl_hdr.cells[1])
shade_cell(c_decl, "F2F2F2")
para(c_decl, lookup["Déclaration du prestataire de services médicaux"], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

r15 = t.add_row()
c_sig_m = r15.cells[0]
p = c_sig_m.paragraphs[0]
r = p.add_run(
    "Nous certifions par la présente que TOUTES les informations mentionnées dans le présent document sont correctes "
    "et que les services médicaux indiqués sur ce formulaire étaient médicalement indiqués et nécessaires pour la prise en charge de ce cas."
)
r.font.size = Pt(7)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

c_doc = r15.cells[1]
p2 = c_doc.paragraphs[0]
r2 = p2.add_run(f"Nom du médecin: {lookup['Nom du médecin']}")
r2.font.size = Pt(9)
add_para(c_doc, f"Signature: {lookup['Signature (Médecin)']}", size=9)
add_para(c_doc, f"N° L. / Cachet & N° de licence: {lookup['N° L. / Cachet & N° de licence']}", size=9)
add_para(c_doc, f"Date (Médecin): {lookup['Date (Médecin)']}", size=9)

r16 = t.add_row()
para(r16.cells[0], f"Signature (Membre): {lookup['Signature (Membre)']}", size=9)
para(r16.cells[1], "", size=9)

r17 = t.add_row()
para(r17.cells[0], f"Date (Membre): {lookup['Date (Membre)']}", size=9)
para(r17.cells[1], "", size=9)

ins_hdr = t.add_row()
c_ins = ins_hdr.cells[0]
c_ins.merge(ins_hdr.cells[1])
shade_cell(c_ins)
para(c_ins, lookup["RÉSERVÉ À L'USAGE DE LA COMPAGNIE D'ASSURANCE"], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

r18 = t.add_row()
c18 = r18.cells[0]
c18.merge(r18.cells[1])
para(c18, f"Commentaires: {lookup['Commentaires']}", size=9)

r19 = t.add_row()
c_left19 = r19.cells[0]
p19 = c_left19.paragraphs[0]
r19r = p19.add_run(
    f"Approuvé {lookup['Approuvé']}    Non approuvé {lookup['Non approuvé']}\n"
    f"N° d'approbation: {dotline(15)}    Validité de l'approbation: {dotline(15)}\n"
    f"Agent d'assurance: {lookup['Agent d\u2019assurance']}    Signature: {lookup['Signature (Agent d\u2019assurance)']}"
)
r19r.font.size = Pt(9)
c_right19 = r19.cells[1]
p19r = c_right19.paragraphs[0]
rr = p19r.add_run(
    f"Date: {lookup['Date (Assurance)']}    N° de SINISTRE: {lookup['N° de SINISTRE']}"
)
rr.font.size = Pt(9)

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    "Hotline et contact de pré-autorisation: +973 17530886  Fax: +973 17530602\n"
    "Pour les approbations ambulatoires: op.approvals@globemedbahrain.com | Hospitalisés: lp.approvals@globemedbahrain.com\n"
    "Pour le remboursement des sinistres: cr@globemedbahrain.com"
)
fr.font.size = Pt(7)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)