import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_04.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p04.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key):
    return lookup.get(key, "")

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

def add_outer_border(doc):
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        pgBorders.append(border)
    sectPr = doc.sections[0]._sectPr
    sectPr.append(pgBorders)

add_outer_border(doc)

def add_empty_para(doc, space_pt=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_pt)
    return p

for i in range(8):
    add_empty_para(doc, space_pt=12)

header_para = doc.add_paragraph()
header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_para.paragraph_format.space_before = Pt(0)
header_para.paragraph_format.space_after = Pt(18)
header_run = header_para.add_run(get_value("PAGE D'IDENTITÉ EXCLUE"))
header_run.bold = True
header_run.font.size = Pt(24)
header_run.font.name = "Arial"
header_run.font.color.rgb = RGBColor(0, 0, 0)

body_fields = [
    "Nom du patient correspondant aux autres pages",
    ("Nom correspondant", "Nom correspondant"),
    ("Type de document", "Type de document"),
    ("Numéro de page original", "Numéro de page original"),
]

info_para1 = doc.add_paragraph()
info_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
info_para1.paragraph_format.space_before = Pt(0)
info_para1.paragraph_format.space_after = Pt(2)
run1 = info_para1.add_run(get_value("Nom du patient correspondant aux autres pages"))
run1.font.size = Pt(10)
run1.font.name = "Arial"
run1.font.color.rgb = RGBColor(0, 0, 0)

info_para2 = doc.add_paragraph()
info_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
info_para2.paragraph_format.space_before = Pt(0)
info_para2.paragraph_format.space_after = Pt(2)
run2 = info_para2.add_run("Nom correspondant: " + get_value("Nom correspondant"))
run2.font.size = Pt(10)
run2.font.name = "Arial"
run2.font.color.rgb = RGBColor(0, 0, 0)

info_para3 = doc.add_paragraph()
info_para3.alignment = WD_ALIGN_PARAGRAPH.LEFT
info_para3.paragraph_format.space_before = Pt(0)
info_para3.paragraph_format.space_after = Pt(2)
run3 = info_para3.add_run("Type de document: " + get_value("Type de document"))
run3.font.size = Pt(10)
run3.font.name = "Arial"
run3.font.color.rgb = RGBColor(0, 0, 0)

info_para4 = doc.add_paragraph()
info_para4.alignment = WD_ALIGN_PARAGRAPH.LEFT
info_para4.paragraph_format.space_before = Pt(0)
info_para4.paragraph_format.space_after = Pt(2)
run4 = info_para4.add_run("Numéro de page original: " + get_value("Numéro de page original"))
run4.font.size = Pt(10)
run4.font.name = "Arial"
run4.font.color.rgb = RGBColor(0, 0, 0)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)