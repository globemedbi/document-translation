import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_01.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p01.json"

with open(DATA_FILE, "r", encoding="utf-8") as f:
    data = json.load(f)

def get_value(key):
    for item in data:
        if item["key"] == key:
            return item["value"]
    return ""

NAVY = RGBColor(0x1e, 0x3a, 0x5f)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x44, 0x44, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Remove default paragraph spacing
style = doc.styles["Normal"]
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

def set_para_spacing(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def add_run_text(para, text, bold=False, italic=False, size=10, color=BLACK, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run

def add_horizontal_rule(doc, color=NAVY, thickness=6):
    para = doc.add_paragraph()
    set_para_spacing(para, 2, 2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(color.red, color.green, color.blue))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def add_section_header(doc, text, before=8, after=4):
    para = doc.add_paragraph()
    set_para_spacing(para, before, after)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    # underline the section header with a bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(NAVY.red, NAVY.green, NAVY.blue))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV", "left", "right"):
        tag = "w:{}".format(edge)
        if edge in kwargs or edge.replace("start","left").replace("end","right") in kwargs:
            val = kwargs.get(edge, kwargs.get(edge.replace("start","left").replace("end","right"), {}))
            el = OxmlElement(tag)
            for k, v in val.items():
                el.set(qn(k), v)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def no_border_cell(cell):
    set_cell_border(cell,
        top={"w:val":"none","w:sz":"0","w:space":"0","w:color":"auto"},
        bottom={"w:val":"none","w:sz":"0","w:space":"0","w:color":"auto"},
        left={"w:val":"none","w:sz":"0","w:space":"0","w:color":"auto"},
        right={"w:val":"none","w:sz":"0","w:space":"0","w:color":"auto"},
    )

def set_table_no_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement("w:{}".format(edge))
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

# ─── HEADER ───────────────────────────────────────────────────────────────────

# Full Name
p_name = doc.add_paragraph()
set_para_spacing(p_name, 0, 2)
run_name = p_name.add_run(get_value("Nom complet"))
run_name.bold = True
run_name.font.size = Pt(28)
run_name.font.color.rgb = DARK

# Title / Role
p_title = doc.add_paragraph()
set_para_spacing(p_title, 0, 2)
run_title = p_title.add_run(get_value("Titre / Rôle"))
run_title.font.size = Pt(11)
run_title.font.color.rgb = NAVY

# Contact line
p_contact = doc.add_paragraph()
set_para_spacing(p_contact, 0, 4)
contact_parts = [
    get_value("Localisation"),
    get_value("Téléphone"),
    get_value("E-mail"),
    get_value("LinkedIn"),
]
contact_text = "  •  ".join(contact_parts)
run_contact = p_contact.add_run(contact_text)
run_contact.font.size = Pt(9)
run_contact.font.color.rgb = GRAY

# Horizontal rule
add_horizontal_rule(doc, NAVY, 8)

# ─── PROFESSIONAL SUMMARY ────────────────────────────────────────────────────

add_section_header(doc, get_value("RÉSUMÉ PROFESSIONNEL"), before=6, after=3)

p_summary = doc.add_paragraph()
set_para_spacing(p_summary, 0, 6)
run_summary = p_summary.add_run(get_value("Texte du résumé professionnel"))
run_summary.font.size = Pt(9.5)
run_summary.font.color.rgb = BLACK

# ─── CORE SKILLS ─────────────────────────────────────────────────────────────

add_section_header(doc, get_value("COMPÉTENCES CLÉS"), before=6, after=3)

# Two-column table for skills
skills_table = doc.add_table(rows=3, cols=2)
set_table_no_border(skills_table)
skills_table.autofit = False

col_width = Inches(3.2)
for row in skills_table.rows:
    for cell in row.cells:
        no_border_cell(cell)
        cell.width = col_width
        set_cell_bg(cell, "F5F7FA")

# Left column skills
left_skills = [
    ("IA / ML", get_value("IA / ML")),
    ("Données et infrastructure", get_value("Données et infrastructure")),
    ("Architecture et leadership", get_value("Architecture et leadership")),
]
right_skills = [
    ("Programmation", get_value("Programmation")),
    ("Frameworks et outils", get_value("Frameworks et outils")),
]

for i, (label, val) in enumerate(left_skills):
    cell = skills_table.rows[i].cells[0]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = BLACK
    run_val = p.add_run(val)
    run_val.font.size = Pt(9)
    run_val.font.color.rgb = BLACK

for i, (label, val) in enumerate(right_skills):
    cell = skills_table.rows[i].cells[1]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = BLACK
    run_val = p.add_run(val)
    run_val.font.size = Pt(9)
    run_val.font.color.rgb = BLACK

doc.add_paragraph()

# ─── LANGUAGES ───────────────────────────────────────────────────────────────

add_section_header(doc, get_value("LANGUES"), before=6, after=3)

p_lang = doc.add_paragraph()
set_para_spacing(p_lang, 0, 6)

lang_name_1 = get_value("Langue 1")
lang_level_1 = get_value("Niveau 1")
lang_name_2 = get_value("Langue 2")
lang_level_2 = get_value("Niveau 2")
lang_name_3 = get_value("Langue 3")
lang_level_3 = get_value("Niveau 3")

def add_lang_entry(para, name, level, separator=True):
    r1 = para.add_run(name)
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = BLACK
    r2 = para.add_run(" \u2014 " + level)
    r2.font.size = Pt(9.5)
    r2.italic = True
    r2.font.color.rgb = GRAY
    if separator:
        r3 = para.add_run("   \u2022   ")
        r3.font.size = Pt(9.5)
        r3.font.color.rgb = GRAY

add_lang_entry(p_lang, lang_name_1, lang_level_1, separator=True)
add_lang_entry(p_lang, lang_name_2, lang_level_2, separator=True)
add_lang_entry(p_lang, lang_name_3, lang_level_3, separator=False)

# ─── EDUCATION ───────────────────────────────────────────────────────────────

add_section_header(doc, get_value("FORMATION"), before=6, after=3)

edu_entries = [
    {
        "degree": get_value("Diplôme 1"),
        "location": get_value("Lieu de formation 1"),
        "institution": get_value("Établissement 1"),
        "date": get_value("Date de formation 1"),
    },
    {
        "degree": get_value("Diplôme 2"),
        "location": get_value("Lieu de formation 2"),
        "institution": get_value("Établissement 2"),
        "date": get_value("Date de formation 2"),
    },
]

for edu in edu_entries:
    edu_table = doc.add_table(rows=1, cols=2)
    set_table_no_border(edu_table)
    edu_table.autofit = False

    left_cell = edu_table.rows[0].cells[0]
    right_cell = edu_table.rows[0].cells[1]
    no_border_cell(left_cell)
    no_border_cell(right_cell)

    left_cell.width = Inches(5.0)
    right_cell.width = Inches(1.5)

    # Degree
    p_deg = left_cell.paragraphs[0]
    p_deg.paragraph_format.space_before = Pt(2)
    p_deg.paragraph_format.space_after = Pt(0)
    r_deg = p_deg.add_run(edu["degree"])
    r_deg.bold = True
    r_deg.font.size = Pt(10)
    r_deg.font.color.rgb = BLACK

    # Location
    p_loc = left_cell.add_paragraph()
    p_loc.paragraph_format.space_before = Pt(0)
    p_loc.paragraph_format.space_after = Pt(0)
    r_loc = p_loc.add_run(edu["location"])
    r_loc.italic = True
    r_loc.font.size = Pt(9)
    r_loc.font.color.rgb = GRAY

    # Institution
    p_inst = left_cell.add_paragraph()
    p_inst.paragraph_format.space_before = Pt(0)
    p_inst.paragraph_format.space_after = Pt(4)
    r_inst = p_inst.add_run(edu["institution"])
    r_inst.font.size = Pt(9)
    r_inst.font.color.rgb = BLACK

    # Date (right-aligned)
    p_date = right_cell.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(2)
    p_date.paragraph_format.space_after = Pt(0)
    r_date = p_date.add_run(edu["date"])
    r_date.bold = True
    r_date.font.size = Pt(9.5)
    r_date.font.color.rgb = NAVY

# ─── PROFESSIONAL EXPERIENCE ─────────────────────────────────────────────────

add_section_header(doc, get_value("EXPÉRIENCE PROFESSIONNELLE"), before=6, after=3)

exp_count_key_pairs = [
    ("Poste 1", "Entreprise 1", "Dates du poste 1", "Lieu du poste 1",
     ["Réalisation 1.1", "Réalisation 1.2", "Réalisation 1.3", "Réalisation 1.4"]),
    ("Poste 2", "Entreprise 2", "Dates du poste 2", "Lieu du poste 2",
     ["Réalisation 2.1", "Réalisation 2.2", "Réalisation 2.3"]),
]

def add_experience_block(doc, title_key, company_key, date_key, loc_key, bullet_keys):
    title_val = get_value(title_key)
    company_val = get_value(company_key)
    date_val = get_value(date_key)
    loc_val = get_value(loc_key)

    # Header row: Title | Company on left, Date on right
    exp_table = doc.add_table(rows=1, cols=2)
    set_table_no_border(exp_table)
    exp_table.autofit = False

    left_cell = exp_table.rows[0].cells[0]
    right_cell = exp_table.rows[0].cells[1]
    no_border_cell(left_cell)
    no_border_cell(right_cell)
    left_cell.width = Inches(5.0)
    right_cell.width = Inches(1.5)

    set_cell_bg(left_cell, "EEF2F7")
    set_cell_bg(right_cell, "EEF2F7")

    p_title_co = left_cell.paragraphs[0]
    p_title_co.paragraph_format.space_before = Pt(3)
    p_title_co.paragraph_format.space_after = Pt(2)
    r_title = p_title_co.add_run(title_val)
    r_title.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = NAVY

    if company_val:
        r_sep = p_title_co.add_run("  |  ")
        r_sep.font.size = Pt(10)
        r_sep.font.color.rgb = GRAY
        r_company = p_title_co.add_run(company_val)
        r_company.font.size = Pt(10)
        r_company.font.color.rgb = NAVY

    p_date_right = right_cell.paragraphs[0]
    p_date_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date_right.paragraph_format.space_before = Pt(3)
    p_date_right.paragraph_format.space_after = Pt(2)
    r_date = p_date_right.add_run(date_val)
    r_date.bold = True
    r_date.font.size = Pt(9.5)
    r_date.font.color.rgb = NAVY

    # Location
    p_loc = doc.add_paragraph()
    p_loc.paragraph_format.space_before = Pt(0)
    p_loc.paragraph_format.space_after = Pt(2)
    r_loc = p_loc.add_run(loc_val)
    r_loc.italic = True
    r_loc.font.size = Pt(9)
    r_loc.font.color.rgb = GRAY

    # Bullets
    for bkey in bullet_keys:
        bval = get_value(bkey)
        if bval:
            p_b = doc.add_paragraph(style="List Bullet")
            p_b.paragraph_format.space_before = Pt(1)
            p_b.paragraph_format.space_after = Pt(1)
            p_b.paragraph_format.left_indent = Inches(0.2)
            r_b = p_b.add_run(bval)
            r_b.font.size = Pt(9)
            r_b.font.color.rgb = BLACK

    # Small spacer
    sp = doc.add_paragraph()
    set_para_spacing(sp, 2, 2)

for args in exp_count_key_pairs:
    add_experience_block(doc, args[0], args[1], args[2], args[3], args[4])

doc.save(OUTPUT_PATH)
print("Saved to", OUTPUT_PATH)