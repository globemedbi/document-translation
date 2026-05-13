import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_05.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p05.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key, default=""):
    return lookup.get(key, default)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_table_borders(table, border_size=4, color="000000"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = OxmlElement("w:" + border_name)
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), str(border_size))
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), color)
        tblBorders.append(border_el)
    tblPr.append(tblBorders)

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="000000", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    sides = {"top": top, "bottom": bottom, "left": left, "right": right}
    for side, active in sides.items():
        el = OxmlElement("w:" + side)
        if active:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def shade_cell(cell, hex_color):
    set_cell_bg(cell, hex_color)

def add_paragraph_with_text(cell, text, bold=False, font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    para.alignment = alignment
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(widths[i].twips)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Remove default paragraph spacing
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)

# ─── HEADER ───
header_p1 = doc.add_paragraph()
header_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = header_p1.add_run(get_value("Dr. Mohsin Kaldas Clinic"))
run1.bold = True
run1.font.size = Pt(14)

header_p2 = doc.add_paragraph()
header_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = header_p2.add_run(get_value("Formulaire d'approbation"))
run2.font.size = Pt(11)

header_p3 = doc.add_paragraph()
header_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = header_p3.add_run(get_value("R\u00e9f\u00e9rence#") + ":  " + get_value("R\u00e9f\u00e9rence#"))
run3.font.size = Pt(10)

# Fix: use correct key for reference
ref_p = doc.add_paragraph()
ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ref_p.clear()
ref_run = ref_p.add_run("R\u00e9f\u00e9rence#:  " + get_value("R\u00e9f\u00e9rence#"))
ref_run.font.size = Pt(10)

# Remove the incorrectly built header_p3
header_p3._element.getparent().remove(header_p3._element)

# Thin horizontal rule
def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

add_horizontal_rule(doc)

# ─── PATIENT INFO TABLE ───
info_table = doc.add_table(rows=7, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(info_table, border_size=0, color="FFFFFF")

col_widths_info = [Cm(7), Cm(11)]
labels = [
    "Num\u00e9ro individuel",
    "Nom du titulaire de la carte",
    "Date de naissance",
    "Compagnie d'assurance",
    "M\u00e9decin",
    "Date d'approbation",
    "Date d'approbation (fin)",
]
values = [
    get_value("Num\u00e9ro individuel"),
    get_value("Nom du titulaire de la carte"),
    get_value("Date de naissance"),
    get_value("Compagnie d'assurance"),
    get_value("M\u00e9decin"),
    get_value("Date d'approbation"),
    get_value("Date d'approbation (fin)"),
]

for i, (label, value) in enumerate(zip(labels, values)):
    row = info_table.rows[i]
    # Label cell
    lc = row.cells[0]
    lc.width = col_widths_info[0]
    lp = lc.paragraphs[0]
    lp.clear()
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Value cell
    vc = row.cells[1]
    vc.width = col_widths_info[1]
    vp = vc.paragraphs[0]
    vp.clear()

    if label == "Nom du titulaire de la carte":
        # Black bar for redacted
        vr = vp.add_run("          ")
        vr.font.highlight_color = None
        vr.font.size = Pt(10)
        from docx.oxml import OxmlElement as OE
        rPr = vr._r.get_or_add_rPr()
        shd = OE("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "000000")
        rPr.append(shd)
    else:
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
    vp.alignment = WD_ALIGN_PARAGRAPH.LEFT

add_horizontal_rule(doc)

# ─── SERVICE TABLE 1 ───
def add_service_table_header(doc):
    tbl = doc.add_table(rows=2, cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    # Header row
    headers = [
        get_value("Service"),
        get_value("Article"),
        get_value("Qt\u00e9 r\u00e9clam\u00e9e"),
        get_value("Qt\u00e9 approuv\u00e9e"),
        get_value("Statut"),
        get_value("Notes"),
        get_value("Motif"),
    ]
    col_w = [Cm(3.5), Cm(4), Cm(2), Cm(2), Cm(2.5), Cm(2), Cm(2)]
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = col_w[i]
        shade_cell(cell, "D9D9D9")
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data row 1
    data_row = tbl.rows[1]
    row_data = [
        get_value("Ligne de service 1 - Service"),
        get_value("Ligne de service 1 - Article"),
        get_value("Ligne de service 1 - Qt\u00e9 r\u00e9clam\u00e9e"),
        get_value("Ligne de service 1 - Qt\u00e9 approuv\u00e9e"),
        get_value("Ligne de service 1 - Statut"),
        "",
        "",
    ]
    for i, val in enumerate(row_data):
        cell = data_row.cells[i]
        cell.width = col_w[i]
        p = cell.paragraphs[0]
        p.clear()
        if i == 1:
            # Item with bullet
            r = p.add_run("\u25a0 " + val)
        elif i == 4:
            r = p.add_run(val)
            r.bold = True
        else:
            r = p.add_run(val)
        r.font.size = Pt(9)
        if i == 4:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return tbl

svc_table1 = add_service_table_header(doc)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─── SERVICE TABLE 2 ───
tbl2 = doc.add_table(rows=3, cols=7)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tbl2)

col_w2 = [Cm(3.5), Cm(4), Cm(2), Cm(2), Cm(2.5), Cm(2), Cm(2)]

# Header row
headers2 = [
    get_value("Service"),
    get_value("Article"),
    get_value("Qt\u00e9 r\u00e9clam\u00e9e"),
    get_value("Qt\u00e9 approuv\u00e9e"),
    get_value("Statut"),
    get_value("Notes"),
    get_value("Motif"),
]
hdr_row2 = tbl2.rows[0]
for i, h in enumerate(headers2):
    cell = hdr_row2.cells[i]
    cell.width = col_w2[i]
    shade_cell(cell, "D9D9D9")
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data row 1 - LINCOCIN
row2_1 = tbl2.rows[1]
row2_1_data = [
    get_value("Ligne de service 2 - Service"),
    get_value("Ligne de service 2 - Article 1"),
    get_value("Ligne de service 2 - Article 1 Qt\u00e9 r\u00e9clam\u00e9e"),
    get_value("Ligne de service 2 - Article 1 Qt\u00e9 approuv\u00e9e"),
    get_value("Ligne de service 2 - Article 1 Statut"),
    "",
    "",
]
for i, val in enumerate(row2_1_data):
    cell = row2_1.cells[i]
    cell.width = col_w2[i]
    p = cell.paragraphs[0]
    p.clear()
    if i == 1:
        r = p.add_run("\u25a0 " + val)
    elif i == 4:
        r = p.add_run(val)
        r.bold = True
    else:
        r = p.add_run(val)
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data row 2 - Depomedrol
row2_2 = tbl2.rows[2]
row2_2_data = [
    "",
    get_value("Ligne de service 2 - Article 2"),
    get_value("Ligne de service 2 - Article 2 Qt\u00e9 r\u00e9clam\u00e9e"),
    get_value("Ligne de service 2 - Article 2 Qt\u00e9 approuv\u00e9e"),
    get_value("Ligne de service 2 - Article 2 Statut"),
    get_value("Ligne de service 2 - Article 2 Remarques"),
    "",
]
for i, val in enumerate(row2_2_data):
    cell = row2_2.cells[i]
    cell.width = col_w2[i]
    p = cell.paragraphs[0]
    p.clear()
    if i == 1:
        r = p.add_run("\u25a0 " + val)
    elif i == 4:
        r = p.add_run(val)
        r.bold = True
    else:
        r = p.add_run(val)
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Coverage Details row - merged
cov_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(cov_tbl)
cov_cell = cov_tbl.rows[0].cells[0]
shade_cell(cov_cell, "D9D9D9")
cp = cov_cell.paragraphs[0]
cp.clear()
cr = cp.add_run(get_value("D\u00e9tails de la couverture"))
cr.bold = True
cr.font.size = Pt(9)
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─── IMPORTANT NOTES SECTION ───
imp_notes_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(imp_notes_tbl)
shade_cell(imp_notes_tbl.rows[0].cells[0], "D9D9D9")
inp = imp_notes_tbl.rows[0].cells[0].paragraphs[0]
inp.clear()
inr = inp.add_run(get_value("Notes importantes"))
inr.bold = True
inr.font.size = Pt(10)

# Notes content
notes_content_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(notes_content_tbl)
nc = notes_content_tbl.rows[0].cells[0]
nc_p = nc.paragraphs[0]
nc_p.clear()

# Build notes text
notes_intro = get_value("Veuillez joindre les documents requis lors de la soumission de la demande de remboursement")
if not notes_intro:
    notes_intro = "Veuillez joindre les documents requis lors de la soumission de la demande de remboursement"

nc_run = nc_p.add_run(notes_intro)
nc_run.font.size = Pt(9)
nc_run.italic = True
nc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

notes_items = [
    get_value("1-Formulaire d'approbation") or "1-Formulaire d'approbation",
    get_value("2-Copie de la carte d'identit\u00e9 et/ou passeport, CPR, permis de conduire") or "2-Copie de la carte d'identit\u00e9 et/ou passeport, CPR, permis de conduire",
    get_value("3-Formulaire de demande de remboursement") or "3-Formulaire de demande de remboursement",
    get_value("4-Copie de la carte d'assurance") or "4-Copie de la carte d'assurance",
    get_value("5-Rapport m\u00e9dical et r\u00e9sultats du test demand\u00e9, le cas \u00e9ch\u00e9ant") or "5-Rapport m\u00e9dical et r\u00e9sultats du test demand\u00e9, le cas \u00e9ch\u00e9ant",
    get_value("6-R\u00e9sum\u00e9 de sortie") or "6-R\u00e9sum\u00e9 de sortie",
]

for item_text in notes_items:
    nc.add_paragraph()
    item_p = nc.paragraphs[-1]
    item_p.clear()
    item_r = item_p.add_run(item_text)
    item_r.font.size = Pt(9)
    item_p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─── DIAGNOSIS SECTION ───
diag_hdr_table = doc.add_table(rows=1, cols=1)
set_table_borders(diag_hdr_table)
shade_cell(diag_hdr_table.rows[0].cells[0], "D9D9D9")
dh = diag_hdr_table.rows[0].cells[0].paragraphs[0]
dh.clear()
dhr = dh.add_run(get_value("Diagnostic") or "Diagnostic")
dhr.bold = True
dhr.font.size = Pt(10)

diag_content = doc.add_table(rows=2, cols=1)
set_table_borders(diag_content)
diag_val1 = get_value("Diagnostic 1") or "Otite s\u00e9reuse aigu\u00eb"
diag_val2 = get_value("Diagnostic 2") or "Otite s\u00e9reuse aigu\u00eb"

for i, dval in enumerate([diag_val1, diag_val2]):
    dp = diag_content.rows[i].cells[0].paragraphs[0]
    dp.clear()
    dr = dp.add_run(dval)
    dr.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─── MEDICAL INFORMATION SECTION ───
med_hdr_table = doc.add_table(rows=1, cols=1)
set_table_borders(med_hdr_table)
shade_cell(med_hdr_table.rows[0].cells[0], "D9D9D9")
mh = med_hdr_table.rows[0].cells[0].paragraphs[0]
mh.clear()
mhr = mh.add_run(get_value("Informations m\u00e9dicales") or "Informations m\u00e9dicales")
mhr.bold = True
mhr.font.size = Pt(10)

med_tbl = doc.add_table(rows=2, cols=7)
set_table_borders(med_tbl)
med_headers = ["Date", "TA", "Pouls", "Temp\u00e9rature", "Dur\u00e9e de la maladie", "Plaintes", "Signes", "Autres"]
med_headers_short = ["Date", "TA", "Pouls", "Temp\u00e9rature", "Dur\u00e9e de la maladie", "Plaintes", "Signes"]

shade_cell(med_tbl.rows[0].cells[0], "D9D9D9")
for i, mh_text in enumerate(med_headers_short):
    cell = med_tbl.rows[0].cells[i] if i < len(med_tbl.rows[0].cells) else None
    if cell:
        shade_cell(cell, "D9D9D9")
        mp = cell.paragraphs[0]
        mp.clear()
        mr = mp.add_run(mh_text)
        mr.bold = True
        mr.font.size = Pt(8)
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Empty data row
for cell in med_tbl.rows[1].cells:
    cp = cell.paragraphs[0]
    cp.clear()
    cp.add_run(" ").font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─── UPLOADED FILES SECTION ───
upload_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(upload_tbl)
shade_cell(upload_tbl.rows[0].cells[0], "D9D9D9")
uh = upload_tbl.rows[0].cells[0].paragraphs[0]
uh.clear()
uhr = uh.add_run(get_value("Fichiers t\u00e9l\u00e9charg\u00e9s") or "Fichiers t\u00e9l\u00e9charg\u00e9s")
uhr.bold = True
uhr.font.size = Pt(10)

upload_content = doc.add_table(rows=1, cols=1)
set_table_borders(upload_content)
uc = upload_content.rows[0].cells[0].paragraphs[0]
uc.clear()
uc.add_run(" ").font.size = Pt(9)

# ─── ADDITIONAL INFO SECTION ───
add_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(add_tbl)
shade_cell(add_tbl.rows[0].cells[0], "D9D9D9")
ah = add_tbl.rows[0].cells[0].paragraphs[0]
ah.clear()
ahr = ah.add_run(get_value("Informations suppl\u00e9mentaires") or "Informations suppl\u00e9mentaires")
ahr.bold = True
ahr.font.size = Pt(10)

add_content = doc.add_table(rows=2, cols=1)
set_table_borders(add_content)
for row in add_content.rows:
    rp = row.cells[0].paragraphs[0]
    rp.clear()
    rp.add_run(" ").font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── FOOTER ───
footer_code = doc.add_paragraph()
footer_code.alignment = WD_ALIGN_PARAGRAPH.LEFT
fc_r = footer_code.add_run("%09%09")
fc_r.font.size = Pt(9)

footer_disclaimer_tbl = doc.add_table(rows=1, cols=1)
set_table_borders(footer_disclaimer_tbl)
fd_cell = footer_disclaimer_tbl.rows[0].cells[0]
fd_p = fd_cell.paragraphs[0]
fd_p.clear()
fd_text = get_value("En cas d'approbation, cette demande est soumise \u00e0 un examen des tarifs, des remises et de la n\u00e9cessit\u00e9 m\u00e9dicale \u2013 Habituelle, Raisonnable et Personnalis\u00e9e.")
if not fd_text:
    fd_text = "En cas d'approbation, cette demande est soumise \u00e0 un examen des tarifs, des remises et de la n\u00e9cessit\u00e9 m\u00e9dicale \u2013 Habituelle, Raisonnable et Personnalis\u00e9e."
fd_r = fd_p.add_run(fd_text)
fd_r.font.size = Pt(9)
fd_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Thank you
ty_p = fd_cell.add_paragraph()
ty_r = ty_p.add_run(get_value("Merci.") or "Merci.")
ty_r.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# VAT statement
vat_p = doc.add_paragraph()
vat_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
vat_text = get_value("Cette demande est soumise aux conditions de TVA du Royaume de Bahr\u00efn")
if not vat_text:
    vat_text = "Cette demande est soumise aux conditions de TVA du Royaume de Bahr\u00efn"
vat_r = vat_p.add_run(vat_text)
vat_r.bold = True
vat_r.font.size = Pt(10)

doc.save(OUTPUT_PATH)
print("Saved:", OUTPUT_PATH)