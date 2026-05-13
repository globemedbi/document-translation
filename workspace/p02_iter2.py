import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_02.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p02.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key):
    return lookup.get(key, "")

def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_run_with_format(para, text, bold=False, italic=False, font_size=10,
                        color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

def add_job_header(doc, title_key, employer_key, period_key, location_key):
    title = get_value(title_key)
    employer = get_value(employer_key)
    period = get_value(period_key)
    location = get_value(location_key)

    table = doc.add_table(rows=1, cols=2)
    set_table_borders_none(table)
    table.autofit = False

    # Set column widths
    total_width = Inches(6.5)
    left_width = Inches(4.5)
    right_width = Inches(2.0)

    table.columns[0].width = left_width
    table.columns[1].width = right_width

    row = table.rows[0]
    row.cells[0].width = left_width
    row.cells[1].width = right_width

    # Left cell: title | employer
    left_cell = row.cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(left_para, before=0, after=0)

    add_run_with_format(left_para, title, bold=True, font_size=11,
                        color=(31, 56, 100))
    add_run_with_format(left_para, "  |  ", bold=False, font_size=11,
                        color=(180, 180, 180))
    add_run_with_format(left_para, employer, bold=False, font_size=11,
                        color=(31, 56, 100))

    # Right cell: period
    right_cell = row.cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(right_para, before=0, after=0)
    add_run_with_format(right_para, period, bold=True, font_size=10,
                        color=(31, 56, 100))

    # Location row
    loc_para = doc.add_paragraph()
    loc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(loc_para, before=0, after=0)
    add_run_with_format(loc_para, location, italic=True, font_size=9,
                        color=(120, 120, 120))

    return table

def add_bullet(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before=0, after=20, line=276)

    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)

    run_bullet = para.add_run("\u2022  ")
    run_bullet.font.size = Pt(9.5)
    run_bullet.font.color.rgb = RGBColor(0, 0, 0)

    run_text = para.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(0, 0, 0)
    return para

def add_spacer(doc, space_before=80):
    para = doc.add_paragraph()
    set_para_spacing(para, before=space_before, after=0)
    return para

def add_separator_line(doc):
    para = doc.add_paragraph()
    set_para_spacing(para, before=40, after=40)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ─── JOB 1: AI Engineer | Inheaden ───────────────────────────────────────────
add_job_header(doc,
               "Intitulé du poste 1",
               "Employeur 1",
               "Période 1",
               "Lieu 1")

add_bullet(doc, get_value("Point 1.1"))
add_bullet(doc, get_value("Point 1.2"))
add_bullet(doc, get_value("Point 1.3"))
add_bullet(doc, get_value("Point 1.4"))

add_spacer(doc, space_before=60)

# ─── JOB 2: Machine Learning Engineer | GlobeMed ─────────────────────────────
add_job_header(doc,
               "Intitulé du poste 2",
               "Employeur 2",
               "Période 2",
               "Lieu 2")

add_bullet(doc, get_value("Point 2.1"))
add_bullet(doc, get_value("Point 2.2"))
add_bullet(doc, get_value("Point 2.3"))

add_spacer(doc, space_before=60)

# ─── JOB 3: AI Engineer | Augment Digital ────────────────────────────────────
add_job_header(doc,
               "Intitulé du poste 3",
               "Employeur 3",
               "Période 3",
               "Lieu 3")

add_bullet(doc, get_value("Point 3.1"))
add_bullet(doc, get_value("Point 3.2"))

add_spacer(doc, space_before=60)

# ─── JOB 4: Machine Learning Engineer Intern | Quantmetry ────────────────────
add_job_header(doc,
               "Intitulé du poste 4",
               "Employeur 4",
               "Période 4",
               "Lieu 4")

add_bullet(doc, get_value("Point 4.1"))
add_bullet(doc, get_value("Point 4.2"))
add_bullet(doc, get_value("Point 4.3"))

# ─── Page number ─────────────────────────────────────────────────────────────
add_spacer(doc, space_before=120)
page_para = doc.add_paragraph()
page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(page_para, before=0, after=0)
add_run_with_format(page_para, "Page ", font_size=8, color=(150, 150, 150))
add_run_with_format(page_para, "2", bold=False, font_size=8,
                    color=(150, 150, 150))

doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")