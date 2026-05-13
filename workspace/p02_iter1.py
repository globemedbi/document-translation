import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/page_02.docx"
DATA_FILE = "/Users/patricksaade/Desktop/Work/Globemed/document_translation/v3/workspace/data_p02.json"

DATA = json.load(open(DATA_FILE, encoding="utf-8"))
lookup = {d["key"]: d["value"] for d in DATA}

def get_value(key):
    return lookup.get(key, "")

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        val = kwargs.get(edge, {"val": "nil"})
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), val.get("val", "nil"))
        if "sz" in val:
            element.set(qn("w:sz"), str(val["sz"]))
        if "color" in val:
            element.set(qn("w:color"), val["color"])
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_bottom_border_to_row(row):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "CCCCCC")
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

def add_run_with_format(para, text, bold=False, italic=False, size_pt=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_job_block(doc, job_num):
    title_key = "Intitulé du poste " + str(job_num)
    employer_key = "Employeur " + str(job_num)
    period_key = "Période " + str(job_num)
    location_key = "Lieu " + str(job_num)

    title_val = get_value(title_key)
    employer_val = get_value(employer_key)
    period_val = get_value(period_key)
    location_val = get_value(location_key)

    table = doc.add_table(rows=2, cols=2)
    set_table_borders_none(table)
    table.columns[0].width = Cm(12.5)
    table.columns[1].width = Cm(4.5)

    row0 = table.rows[0]
    row0.cells[0].width = Cm(12.5)
    row0.cells[1].width = Cm(4.5)

    left_cell = row0.cells[0]
    right_cell = row0.cells[1]

    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)

    add_run_with_format(left_para, title_val, bold=True, size_pt=11, color=(26, 26, 64))
    add_run_with_format(left_para, "  |  ", bold=False, size_pt=11, color=(150, 150, 150))
    add_run_with_format(left_para, employer_val, bold=False, size_pt=11, color=(26, 26, 64))

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    add_run_with_format(right_para, period_val, bold=True, size_pt=10, color=(26, 26, 64))

    row1 = table.rows[1]
    loc_cell = row1.cells[0]
    loc_para = loc_cell.paragraphs[0]
    loc_para.paragraph_format.space_before = Pt(0)
    loc_para.paragraph_format.space_after = Pt(2)
    add_run_with_format(loc_para, location_val, bold=False, italic=True, size_pt=9, color=(120, 120, 120))

    merged = row1.cells[0].merge(row1.cells[1])

    add_bottom_border_to_row(row0)

    bullet_num = 1
    while True:
        bullet_key = "Point " + str(job_num) + "." + str(bullet_num)
        bullet_val = get_value(bullet_key)
        if not bullet_val:
            break
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(0.5)
        bp.paragraph_format.first_line_indent = Cm(-0.3)
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = Pt(14)
        add_run_with_format(bp, "\u2022  " + bullet_val, bold=False, size_pt=9.5, color=(0, 0, 0))
        bullet_num += 1

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

for job_num in range(1, 5):
    title_key = "Intitulé du poste " + str(job_num)
    if get_value(title_key):
        add_job_block(doc, job_num)

from docx.oxml import OxmlElement as OE

def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)

    run_label = footer_para.add_run("Page ")
    run_label.font.size = Pt(8)
    run_label.font.color.rgb = RGBColor(150, 150, 150)

    fldChar1 = OE("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OE("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OE("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run_field = footer_para.add_run()
    run_field.font.size = Pt(8)
    run_field.font.bold = True
    run_field.font.color.rgb = RGBColor(0, 0, 0)
    run_field._r.append(fldChar1)
    run_field._r.append(instrText)
    run_field._r.append(fldChar2)

add_page_number_footer(doc)

doc.save(OUTPUT_PATH)
print("SUCCESS: saved to " + OUTPUT_PATH)