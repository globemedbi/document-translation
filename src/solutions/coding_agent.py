from __future__ import annotations

"""
Solution 2 — CODING mode (PAO: Plan → Action → Observe per page)
=================================================================
Architecture per page:
  PLAN     — Claude sees the page image + translated fields → layout plan
  ACTION   — Claude generates a python-docx script for that single page
             (syntax-checked, executed; on error the error is fed back and
              the ACTION step is retried up to MAX_AGENT_ITERATIONS times)
  OBSERVE  — Claude compares original page image vs generated DOCX text dump
             → score 1-10 + specific fix instructions
             If score < 7 and iterations remain → loop back to ACTION

All pages run in parallel (up to 3 at a time).
Per-page DOCX files are merged into one final document, then PDF conversion
is attempted via docx2pdf.
"""

import json
import re
import shutil
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any

from loguru import logger
from rich.console import Console
from rich.panel import Panel

from src.config import get_client, settings
from src.llm.schema.layout import LayoutPlan
from src.llm.schema.observer import ObserverResult
from src.llm.structured import chat, image_block_from_bytes, structured_call
from src.models import TranslatedDocument, TranslatedPage

console = Console()

# ── Models ────────────────────────────────────────────────────────────────────

_CODER = settings.coding_model  # code generation (o3)
_SONNET = settings.model        # observe + misc (gpt-5.5)
_OPUS = settings.heavy_model    # plan (gpt-5.5)


# ── System prompts ────────────────────────────────────────────────────────────

_PLANNER_SYSTEM = """\
You are a document layout expert and visual designer. You receive:
  - An image of ONE page of a medical/insurance form
  - The translated field data for that page (JSON list)

Analyse the visual structure and fill every field of the structured layout plan:

- header: describe the page header — title text, logo position (left/right/center), any reference numbers shown, background color, font color
- footer: describe the footer content and styling, or write "none"
- layout: overall page structure — number of sections, LTR or RTL reading direction, column layout, approximate margins
- colors: one entry per distinct colored element (page header, section headers, footer, highlighted rows, etc.) — use exact hex codes where visible, or best estimates
- tables: ALL tables top to bottom — for each: exact row/col counts, column widths as percentages (must sum to 100), header row colors, which field keys belong in which table (use EXACT key strings from the JSON)
- images: logo or image placements with position, or "none"
- content: anything outside tables — standalone labels, signature lines, checkboxes not in a table, divider lines, instructional text

Be PRECISE. Every field key from the JSON must appear in exactly one table's fields list or in content.
"""

_CODER_SYSTEM = """\
Write a compact python-docx script. Simple, no comments, no docstrings.

INPUT: page image (visual reference), layout plan, translated field data (JSON), OUTPUT_PATH, DATA_FILE.

SCRIPT STRUCTURE:
  import json
  from docx import Document
  from docx.shared import Pt, Cm, RGBColor
  from docx.enum.text import WD_ALIGN_PARAGRAPH
  from docx.oxml.ns import qn
  from docx.oxml import OxmlElement
  DATA_FILE = "..."
  OUTPUT_PATH = "..."
  DATA = json.load(open(DATA_FILE, encoding="utf-8"))
  lookup = {d["key"]: d["value"] for d in DATA}
  doc = Document()
  # ... build document ...
  doc.save(OUTPUT_PATH)
  print("SUCCESS: saved to " + OUTPUT_PATH)

COLORS — always reproduce the original design:
- Set cell background: use this helper:
    def set_cell_bg(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)
- Set font color: run.font.color.rgb = RGBColor(0xRR, 0xGG, 0xBB)
- Header cells with dark background MUST have white font color
- Apply colors as described in the layout plan — never leave headers plain white if the original is colored

RULES:
- Match the visual layout: same tables, same columns, same font sizes (~12pt headers, ~9pt body)
- Use short variable names: doc, t (table), r (row), c (cell), s (section)
- Use helper functions to avoid repetition: def row(t, k, v): ...
- Loop over DATA for repetitive fields — never hard-code each one individually
- Use exact key strings from DATA_SAMPLE — never invent names
- Double-quoted strings only
- RTL text: par.paragraph_format.bidi = True; par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
- add_table() ALWAYS requires a width: doc.add_table(rows=N, cols=M) followed by setting width via table.width = Cm(17), OR use a helper. Never call add_table() inside a footer/header without Cm width.
- No inline comments, no explanations — output ONLY the Python code
"""

_CODER_RETRY_SYSTEM = """\
Fix the broken python-docx script. Output ONLY the corrected Python code.
- Fix only what the error/feedback describes. Keep OUTPUT_PATH and DATA_FILE unchanged.
- Double-quoted strings only.
- If fields are empty: use the EXACT key strings from DATA_SAMPLE in the lookup.
- add_table() inside a footer or header requires a width argument: section.footer.add_table(rows=1, cols=2, width=Cm(17))
- No comments, no explanations.
"""

_OBSERVER_SYSTEM = """\
You are a quality-control expert for document reconstruction.
The document has been TRANSLATED — all labels and values in the output are in the TARGET LANGUAGE.
This is CORRECT and expected. Do NOT penalise for translation.

You receive:
  - The ORIGINAL page image (in the source language)
  - A text dump of the GENERATED DOCX content (in the target language)
  - The target language name

Score the reconstruction on a scale of 1-10:
  10 = perfect — all fields present, correct values, proper structure
   8 = good — minor formatting issues, all important data is there
   5 = acceptable — some fields missing or misplaced
   3 = poor — many fields wrong or missing
   1 = failed — completely wrong or empty content

Set is_done to true if score >= 8.
List any original-language field keys not found in the output in missing_fields.
List fields with incorrect values in wrong_values.
Write 2-3 sentences of actionable feedback using only ASCII characters.
"""


# ── DOCX helpers ──────────────────────────────────────────────────────────────


def _merge_docx_pages(page_paths: list[Path]) -> Any:
    """
    Merge multiple single-page DOCX files into one Document.
    Each page is separated by a hard page break inserted at the XML level.
    The final sectPr of each sub-doc is skipped (only the merged doc's is kept).
    """
    import copy

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not page_paths:
        doc = Document()
        doc.add_paragraph("No pages generated.")
        return doc

    merged = Document(str(page_paths[0]))

    for path in page_paths[1:]:
        next_doc = Document(str(path))

        # Insert a hard page break as a bare paragraph before appending body elements
        br = OxmlElement("w:p")
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        page_br = OxmlElement("w:br")
        page_br.set(qn("w:type"), "page")
        r.append(rPr)
        r.append(page_br)
        br.append(r)
        merged.element.body.append(br)

        for elem in next_doc.element.body:
            # Skip the section properties of each sub-doc; keep only the merged doc's
            if elem.tag.endswith("}sectPr"):
                continue
            merged.element.body.append(copy.deepcopy(elem))

    return merged


def _dump_docx_text(docx_path: Path) -> str:
    """
    Produce a plain-text summary of a DOCX file for the Observer to review.
    Lists all non-empty paragraphs and all table cells (row | col | col...).
    This is what the Observer compares against the original page image.
    """
    try:
        from docx import Document

        doc = Document(str(docx_path))
        lines: list[str] = []

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                lines.append(f"[P{i}] {para.text.strip()}")

        for t_idx, table in enumerate(doc.tables):
            lines.append(f"[Table {t_idx}: {len(table.rows)}r × {len(table.columns)}c]")
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells)
                if row_text.strip():
                    lines.append(f"  {row_text}")

        return "\n".join(lines) if lines else "(empty document)"
    except Exception as exc:
        return f"(could not read DOCX: {exc})"


def _word_installed() -> bool:
    """Check if Microsoft Word is available (required by docx2pdf on macOS/Windows)."""
    import platform
    if platform.system() == "Darwin":
        return Path("/Applications/Microsoft Word.app").exists()
    if platform.system() == "Windows":
        import shutil
        return shutil.which("WINWORD.EXE") is not None
    return False


def _convert_docx_to_pdf(docx_path: Path) -> Path | None:
    """Convert DOCX to PDF via docx2pdf (needs Microsoft Word on macOS/Windows)."""
    if not _word_installed():
        logger.info("DOCX→PDF skipped: Microsoft Word not installed — output is DOCX")
        return None
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            logger.info(f"DOCX → PDF: {pdf_path} ({pdf_path.stat().st_size // 1024} KB)")
            return pdf_path
        logger.warning("docx2pdf ran but PDF not found — returning DOCX")
        return None
    except Exception as exc:
        logger.warning(f"DOCX→PDF failed: {exc}")
        return None


# ── Code helpers ──────────────────────────────────────────────────────────────


def _strip_fences(text: str) -> str:
    """Remove markdown code fences that Claude sometimes wraps output in."""
    text = text.strip()
    text = re.sub(r"^```(?:python)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


_CODE_START_TOKENS = (
    "import ",
    "from ",
    "#",
    "def ",
    "class ",
    "OUTPUT_PATH",
    "DATA_FILE",
    "DATA",
    "doc ",
)


def _looks_like_prose(code: str) -> str | None:
    """
    Detect when Claude returns prose/explanation instead of Python code.
    Returns an error string if prose is detected, None if the response looks like code.
    """
    first = code.strip().split("\n")[0].strip()
    if not first:
        return "Empty response — no code generated"
    if not any(first.startswith(tok) for tok in _CODE_START_TOKENS):
        return f"Response looks like prose, not code (first line: {first[:80]!r})"
    return None


def _check_syntax(code: str) -> str | None:
    """
    Pre-validate generated code with Python's compile() before subprocess execution.
    Catches truncation issues (unclosed parens, unterminated strings) early
    and feeds the error back to Claude for a retry without wasting a subprocess call.
    """
    try:
        compile(code, "<generated>", "exec")
        return None
    except SyntaxError as exc:
        return f"SyntaxError line {exc.lineno}: {exc.msg} — {exc.text}"


def _run_script(script_path: Path) -> tuple[bool, str]:
    """
    Execute a generated python-docx script in a subprocess.
    Uses the same Python interpreter that's running this process.
    Returns (success, output_or_error_text).
    """
    try:
        result = subprocess.run(
            [sys.executable, str(script_path)],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode == 0:
            return True, result.stdout.strip() or "OK"
        combined = (result.stderr + "\n" + result.stdout).strip()
        return False, combined[:800]
    except subprocess.TimeoutExpired:
        return False, "Timeout after 60s"
    except Exception as exc:
        return False, str(exc)


# ── PAO per-page builder ──────────────────────────────────────────────────────


# ── PAO per-page builder ──────────────────────────────────────────────────────


class _PagePAO:
    """
    Runs a full Plan-Action-Observe loop for a single page.
    """

    def __init__(
        self,
        client: Any,
        workspace: Path,
        max_iterations: int,
    ) -> None:
        self.client = client
        self.workspace = workspace
        self.max_iterations = max_iterations

    def _img_block(self, path: Path) -> dict:
        return image_block_from_bytes(self.client, path.read_bytes())

    # ── Public entry point ────────────────────────────────────────────────────

    def run(
        self,
        page_num: int,
        page_fields: list[dict],
        img_path: Path,
        output_path: Path,
        target_language: str = "French",
    ) -> Path | None:
        """
        Full PAO loop for one page.
        Returns the output_path on success, None on failure.
        """
        data_file = (self.workspace / f"data_p{page_num:02d}.json").resolve()
        data_file.write_text(
            json.dumps(page_fields, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        logger.info(f"[P{page_num}] PLAN — generating layout plan …")
        plan = self._plan(page_num, img_path, page_fields)
        logger.debug(f"[P{page_num}] Plan:\n{plan[:300]}")

        prev_code = ""
        prev_error = ""
        observer_feedback = ""
        best_score = -1
        best_output: Path | None = None

        accept_threshold = 8

        for iteration in range(1, self.max_iterations + 1):
            logger.info(f"[P{page_num}] ACTION iteration {iteration}/{self.max_iterations}")

            # ── ACTION ───────────────────────────────────────────────────────
            code = self._generate_code(
                page_num,
                img_path,
                data_file,
                output_path,
                plan,
                prev_code,
                prev_error,
                observer_feedback,
            )

            prose_err = _looks_like_prose(code)
            if prose_err:
                logger.warning(f"[P{page_num}] Prose response: {prose_err[:80]}")
                prev_error = (
                    f"{prose_err}. Output ONLY Python code starting with 'import' or 'from'."
                )
                prev_code = ""
                continue

            syntax_err = _check_syntax(code)
            if syntax_err:
                logger.warning(f"[P{page_num}] Syntax error: {syntax_err[:80]}")
                prev_error = syntax_err
                prev_code = code
                continue

            script = (self.workspace / f"p{page_num:02d}_iter{iteration}.py").resolve()
            script.write_text(code, encoding="utf-8")
            success, run_output = _run_script(script)

            if not success:
                logger.warning(f"[P{page_num}] Runtime error: {run_output[:120]}")
                prev_error = run_output
                prev_code = code
                continue

            logger.info(f"[P{page_num}] Script OK — {run_output[:60]}")
            prev_error = ""
            prev_code = code

            # ── OBSERVE ──────────────────────────────────────────────────────
            logger.info(f"[P{page_num}] OBSERVE — checking output …")
            obs = self._observe(page_num, img_path, output_path, target_language, iteration)

            if obs is None:
                # Observer failed — keep best result, skip this iteration without new feedback
                logger.warning(f"[P{page_num}] Observer failed — keeping best result so far")
                continue

            score = obs.score
            feedback = obs.feedback
            missing = obs.missing_fields
            wrong = obs.wrong_values

            logger.info(
                f"[P{page_num}] Score={score}/10 | "
                f"missing={len(missing)} wrong={len(wrong)} | {feedback[:80]}"
            )

            # Track best result — return it if we never reach the threshold
            if score > best_score:
                best_score = score
                best_path = output_path.with_suffix(f".best{output_path.suffix}")
                shutil.copy2(output_path, best_path)
                best_output = best_path

            if obs.is_done or score >= accept_threshold:
                logger.info(
                    f"[P{page_num}] ✓ Accepted (score={score}, threshold={accept_threshold})"
                )
                return output_path

            # Regression guard: if score dropped with no concrete issues, revert to best
            if score < best_score and not missing and not wrong:
                logger.warning(
                    f"[P{page_num}] Score regression ({best_score}→{score}) with no specific issues — reverting to best"
                )
                if best_output and best_output.exists():
                    shutil.copy2(best_output, output_path)
                break

            # Build targeted feedback listing ALL missing keys so Claude can add them
            observer_feedback = (
                f"Score: {score}/10 (need {accept_threshold}+). Fix specifically:\n"
                f"Missing fields: {', '.join(missing)}\n"
                f"Wrong values: {', '.join(wrong)}\n"
                f"Feedback: {feedback}\n\n"
                f"IMPORTANT: Keep all working sections. Only ADD the missing fields above."
            )
            prev_error = ""

        logger.error(
            f"[P{page_num}] ✗ Failed after {self.max_iterations} iterations (best={best_score})"
        )
        if best_output and best_output.exists():
            shutil.copy2(best_output, output_path)
            logger.info(f"[P{page_num}] Using best result (score={best_score})")
        return output_path if output_path.exists() else None

    # ── PLAN ──────────────────────────────────────────────────────────────────

    def _plan(self, page_num: int, img_path: Path, fields: list[dict]) -> str:
        field_summary = json.dumps(
            [{"key": f["key"], "type": f.get("type", "text")} for f in fields],
            ensure_ascii=False,
        )
        try:
            plan = structured_call(
                self.client,
                _OPUS,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            self._img_block(img_path),
                            {
                                "type": "text",
                                "text": (
                                    f"PAGE {page_num} — {len(fields)} translated fields.\n"
                                    f"Fields (key + type): {field_summary}\n\n"
                                    f"Analyse the page image and fill the structured layout plan."
                                ),
                            },
                        ],
                    }
                ],
                schema=LayoutPlan,
                system=_PLANNER_SYSTEM,
                max_tokens=8192,
                reasoning_effort="medium",
            )
            return plan.model_dump_json(indent=2)
        except Exception as exc:
            logger.warning(f"[P{page_num}] Plan failed: {exc} — using minimal plan")
            return f"Page {page_num}: {len(fields)} fields in a form layout. Use a table-based structure."

    # ── ACTION ────────────────────────────────────────────────────────────────

    def _generate_code(
        self,
        page_num: int,
        img_path: Path,
        data_file: Path,
        output_path: Path,
        plan: str,
        prev_code: str,
        prev_error: str,
        observer_feedback: str,
    ) -> str:
        # Always use the main system prompt — full rewrite each time with feedback
        system = _CODER_SYSTEM

        # Pass ALL data without truncation — missing fields are caused by cut-off data
        data_raw = json.loads(data_file.read_text())
        data_sample = json.dumps(data_raw, ensure_ascii=False, indent=2)

        feedback_block = ""
        if prev_error:
            feedback_block += f"=== PREVIOUS ERROR TO FIX ===\n{prev_error[:800]}\n\n"
        if observer_feedback:
            feedback_block += f"=== OBSERVER FEEDBACK (rewrite the full script addressing this) ===\n{observer_feedback}\n\n"

        content: list[dict] = [
            self._img_block(img_path),
            {
                "type": "text",
                "text": (
                    f"{feedback_block}"
                    f"=== LAYOUT PLAN (JSON) ===\n{plan}\n\n"
                    f"The plan above defines: header, footer, layout, colors (hex per section), "
                    f"tables (rows/cols/widths/field keys), images, and non-table content.\n"
                    f"Follow it precisely — use the colors, table dimensions, and field assignments exactly.\n\n"
                    f"PAGE {page_num} — {len(data_raw)} fields\n"
                    f"OUTPUT_PATH = {str(output_path)!r}\n"
                    f"DATA_FILE = {str(data_file)!r}\n\n"
                    f"=== DATA (all {len(data_raw)} fields — use exact key strings verbatim) ===\n"
                    f"{data_sample}\n"
                    f"=== END DATA ===\n\n"
                    f"Write the complete python-docx script. "
                    f"Every field in DATA must appear. Every table in the plan must be built. "
                    f"Every color in the plan must be applied."
                ),
            },
        ]

        code = chat(
            self.client,
            _CODER,
            messages=[{"role": "user", "content": content}],
            system=system,
            max_tokens=16000,
            reasoning_effort="medium",
        )
        return _strip_fences(code)

    # ── OBSERVE ───────────────────────────────────────────────────────────────

    def _observe(
        self,
        page_num: int,
        img_path: Path,
        output_path: Path,
        target_language: str,
        iteration: int,  # noqa: ARG002
    ) -> ObserverResult | None:
        if not output_path.exists():
            return ObserverResult(
                score=1,
                is_done=False,
                feedback="Output file does not exist.",
            )

        docx_text = _dump_docx_text(output_path)

        content: list[dict] = [
            self._img_block(img_path),
            {
                "type": "text",
                "text": (
                    f"TARGET LANGUAGE: {target_language} "
                    f"(all text in the DOCX should be in {target_language} — this is CORRECT)\n\n"
                    f"=== ORIGINAL PAGE IMAGE (source language, shown above) ===\n\n"
                    f"=== GENERATED DOCX CONTENT (translated to {target_language}) ===\n"
                    f"{docx_text[:3000]}\n\n"
                    f"Score the reconstruction quality. "
                    f"Check that all fields from the original appear (translated) in the output. "
                    f"Penalise empty fields and missing data. "
                    f"Do NOT penalise for the language being {target_language}."
                ),
            },
        ]

        try:
            return structured_call(
                self.client,
                _SONNET,
                messages=[{"role": "user", "content": content}],
                schema=ObserverResult,
                system=_OBSERVER_SYSTEM,
                max_tokens=4096,
                reasoning_effort="low",
            )
        except Exception as exc:
            logger.warning(f"[P{page_num}] Observer error: {exc}")
            return None


# ── Main agent ────────────────────────────────────────────────────────────────


class PAOCodingAgent:
    """
    Builds a translated DOCX using a per-page Plan→Action→Observe loop,
    running all pages in parallel, then merging into one final document.
    """

    def __init__(self) -> None:
        self.client = get_client()
        self.workspace = settings.agent_workspace.resolve()
        self.workspace.mkdir(parents=True, exist_ok=True)
        self.max_iterations = settings.max_agent_iterations
        logger.info(
            f"PAOCodingAgent ready (provider={settings.llm_provider}, model={settings.model}, "
            f"heavy={settings.heavy_model}, workspace={self.workspace}, max_iter={self.max_iterations})"
        )

    def run(
        self,
        translated_doc: TranslatedDocument,
        page_image_paths: list[Path],
        output_docx: Path,
    ) -> Path:
        output_docx = output_docx.resolve()
        output_docx.parent.mkdir(parents=True, exist_ok=True)

        # Per-PDF workspace subfolder keeps runs isolated and easy to inspect
        pdf_stem = output_docx.stem.replace("_translated", "") or "job"
        job_workspace = self.workspace / pdf_stem
        job_workspace.mkdir(parents=True, exist_ok=True)

        target_language = translated_doc.target_language
        total_fields = sum(len(p.translated_fields) for p in translated_doc.pages)
        logger.info(
            f"=== PAO CODING AGENT START: {len(translated_doc.pages)} pages, "
            f"{total_fields} fields → {target_language}, max_iter={self.max_iterations} ==="
        )
        logger.info(f"  workspace: {job_workspace}")

        img_map = {i + 1: p for i, p in enumerate(page_image_paths)}
        page_results: dict[int, Path | None] = {}

        # ── Run PAO loop per page in parallel (up to 3 at once) ──────────────
        with ThreadPoolExecutor(max_workers=3) as pool:
            futures: dict[Any, int] = {}

            for tpage in translated_doc.pages:
                page_num = tpage.page_number
                fields = self._page_fields(tpage)
                img_path = img_map.get(page_num, page_image_paths[0])
                page_out = (job_workspace / f"page_{page_num:02d}.docx").resolve()

                # Complex pages (>40 fields) get extra iterations automatically
                page_max_iter = self.max_iterations + (2 if len(fields) > 40 else 0)
                pao = _PagePAO(self.client, job_workspace, page_max_iter)
                console.print(
                    f"[dim]Queuing page {page_num} ({len(fields)} fields, "
                    f"up to {page_max_iter} PAO iterations)…[/dim]"
                )
                f = pool.submit(pao.run, page_num, fields, img_path, page_out, target_language)
                futures[f] = page_num

            for future in as_completed(futures):
                page_num = futures[future]
                try:
                    result = future.result()
                    page_results[page_num] = result
                    status = "✓" if result else "✗"
                    color = "green" if result else "red"
                    console.print(f"[{color}]{status} Page {page_num} done[/]")
                except Exception as exc:
                    logger.error(f"Page {page_num} raised exception: {exc}")
                    page_results[page_num] = None

        # ── Merge pages in order ──────────────────────────────────────────────
        ordered = [
            page_results.get(tpage.page_number)
            for tpage in translated_doc.pages
            if page_results.get(tpage.page_number)
        ]

        if not ordered:
            logger.error("No pages were built — cannot create DOCX")
            return output_docx

        n_success = len(ordered)
        n_total = len(translated_doc.pages)
        console.print(
            Panel(
                f"Merging {n_success}/{n_total} pages…",
                title="[PAO Coding Agent]",
                expand=False,
            )
        )

        merged = _merge_docx_pages(ordered)
        merged.save(str(output_docx))
        logger.info(f"DOCX saved: {output_docx} ({output_docx.stat().st_size // 1024} KB)")

        # ── Convert to PDF ────────────────────────────────────────────────────
        pdf = _convert_docx_to_pdf(output_docx)
        logger.info("=== PAO CODING AGENT DONE ===")
        return pdf if pdf else output_docx

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _page_fields(tpage: TranslatedPage) -> list[dict]:
        return [
            {
                "key": tf.translated_key,
                "value": tf.translated_value,
                "original_key": tf.original.key,
                "original_value": tf.original.value,
                "type": tf.original.field_type,
            }
            for tf in tpage.translated_fields
        ]
